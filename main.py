import os
os.environ["STREAMLIT_WATCHDOG"] = "false"
import json
import random
import string
import re
import asyncio
import io
import urllib.parse
import base64
from io import BytesIO
from collections import Counter
from datetime import datetime
import time

# Third-party library imports
import streamlit as st
import streamlit.components.v1 as components
from base64 import b64encode
import requests
import fitz
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import altair as alt
from PIL import Image
from pdf2image import convert_from_path
from dotenv import load_dotenv
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from xhtml2pdf import pisa
from pydantic import BaseModel
from streamlit_pdf_viewer import pdf_viewer

# Heavy libraries - loaded with caching
import torch

# Langchain & Embeddings

from langchain_text_splitters import CharacterTextSplitter 
from langchain_community.vectorstores import FAISS 
from langchain_community.embeddings import HuggingFaceEmbeddings 
from langchain_groq import ChatGroq  # optional if you're using it













# Local project imports
from llm_manager import call_llm, load_groq_api_keys
from db_manager import (
    db_manager,
    insert_candidate,
    get_top_domains_by_score,
    get_database_stats,
    detect_domain_from_title_and_description,
    get_domain_similarity
)
from user_login import (
    create_user_table,
    add_user,
    complete_registration,
    verify_user,
    get_logins_today,
    get_total_registered_users,
    log_user_action,
    username_exists,
    email_exists,
    is_valid_email,
    save_user_api_key,
    get_user_api_key,
    get_all_user_logs,
    generate_otp,
    send_email_otp,
    get_user_by_email,
    update_password_by_email,
    is_strong_password,
    domain_has_mx_record
)

# ============================================================
# 💾 Persistent Storage Configuration for Streamlit Cloud
# ============================================================
os.makedirs(".streamlit_storage", exist_ok=True)
DB_PATH = os.path.join(".streamlit_storage", "resume_data.db")

def html_to_pdf_bytes(html_string):
    styled_html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: 400mm 297mm;  /* Original custom large page size */
                margin-top: 10mm;
                margin-bottom: 10mm;
                margin-left: 10mm;
                margin-right: 10mm;
            }}
            body {{
                font-size: 14pt;
                font-family: "Segoe UI", "Helvetica", sans-serif;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2, h3 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 10px;
            }}
            .box {{
                padding: 8px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;  /* More elegant than full border */
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.5em;
            }}
            li {{
                margin-bottom: 5px;
            }}
        </style>
    </head>
    <body>
        {html_string}
    </body>
    </html>
    """

    pdf_io = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_io)
    pdf_io.seek(0)
    return pdf_io

def generate_cover_letter_from_resume_builder():
    name = st.session_state.get("name", "")
    job_title = st.session_state.get("job_title", "")
    summary = st.session_state.get("summary", "")
    skills = st.session_state.get("skills", "")
    location = st.session_state.get("location", "")
    today_date = datetime.today().strftime("%B %d, %Y")

    # ✅ Input boxes for contact info
    company = st.text_input("🏢 Target Company", placeholder="e.g., Google")
    linkedin = st.text_input("🔗 LinkedIn URL", placeholder="e.g., https://linkedin.com/in/username")
    email = st.text_input("📧 Email", placeholder="e.g., you@example.com")
    mobile = st.text_input("📞 Mobile Number", placeholder="e.g., +91 9876543210")

    # ✅ Button to prevent relooping
    if st.button("✉️ Generate Cover Letter"):
        # ✅ Validate input before generating
        if not all([name, job_title, summary, skills, company, linkedin, email, mobile]):
            st.warning("⚠️ Please fill in all fields including LinkedIn, email, and mobile.")
            return

        prompt = f"""
You are a professional cover letter writer.

Write a formal and compelling cover letter using the information below. 
Format it as a real letter with:
1. Date
2. Recipient heading
3. Proper salutation
4. Three short paragraphs
5. Professional closing

Ensure you **only include the company name once** in the header or salutation, 
and avoid repeating it redundantly in the body.

### Heading Info:
{today_date}
Hiring Manager, {company}, {location}

### Candidate Info:
- Name: {name}
- Job Title: {job_title}
- Summary: {summary}
- Skills: {skills}
- Location: {location}

### Instructions:
- Do not use HTML tags. 
- Return plain text only.
"""

        # ✅ Call LLM
        cover_letter = call_llm(prompt, session=st.session_state).strip()

        # ✅ Store plain text
        st.session_state["cover_letter"] = cover_letter

        # ✅ Build HTML wrapper for preview (safe)
        cover_letter_html = f"""
        <div style="font-family: Georgia, serif; font-size: 13pt; line-height: 1.6; 
                    color: #000; background: #fff; padding: 25px; 
                    border-radius: 8px; box-shadow: 0px 2px 6px rgba(0,0,0,0.1); 
                    max-width: 800px; margin: auto;">
            <div style="text-align:center; margin-bottom:15px;">
                <div style="font-size:18pt; font-weight:bold; color:#003366;">{name}</div>
                <div style="font-size:14pt; color:#555;">{job_title}</div>
                <div style="font-size:10pt; margin-top:5px;">
                    <a href="{linkedin}" style="color:#003366;">{linkedin}</a><br/>
                    📧 {email} | 📞 {mobile}
                </div>
            </div>
            <hr/>
            <pre style="white-space: pre-wrap; font-family: Georgia, serif; font-size: 12pt; color:#000;">
{cover_letter}
            </pre>
        </div>
        """

        st.session_state["cover_letter_html"] = cover_letter_html

        # ✅ Show nicely in Streamlit
        st.markdown(cover_letter_html, unsafe_allow_html=True)

# ------------------- Initialize -------------------
# ✅ Initialize database in persistent storage
create_user_table()

# ------------------- Tab-Specific Notification System -------------------
if "login_notification" not in st.session_state:
    st.session_state.login_notification = {"type": None, "text": None, "expires": 0.0}
if "register_notification" not in st.session_state:
    st.session_state.register_notification = {"type": None, "text": None, "expires": 0.0}

def notify(tab, msg_type, text, duration=3.0):
    """Show auto-disappearing message for specific tab (login/register)."""
    notification_key = f"{tab}_notification"
    st.session_state[notification_key] = {
        "type": msg_type,
        "text": text,
        "expires": time.time() + duration,
    }

def render_notification(tab):
    """Render notification in a fixed center slot for specific tab (prevents button shifting)."""
    notification_key = f"{tab}_notification"
    notif = st.session_state[notification_key]

    # Always reserve space for notification (60px height)
    if notif["type"] and time.time() < notif["expires"]:
        # Show active notification
        if notif["type"] == "success":
            st.success(notif["text"])
        elif notif["type"] == "error":
            st.error(notif["text"])
        elif notif["type"] == "warning":
            st.warning(notif["text"])
        elif notif["type"] == "info":
            st.info(notif["text"])
    else:
        # Reserve space with empty div to prevent layout shift
        st.markdown("<div style='height:60px;'></div>", unsafe_allow_html=True)

def display_timer(remaining_seconds, expired=False, key_suffix=""):
    """
    Display a server-synced timer with glassmorphism styling.
    Server-side validation ensures OTP expiry is accurately enforced.

    Args:
        remaining_seconds: Time remaining in seconds (server-calculated)
        expired: Whether the timer has expired
        key_suffix: Unique suffix for the timer component
    """
    minutes = remaining_seconds // 60
    seconds = remaining_seconds % 60

    if expired or remaining_seconds <= 0:
        st.markdown("""
        <div class='timer-display timer-expired' style="
            background: linear-gradient(135deg, rgba(255, 99, 71, 0.18) 0%, rgba(255, 99, 71, 0.08) 100%);
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border: 2px solid rgba(255, 99, 71, 0.4);
            border-radius: 14px;
            padding: 16px 24px;
            margin: 20px 0;
            text-align: center;
            box-shadow: 0 4px 20px rgba(255, 99, 71, 0.15), inset 0 1px 0 rgba(255, 255, 255, 0.1);
        ">
            <span class='timer-text' style="
                color: #FF6347;
                font-size: 1.15em;
                font-weight: bold;
                font-family: 'Orbitron', sans-serif;
                text-shadow: 0 0 18px rgba(255, 99, 71, 0.5);
            ">⏱️ OTP Expired</span>
        </div>
        """, unsafe_allow_html=True)
    else:
        # Client-side countdown for UX, but server validates on action
        st.components.v1.html(f"""
        <div class='timer-display' id='timer-{key_suffix}' style="
            background: linear-gradient(135deg, rgba(255, 215, 0, 0.18) 0%, rgba(255, 165, 0, 0.08) 100%);
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border: 2px solid rgba(255, 215, 0, 0.4);
            border-radius: 14px;
            padding: 16px 24px;
            margin: 20px 0;
            text-align: center;
            box-shadow: 0 4px 20px rgba(255, 215, 0, 0.15), inset 0 1px 0 rgba(255, 255, 255, 0.1);
        ">
            <span class='timer-text' style="
                color: #FFD700;
                font-size: 1.15em;
                font-weight: bold;
                font-family: 'Orbitron', sans-serif;
                text-shadow: 0 0 18px rgba(255, 215, 0, 0.5);
            ">⏱️ Time Remaining: <span id='countdown-{key_suffix}'>{minutes:02d}:{seconds:02d}</span></span>
        </div>
        <script>
        (function() {{
            let remaining = {remaining_seconds};
            const countdownEl = document.getElementById('countdown-{key_suffix}');
            const timerEl = document.getElementById('timer-{key_suffix}');

            const interval = setInterval(() => {{
                remaining--;
                if (remaining <= 0) {{
                    clearInterval(interval);
                    if (timerEl) {{
                        timerEl.style.background = 'linear-gradient(135deg, rgba(255, 99, 71, 0.18) 0%, rgba(255, 99, 71, 0.08) 100%)';
                        timerEl.style.border = '2px solid rgba(255, 99, 71, 0.4)';
                        timerEl.innerHTML = "<span style='color: #FF6347; font-size: 1.15em; font-weight: bold; font-family: Orbitron, sans-serif; text-shadow: 0 0 18px rgba(255, 99, 71, 0.5);'>⏱️ OTP Expired</span>";
                    }}
                }} else {{
                    const mins = Math.floor(remaining / 60);
                    const secs = remaining % 60;
                    if (countdownEl) {{
                        countdownEl.textContent = `${{mins.toString().padStart(2, '0')}}:${{secs.toString().padStart(2, '0')}}`;
                    }}
                }}
            }}, 1000);
        }})();
        </script>
        """, height=80)

# ------------------- Initialize Session State -------------------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "username" not in st.session_state:
    st.session_state.username = None
if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

# Forgot password session states
if "reset_stage" not in st.session_state:
    st.session_state.reset_stage = "none"
if "reset_email" not in st.session_state:
    st.session_state.reset_email = ""
if "reset_otp" not in st.session_state:
    st.session_state.reset_otp = ""
if "reset_otp_time" not in st.session_state:
    st.session_state.reset_otp_time = 0

# Live validation session states for register tab
if "last_validated_email" not in st.session_state:
    st.session_state.last_validated_email = ""
if "last_validated_username" not in st.session_state:
    st.session_state.last_validated_username = ""
if "last_validated_password" not in st.session_state:
    st.session_state.last_validated_password = ""

# ------------------- CSS Styling -------------------
st.markdown("""
<style>
body, .main {
    background-color: #0d1117;
    color: white;
}

/* Smooth fade animation for notifications */
div.stAlert {
    border-radius: 12px;
    padding: 10px 14px;
    animation: fadein 0.3s, fadeout 0.3s 2.7s;
    text-align: center;
}
@keyframes fadein { from {opacity: 0;} to {opacity: 1;} }
@keyframes fadeout { from {opacity: 1;} to {opacity: 0;} }

.login-card {
    background: #161b22;
    padding: 30px;
    border-radius: 20px;
    box-shadow: 0 0 25px rgba(0,0,0,0.3);
    transition: all 0.4s ease;
}
.login-card:hover {
    transform: translateY(-6px) scale(1.01);
    box-shadow: 0 0 45px rgba(0,255,255,0.25);
}
.stTextInput > div > input {
    background-color: #0d1117;
    color: white;
    border: 1px solid #30363d;
    border-radius: 10px;
    padding: 0.6em;
}
.stTextInput > div > input:hover {
    border: 1px solid #00BFFF;
    box-shadow: 0 0 8px rgba(0,191,255,0.2);
}
.stTextInput > label {
    color: #c9d1d9;
}
.stButton > button {
    background-color: #238636;
    color: white;
    border-radius: 10px;
    padding: 0.6em 1.5em;
    border: none;
    font-weight: bold;
}
.stButton > button:hover {
    background-color: #2ea043;
    box-shadow: 0 0 10px rgba(46,160,67,0.4);
    transform: scale(1.02);
}
.feature-card {
    background: radial-gradient(circle at top left, #1f2937, #111827);
    padding: 20px;
    border-radius: 15px;
    box-shadow: 0 0 20px rgba(0,255,255,0.1);
    text-align: center;
    transition: transform 0.3s ease, box-shadow 0.3s ease;
    color: #fff;
    margin-bottom: 20px;
}
.feature-card:hover {
    transform: translateY(-10px);
    box-shadow: 0 0 30px rgba(0,255,255,0.4);
}
.feature-card h3 {
    color: #00BFFF;
}
.feature-card p {
    color: #c9d1d9;
}
</style>
""", unsafe_allow_html=True)
# 🔹 VIDEO BACKGROUND & GLOW TEXT

# ------------------- BEFORE LOGIN -------------------
if not st.session_state.authenticated:
    

    # -------- Sidebar --------
    with st.sidebar:
        st.markdown("<h1 style='color:#00BFFF;'>Smart Resume AI</h1>", unsafe_allow_html=True)
        st.markdown("<p style='color:#c9d1d9;'>Transform your career with AI-powered resume analysis, job matching, and smart insights.</p>", unsafe_allow_html=True)

        features = [
            ("https://img.icons8.com/fluency/48/resume.png", "Resume Analyzer", "Get feedback, scores, and tips powered by AI along with the biased words detection and rewriting the resume in an inclusive way."),
            ("https://img.icons8.com/fluency/48/resume-website.png", "Resume Builder", "Build modern, eye-catching resumes easily."),
            ("https://img.icons8.com/fluency/48/job.png", "Job Search", "Find tailored job matches."),
            ("https://img.icons8.com/fluency/48/classroom.png", "Course Suggestions", "Get upskilling recommendations based on your goals."),
            ("https://img.icons8.com/fluency/48/combo-chart.png", "Interactive Dashboard", "Visualize trends, scores, and analytics."),
        ]

        for icon, title, desc in features:
            st.markdown(f"""
            <div class="feature-card">
                <img src="{icon}" width="40"/>
                <h3>{title}</h3>
                <p>{desc}</p>
            </div>
            """, unsafe_allow_html=True)

    # -------- Animated Cards --------
    image_url = "https://cdn-icons-png.flaticon.com/512/3135/3135768.png"
    response = requests.get(image_url)
    img_base64 = b64encode(response.content).decode()

    st.markdown(f"""
    <style>
    .animated-cards {{
      margin-top: 30px;
      display: flex;
      justify-content: center;
      position: relative;
      height: 300px;
    }}
    .animated-cards img {{
      position: absolute;
      width: 240px;
      animation: splitCards 2.5s ease-in-out infinite alternate;
      z-index: 1;
    }}
    .animated-cards img:nth-child(1) {{ animation-delay: 0s; z-index: 3; }}
    .animated-cards img:nth-child(2) {{ animation-delay: 0.3s; z-index: 2; }}
    .animated-cards img:nth-child(3) {{ animation-delay: 0.6s; z-index: 1; }}
    @keyframes splitCards {{
      0% {{ transform: scale(1) translateX(0) rotate(0deg); opacity: 1; }}
      100% {{ transform: scale(1) translateX(var(--x-offset)) rotate(var(--rot)); opacity: 1; }}
    }}
    .card-left {{ --x-offset: -80px; --rot: -5deg; }}
    .card-center {{ --x-offset: 0px; --rot: 0deg; }}
    .card-right {{ --x-offset: 80px; --rot: 5deg; }}
    </style>
    <div class="animated-cards">
        <img class="card-left" src="data:image/png;base64,{img_base64}" />
        <img class="card-center" src="data:image/png;base64,{img_base64}" />
        <img class="card-right" src="data:image/png;base64,{img_base64}" />
    </div>
    """, unsafe_allow_html=True)

    # -------- Counter Section (Updated Layout & Style with glassmorphism and shimmer) --------

    # Fetch counters
    total_users = get_total_registered_users()
    active_logins = get_logins_today()
    stats = get_database_stats()

# Replace static 15 with dynamic count
    resumes_uploaded = stats.get("total_candidates", 0)

    states_accessed = 29

    glassmorphism_counter_style = """
    <style>
    @keyframes shimmer {
        0% { background-position: -200% 0; }
        100% { background-position: 200% 0; }
    }
    
    @keyframes float {
        0%, 100% { transform: translateY(0px); }
        50% { transform: translateY(-5px); }
    }

    .counter-grid {
        display: grid;
        grid-template-columns: repeat(2, 250px);
        column-gap: 40px;
        row-gap: 25px;
        justify-content: center;
        padding: 30px 10px;
        max-width: 600px;
        margin: 0 auto;
    }

    .counter-box {
        background: linear-gradient(135deg, 
            rgba(0, 191, 255, 0.1) 0%, 
            rgba(30, 144, 255, 0.05) 50%, 
            rgba(0, 191, 255, 0.1) 100%);
        backdrop-filter: blur(15px);
        -webkit-backdrop-filter: blur(15px);
        border: 1px solid rgba(0, 191, 255, 0.2);
        border-radius: 16px;
        width: 100%;
        height: 120px;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        position: relative;
        overflow: hidden;
        transition: all 0.3s ease;
        animation: float 3s ease-in-out infinite;
    }

    .counter-box::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(
            90deg,
            transparent,
            rgba(0, 191, 255, 0.3),
            transparent
        );
        animation: shimmer 2s infinite;
    }

    .counter-box:hover {
        transform: translateY(-8px) scale(1.02);
        background: linear-gradient(135deg, 
            rgba(0, 191, 255, 0.15) 0%, 
            rgba(30, 144, 255, 0.08) 50%, 
            rgba(0, 191, 255, 0.15) 100%);
        border: 1px solid rgba(0, 191, 255, 0.4);
        box-shadow: 
            0 20px 40px rgba(0, 191, 255, 0.1),
            inset 0 1px 0 rgba(255, 255, 255, 0.1);
    }

    .counter-box:nth-child(1) { animation-delay: 0s; }
    .counter-box:nth-child(2) { animation-delay: 0.5s; }
    .counter-box:nth-child(3) { animation-delay: 1s; }
    .counter-box:nth-child(4) { animation-delay: 1.5s; }

    .counter-number {
        font-size: 2.2em;
        font-weight: bold;
        color: #00BFFF;
        margin: 0;
        position: relative;
        z-index: 2;
        text-shadow: 0 0 20px rgba(0, 191, 255, 0.5);
    }

    .counter-label {
        margin-top: 8px;
        font-size: 1em;
        color: #c9d1d9;
        position: relative;
        z-index: 2;
        text-shadow: 0 1px 2px rgba(0, 0, 0, 0.3);
    }
    </style>
    """

    st.markdown(glassmorphism_counter_style, unsafe_allow_html=True)

    st.markdown(f"""
    <div class="counter-grid">
        <div class="counter-box">
            <div class="counter-number">{total_users}</div>
            <div class="counter-label">Total Users</div>
        </div>
        <div class="counter-box">
            <div class="counter-number">{states_accessed}</div>
            <div class="counter-label">States Accessed</div>
        </div>
        <div class="counter-box">
            <div class="counter-number">{resumes_uploaded}</div>
            <div class="counter-label">Resumes Uploaded</div>
        </div>
        <div class="counter-box">
            <div class="counter-number">{active_logins}</div>
            <div class="counter-label">Active Sessions</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

if not st.session_state.get("authenticated", False):

    # ✅ Futuristic silhouette
    image_url = "https://cdn-icons-png.flaticon.com/512/4140/4140047.png"
    response = requests.get(image_url)
    img_base64 = b64encode(response.content).decode()

    # ✅ Inject glassmorphism CSS with shimmer effects
    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@600&display=swap');

    @keyframes shimmer {{
        0% {{ background-position: -200% 0; }}
        100% {{ background-position: 200% 0; }}
    }}

    @keyframes glassShimmer {{
        0% {{ transform: translateX(-100%) skewX(-15deg); }}
        100% {{ transform: translateX(200%) skewX(-15deg); }}
    }}

    /* ===== Card Shuffle Animation ===== */
    .animated-cards {{
      margin-top: 40px;
      display: flex;
      justify-content: center;
      position: relative;
      height: 260px;
    }}
    .animated-cards img {{
      position: absolute;
      width: 220px;
      animation: splitCards 2.5s ease-in-out infinite alternate;
      z-index: 1;
      filter: drop-shadow(0 0 15px rgba(0,191,255,0.3));
    }}
    .animated-cards img:nth-child(1) {{ animation-delay: 0s; z-index: 3; }}
    .animated-cards img:nth-child(2) {{ animation-delay: 0.3s; z-index: 2; }}
    .animated-cards img:nth-child(3) {{ animation-delay: 0.6s; z-index: 1; }}

    @keyframes splitCards {{
      0%   {{ transform: scale(1) translateX(0) rotate(0deg); opacity: 1; }}
      100% {{ transform: scale(1) translateX(var(--x-offset)) rotate(var(--rot)); opacity: 1; }}
    }}
    .card-left   {{ --x-offset: -80px; --rot: -4deg; }}
    .card-center {{ --x-offset: 0px;  --rot: 0deg;  }}
    .card-right  {{ --x-offset: 80px;  --rot: 4deg;  }}

    /* ===== Glassmorphism Login Card ===== */
    .login-card {{
      background: linear-gradient(135deg,
        rgba(0, 191, 255, 0.1) 0%,
        rgba(30, 144, 255, 0.05) 50%,
        rgba(0, 191, 255, 0.1) 100%);
      backdrop-filter: blur(20px);
      -webkit-backdrop-filter: blur(20px);
      border: 1px solid rgba(0, 191, 255, 0.2);
      border-radius: 20px;
      padding: 25px;
      box-shadow:
        0 8px 32px rgba(0, 191, 255, 0.1),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      font-family: 'Orbitron', sans-serif;
      color: white;
      margin-top: 20px;
      opacity: 0;
      transform: translateX(-120%);
      animation: slideInLeft 1.2s ease-out forwards;
      position: relative;
      overflow: hidden;
    }}

    .login-card::before {{
      content: '';
      position: absolute;
      top: 0;
      left: -100%;
      width: 100%;
      height: 100%;
      background: linear-gradient(
        90deg,
        transparent,
        rgba(0, 191, 255, 0.2),
        transparent
      );
      animation: glassShimmer 3s infinite;
    }}

    @keyframes slideInLeft {{
      0%   {{ transform: translateX(-120%); opacity: 0; }}
      100% {{ transform: translateX(0); opacity: 1; }}
    }}

    .login-card h2 {{
      text-align: center;
      font-size: 1.6rem;
      text-shadow: 0 0 15px rgba(0, 191, 255, 0.5);
      margin-bottom: 15px;
      position: relative;
      z-index: 2;
    }}
    .login-card h2 span {{ color: #00BFFF; }}

    /* ===== Enhanced Message Cards with Consistent Layout ===== */
    .slide-message {{
      position: relative;
      overflow: hidden;
      margin: 16px 0;
      padding: 14px 20px;
      border-radius: 14px;
      font-weight: 600;
      font-size: 0.95em;
      display: flex;
      align-items: center;
      justify-content: flex-start;
      gap: 12px;
      animation: slideIn 0.6s cubic-bezier(0.34, 1.56, 0.64, 1) forwards;
      backdrop-filter: blur(15px);
      -webkit-backdrop-filter: blur(15px);
      box-shadow:
        0 4px 20px rgba(0, 0, 0, 0.15),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      width: 100%;
      max-width: 100%;
      box-sizing: border-box;
      line-height: 1.5;
      font-family: 'Orbitron', sans-serif;
      transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
      min-height: 50px;
    }}

    .slide-message:hover {{
      transform: translateY(-3px) scale(1.01);
      box-shadow:
        0 8px 30px rgba(0, 0, 0, 0.25),
        inset 0 1px 0 rgba(255, 255, 255, 0.15);
    }}

    .slide-message::before {{
      content: '';
      position: absolute;
      top: 0;
      left: -100%;
      width: 100%;
      height: 100%;
      background: linear-gradient(
        90deg,
        transparent,
        rgba(255, 255, 255, 0.1),
        transparent
      );
      transition: left 0.5s;
    }}

    .slide-message:hover::before {{
      left: 100%;
    }}

    .slide-message svg {{
      width: 22px;
      height: 22px;
      flex-shrink: 0;
      filter: drop-shadow(0 0 6px currentColor);
      z-index: 2;
    }}

    .slide-message-text {{
      flex: 1;
      z-index: 2;
      position: relative;
      word-wrap: break-word;
      overflow-wrap: break-word;
      white-space: normal;
    }}

    .success-msg {{
      background: linear-gradient(135deg,
        rgba(0, 255, 127, 0.20) 0%,
        rgba(0, 255, 127, 0.08) 100%);
      border: 2px solid rgba(0, 255, 127, 0.4);
      color: #00FF7F;
      text-shadow: 0 0 12px rgba(0, 255, 127, 0.4);
    }}

    .error-msg {{
      background: linear-gradient(135deg,
        rgba(255, 99, 71, 0.20) 0%,
        rgba(255, 99, 71, 0.08) 100%);
      border: 2px solid rgba(255, 99, 71, 0.4);
      color: #FF6347;
      text-shadow: 0 0 12px rgba(255, 99, 71, 0.4);
    }}

    .info-msg {{
      background: linear-gradient(135deg,
        rgba(30, 144, 255, 0.20) 0%,
        rgba(30, 144, 255, 0.08) 100%);
      border: 2px solid rgba(30, 144, 255, 0.4);
      color: #1E90FF;
      text-shadow: 0 0 12px rgba(30, 144, 255, 0.4);
    }}

    .warn-msg {{
      background: linear-gradient(135deg,
        rgba(255, 215, 0, 0.20) 0%,
        rgba(255, 215, 0, 0.08) 100%);
      border: 2px solid rgba(255, 215, 0, 0.4);
      color: #FFD700;
      text-shadow: 0 0 12px rgba(255, 215, 0, 0.4);
    }}

    @keyframes slideIn {{
      0%   {{
        transform: translateX(-50px);
        opacity: 0;
      }}
      100% {{
        transform: translateX(0);
        opacity: 1;
      }}
    }}

    /* ===== Improved Timer Display ===== */
    .timer-display {{
      background: linear-gradient(135deg,
        rgba(255, 215, 0, 0.18) 0%,
        rgba(255, 165, 0, 0.08) 100%);
      backdrop-filter: blur(15px);
      -webkit-backdrop-filter: blur(15px);
      border: 2px solid rgba(255, 215, 0, 0.4);
      border-radius: 14px;
      padding: 16px 24px;
      margin: 20px 0;
      text-align: center;
      box-shadow:
        0 4px 20px rgba(255, 215, 0, 0.15),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
      position: relative;
      overflow: hidden;
    }}

    .timer-display::before {{
      content: '';
      position: absolute;
      top: 0;
      left: -100%;
      width: 100%;
      height: 100%;
      background: linear-gradient(
        90deg,
        transparent,
        rgba(255, 215, 0, 0.2),
        transparent
      );
      animation: glassShimmer 3s infinite;
    }}

    .timer-display:hover {{
      box-shadow:
        0 8px 30px rgba(255, 215, 0, 0.25),
        inset 0 1px 0 rgba(255, 255, 255, 0.15);
      transform: translateY(-3px);
    }}

    .timer-text {{
      color: #FFD700;
      font-size: 1.15em;
      font-weight: bold;
      font-family: 'Orbitron', sans-serif;
      text-shadow: 0 0 18px rgba(255, 215, 0, 0.5);
      position: relative;
      z-index: 2;
    }}

    .timer-expired {{
      background: linear-gradient(135deg,
        rgba(255, 99, 71, 0.18) 0%,
        rgba(255, 99, 71, 0.08) 100%);
      border: 2px solid rgba(255, 99, 71, 0.4);
    }}

    .timer-expired .timer-text {{
      color: #FF6347;
      text-shadow: 0 0 18px rgba(255, 99, 71, 0.5);
    }}

    /* ===== Glassmorphism Buttons ===== */
    .stButton>button {{
      background: linear-gradient(135deg, 
        rgba(0, 191, 255, 0.2) 0%, 
        rgba(30, 144, 255, 0.1) 100%);
      backdrop-filter: blur(15px);
      -webkit-backdrop-filter: blur(15px);
      color: white;
      border: 1px solid rgba(0, 191, 255, 0.3);
      border-radius: 12px;
      font-family: 'Orbitron', sans-serif;
      font-weight: bold;
      padding: 8px 20px;
      box-shadow: 
        0 4px 16px rgba(0, 191, 255, 0.1),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      transition: all 0.3s ease;
      position: relative;
      overflow: hidden;
    }}
    
    .stButton>button::before {{
      content: '';
      position: absolute;
      top: 0;
      left: -100%;
      width: 100%;
      height: 100%;
      background: linear-gradient(
        90deg,
        transparent,
        rgba(255, 255, 255, 0.2),
        transparent
      );
      transition: left 0.5s;
    }}
    
    .stButton>button:hover {{
      transform: translateY(-2px);
      background: linear-gradient(135deg, 
        rgba(0, 191, 255, 0.3) 0%, 
        rgba(30, 144, 255, 0.15) 100%);
      border: 1px solid rgba(0, 191, 255, 0.5);
      box-shadow: 
        0 8px 25px rgba(0, 191, 255, 0.2),
        inset 0 1px 0 rgba(255, 255, 255, 0.2);
    }}
    
    .stButton>button:hover::before {{
      left: 100%;
    }}

    /* ===== Glassmorphism Input Fields ===== */
    .stTextInput input {{
      background: linear-gradient(135deg, 
        rgba(0, 191, 255, 0.08) 0%, 
        rgba(30, 144, 255, 0.04) 100%);
      backdrop-filter: blur(15px);
      -webkit-backdrop-filter: blur(15px);
      border: 1px solid rgba(0, 191, 255, 0.2);
      border-radius: 10px;
      padding: 10px;
      color: #E0F7FF;
      font-family: 'Orbitron', sans-serif;
      box-shadow: 
        0 4px 16px rgba(0, 191, 255, 0.05),
        inset 0 1px 0 rgba(255, 255, 255, 0.05);
      transition: all 0.3s ease-in-out;
    }}
    .stTextInput input:focus {{
      outline: none !important;
      background: linear-gradient(135deg, 
        rgba(0, 191, 255, 0.12) 0%, 
        rgba(30, 144, 255, 0.06) 100%);
      border: 1px solid rgba(0, 191, 255, 0.4);
      box-shadow: 
        0 8px 25px rgba(0, 191, 255, 0.15),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      transform: translateY(-1px);
    }}
    .stTextInput label {{
      font-family: 'Orbitron', sans-serif;
      color: #00BFFF !important;
      text-shadow: 0 0 10px rgba(0, 191, 255, 0.3);
    }}
    </style>

    <!-- Animated Cards -->
    <div class="animated-cards">
        <img class="card-left" src="data:image/png;base64,{img_base64}" />
        <img class="card-center" src="data:image/png;base64,{img_base64}" />
        <img class="card-right" src="data:image/png;base64,{img_base64}" />
    </div>
    """, unsafe_allow_html=True)

    # -------- Login/Register Layout --------
    left, center, right = st.columns([1, 2, 1])

    with center:
        st.markdown(
            "<div class='login-card'><h2 style='text-align:center;'>🔐 Login to <span style='color:#00BFFF;'>HIRELYZER</span></h2>",
            unsafe_allow_html=True,
        )

        login_tab, register_tab = st.tabs(["Login", "Register"])

        # ---------------- LOGIN TAB ----------------
        with login_tab:
            # Show login or forgot password flow based on reset_stage
            if st.session_state.reset_stage == "none":
                # Normal Login UI
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>🔐 Login to Your Account</h3>", unsafe_allow_html=True)

                user = st.text_input("👤 Username or Email", key="login_user")
                pwd = st.text_input("🔑 Password", type="password", key="login_pass")

                # Render notification area (reserves space)
                render_notification("login")

                if st.button("🚀 Login", key="login_btn", use_container_width=True):
                    success, saved_key = verify_user(user.strip(), pwd.strip())
                    if success:
                        st.session_state.authenticated = True
                        # username is already set in session by verify_user()
                        if saved_key:
                            st.session_state["user_groq_key"] = saved_key
                        log_user_action(st.session_state.username, "login")

                        notify("login", "success", "✅ Login successful!")
                        time.sleep(3.0)
                        st.rerun()
                    else:
                        notify("login", "error", "❌ Invalid credentials. Please try again.")
                        st.rerun()

                st.markdown("<br>", unsafe_allow_html=True)

                # Forgot Password Link
                if st.button("🔑 Forgot Password?", key="forgot_pw_link"):
                    st.session_state.reset_stage = "request_email"
                    st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 1: Request Email
            # ============================================================
            elif st.session_state.reset_stage == "request_email":
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>🔐 Reset Password</h3>", unsafe_allow_html=True)
                st.markdown("<p style='color:#c9d1d9; text-align:center;'>Enter your registered email to receive an OTP</p>", unsafe_allow_html=True)

                email_input = st.text_input("📧 Email Address", key="reset_email_input")

                # Render notification area (reserves space)
                render_notification("login")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("📤 Send OTP", key="send_otp_btn", use_container_width=True):
                        if email_input.strip():
                            if get_user_by_email(email_input.strip()):
                                # Generate and send OTP
                                otp = generate_otp()
                                success = send_email_otp(email_input.strip(), otp)

                                if success:
                                    st.session_state.reset_email = email_input.strip()
                                    st.session_state.reset_otp = otp
                                    st.session_state.reset_otp_time = time.time()
                                    st.session_state.reset_stage = "verify_otp"

                                    notify("login", "success", "✅ OTP sent successfully to your email!")
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    notify("login", "error", "❌ Failed to send OTP. Please try again.")
                                    st.rerun()
                            else:
                                notify("login", "error", "❌ Email not found. Please register first.")
                                st.rerun()
                        else:
                            notify("login", "warning", "⚠️ Please enter your email address.")
                            st.rerun()

                with col2:
                    if st.button("↩️ Back to Login", key="back_to_login_1", use_container_width=True):
                        st.session_state.reset_stage = "none"
                        st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 2: Verify OTP
            # ============================================================
            elif st.session_state.reset_stage == "verify_otp":
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>📩 Verify OTP</h3>", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#c9d1d9; text-align:center;'>Enter the 6-digit OTP sent to <strong>{st.session_state.reset_email}</strong></p>", unsafe_allow_html=True)

                # Calculate elapsed and remaining time (server-side)
                elapsed_time = time.time() - st.session_state.reset_otp_time
                remaining_time = max(0, int(180 - elapsed_time))

                # Display timer
                display_timer(remaining_time, expired=(remaining_time == 0), key_suffix="forgot_pw")

                # Check if OTP expired (3 minutes)
                if remaining_time == 0:
                    # OTP Expired - Show resend option
                    render_notification("login")
                    notify("login", "error", "⏱️ OTP expired. Please request a new one.")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("🔄 Resend OTP", key="resend_otp_btn", use_container_width=True):
                            # Generate new OTP
                            otp = generate_otp()
                            success = send_email_otp(st.session_state.reset_email, otp)

                            if success:
                                st.session_state.reset_otp = otp
                                st.session_state.reset_otp_time = time.time()
                                notify("login", "info", "📨 New OTP sent!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("login", "error", "❌ Failed to send OTP. Please try again.")
                                st.rerun()

                    with col2:
                        if st.button("↩️ Back to Login", key="back_to_login_expired", use_container_width=True):
                            st.session_state.reset_stage = "none"
                            st.rerun()
                else:
                    # OTP still valid - Show verification form
                    otp_input = st.text_input("🔢 Enter 6-Digit OTP", key="otp_input", max_chars=6)

                    # Render notification area (reserves space)
                    render_notification("login")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("✅ Verify OTP", key="verify_otp_btn", use_container_width=True):
                            # Re-check expiry on server side before verifying
                            current_elapsed = time.time() - st.session_state.reset_otp_time
                            if current_elapsed >= 180:
                                notify("login", "error", "⏱️ OTP has expired. Please request a new one.")
                                st.rerun()
                            elif otp_input.strip() == st.session_state.reset_otp:
                                st.session_state.reset_stage = "reset_password"
                                notify("login", "success", "✅ OTP verified successfully!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("login", "error", "❌ Invalid OTP. Please try again.")
                                st.rerun()

                    with col2:
                        if st.button("↩️ Back to Login", key="back_to_login_2", use_container_width=True):
                            st.session_state.reset_stage = "none"
                            st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 3: Reset Password
            # ============================================================
            elif st.session_state.reset_stage == "reset_password":
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>🔐 Reset Password</h3>", unsafe_allow_html=True)
                st.markdown("<p style='color:#c9d1d9; text-align:center;'>Enter your new password</p>", unsafe_allow_html=True)

                new_password = st.text_input("🔑 New Password", type="password", key="new_password_input")
                confirm_password = st.text_input("🔑 Confirm Password", type="password", key="confirm_password_input")

                st.caption("Password must be at least 8 characters, include uppercase, lowercase, number, and special character.")

                # Render notification area (reserves space)
                render_notification("login")

                if st.button("✅ Reset Password", key="reset_password_btn", use_container_width=True):
                    if new_password.strip() and confirm_password.strip():
                        if new_password == confirm_password:
                            success = update_password_by_email(st.session_state.reset_email, new_password)

                            if success:
                                notify("login", "success", "✅ Password reset successful! Please log in again.")

                                # Log the password reset action
                                log_user_action(st.session_state.reset_email, "password_reset")

                                # Reset all forgot password session states
                                st.session_state.reset_stage = "none"
                                st.session_state.reset_email = ""
                                st.session_state.reset_otp = ""
                                st.session_state.reset_otp_time = 0

                                time.sleep(1)
                                st.rerun()
                            else:
                                notify("login", "error", "❌ Failed to reset password. Please try again.")
                                st.rerun()
                        else:
                            notify("login", "error", "❌ Passwords do not match.")
                            st.rerun()
                    else:
                        notify("login", "warning", "⚠️ Please fill in both password fields.")
                        st.rerun()

                if st.button("↩️ Back to Login", key="back_to_login_3"):
                    st.session_state.reset_stage = "none"
                    st.rerun()

        # ---------------- REGISTER TAB ----------------
        with register_tab:
            # Check if OTP was sent and pending verification
            if 'pending_registration' in st.session_state:
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>📧 Verify Your Email</h3>", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#c9d1d9; text-align:center;'>Enter the 6-digit OTP sent to <strong>{st.session_state.pending_registration['email']}</strong></p>", unsafe_allow_html=True)

                # Calculate remaining time
                from datetime import datetime
                elapsed = (datetime.now(st.session_state.pending_registration['timestamp'].tzinfo) - st.session_state.pending_registration['timestamp']).total_seconds()
                remaining = max(0, 180 - int(elapsed))

                # Display timer
                display_timer(remaining, expired=(remaining == 0), key_suffix="register")

                if remaining == 0:
                    # OTP Expired
                    render_notification("register")
                    notify("register", "error", "⏱️ OTP expired. Please request a new one.")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("🔄 Resend OTP", key="reg_resend_expired_btn", use_container_width=True):
                            pending = st.session_state.pending_registration
                            success, message = add_user(pending['username'], pending['password'], pending['email'])
                            if success:
                                notify("register", "success", "✅ New OTP sent!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("register", "error", f"❌ {message}")
                                st.rerun()
                    with col2:
                        if st.button("↩️ Start Over", key="reg_start_over_btn", use_container_width=True):
                            del st.session_state.pending_registration
                            st.rerun()
                else:
                    # OTP still valid
                    otp_input = st.text_input("🔢 Enter 6-Digit OTP", key="reg_otp_input", max_chars=6)

                    # Render notification area (reserves space)
                    render_notification("register")

                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.button("✅ Verify", key="verify_reg_otp_btn", use_container_width=True):
                            # Cache username BEFORE calling complete_registration
                            cached_username = st.session_state.pending_registration['username']

                            # Re-check expiry before verification
                            current_elapsed = (datetime.now(st.session_state.pending_registration['timestamp'].tzinfo) - st.session_state.pending_registration['timestamp']).total_seconds()
                            if current_elapsed >= 180:
                                notify("register", "error", "⏱️ OTP has expired. Please request a new one.")
                                st.rerun()
                            else:
                                success, message = complete_registration(otp_input.strip())
                                if success:
                                    notify("register", "success", message)
                                    log_user_action(cached_username, "register")
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    notify("register", "error", message)
                                    st.rerun()

                    with col2:
                        if st.button("🔄 Resend", key="resend_reg_otp_btn", use_container_width=True):
                            pending = st.session_state.pending_registration
                            success, message = add_user(pending['username'], pending['password'], pending['email'])
                            if success:
                                notify("register", "info", "📨 New OTP sent successfully!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("register", "error", f"❌ {message}")
                                st.rerun()

                    with col3:
                        if st.button("↩️ Back", key="back_to_reg_btn", use_container_width=True):
                            del st.session_state.pending_registration
                            st.rerun()

            else:
                # Normal registration form
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>🧾 Register New User</h3>", unsafe_allow_html=True)

                # Email input with live validation
                new_email = st.text_input("📧 Email", key="reg_email", placeholder="your@email.com")

                # Email validation placeholder (using st.empty for dynamic updates)
                email_validation_placeholder = st.empty()

                # Check if email changed and validate
                if new_email and new_email != st.session_state.last_validated_email:
                    if not is_valid_email(new_email.strip()):
                        with email_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message warn-msg"><span class="slide-message-text">⚠️ Invalid email format.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_email = new_email
                    elif email_exists(new_email.strip()):
                        with email_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message error-msg"><span class="slide-message-text">❌ Email already registered.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_email = new_email
                    else:
                        with email_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message success-msg"><span class="slide-message-text">✅ Email is available.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_email = new_email
                        # Auto-hide after 3 seconds by clearing after delay
                        time.sleep(3)
                        email_validation_placeholder.empty()
                elif not new_email:
                    email_validation_placeholder.empty()
                    st.session_state.last_validated_email = ""

                # Username input with live validation
                new_user = st.text_input("👤 Username", key="reg_user")

                # Username validation placeholder
                username_validation_placeholder = st.empty()

                # Check if username changed and validate
                if new_user and new_user != st.session_state.last_validated_username:
                    if username_exists(new_user.strip()):
                        with username_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message error-msg"><span class="slide-message-text">❌ Username already exists.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_username = new_user
                    else:
                        with username_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message success-msg"><span class="slide-message-text">✅ Username is available.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_username = new_user
                        time.sleep(3)
                        username_validation_placeholder.empty()
                elif not new_user:
                    username_validation_placeholder.empty()
                    st.session_state.last_validated_username = ""

                # Password input with live validation
                new_pass = st.text_input("🔑 Password", type="password", key="reg_pass")
                st.caption("Password must be at least 8 characters, include uppercase, lowercase, number, and special character.")

                # Password validation placeholder
                password_validation_placeholder = st.empty()

                # Check if password changed and validate
                if new_pass and new_pass != st.session_state.last_validated_password:
                    if not is_strong_password(new_pass):
                        with password_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message warn-msg"><span class="slide-message-text">⚠️ Password must be at least 8 characters and strong.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_password = new_pass
                    else:
                        with password_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message success-msg"><span class="slide-message-text">✅ Strong password.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_password = new_pass
                        time.sleep(3)
                        password_validation_placeholder.empty()
                elif not new_pass:
                    password_validation_placeholder.empty()
                    st.session_state.last_validated_password = ""

                # Render notification area (reserves space)
                render_notification("register")

                if st.button("📧 Register & Send OTP", key="register_btn", use_container_width=True):
                    if new_email.strip() and new_user.strip() and new_pass.strip():
                        # Validate before attempting registration
                        if not is_valid_email(new_email.strip()):
                            notify("register", "warning", "⚠️ Invalid email format.")
                            st.rerun()
                        elif email_exists(new_email.strip()):
                            notify("register", "error", "🚫 Email already registered.")
                            st.rerun()
                        elif username_exists(new_user.strip()):
                            notify("register", "error", "🚫 Username already exists.")
                            st.rerun()
                        else:
                            success, message = add_user(new_user.strip(), new_pass.strip(), new_email.strip())
                            if success:
                                notify("register", "success", message)
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("register", "error", message)
                                st.rerun()
                    else:
                        notify("register", "warning", "⚠️ Please fill in all fields (email, username, and password).")
                        st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)

    st.stop()

# ------------------- AFTER LOGIN -------------------
if st.session_state.get("authenticated"):
    st.markdown(
        f"<h2 style='color:#00BFFF;'>Welcome to HIRELYZER, <span style='color:white;'>{st.session_state.username}</span> 👋</h2>",
        unsafe_allow_html=True,
    )

    # 🔓 LOGOUT BUTTON
    if st.button("🚪 Logout"):
        log_user_action(st.session_state.get("username", "unknown"), "logout")

        # ✅ Clear all session keys safely
        for key in list(st.session_state.keys()):
            del st.session_state[key]

        st.success("✅ Logged out successfully.")
        st.rerun()  # Force rerun to prevent stale UI

    # 🔑 GROQ API KEY SECTION (SIDEBAR)
    st.sidebar.markdown("### 🔑 Groq API Key")

    # ✅ Load saved key from DB
    saved_key = get_user_api_key(st.session_state.username)
    masked_preview = f"****{saved_key[-6:]}" if saved_key else ""

    user_api_key_input = st.sidebar.text_input(
        "Your Groq API Key (Optional)",
        placeholder=masked_preview,
        type="password"
    )

    # ✅ Save or reuse key
    if user_api_key_input:
        st.session_state["user_groq_key"] = user_api_key_input
        save_user_api_key(st.session_state.username, user_api_key_input)
        st.sidebar.success("✅ New key saved and in use.")
    elif saved_key:
        st.session_state["user_groq_key"] = saved_key
        st.sidebar.info(f"ℹ️ Using your previously saved API key ({masked_preview})")
    else:
        st.sidebar.warning("⚠ Using shared admin key with possible usage limits")

    # 🧹 Clear saved key
    if st.sidebar.button("🗑️ Clear My API Key"):
        st.session_state["user_groq_key"] = None
        save_user_api_key(st.session_state.username, None)
        st.sidebar.success("✅ Cleared saved Groq API key. Now using shared admin key.")

if st.session_state.username == "admin":
    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#00BFFF;'>📊 Admin Dashboard</h2>", unsafe_allow_html=True)

    # Metrics row
    col1, col2 = st.columns(2)
    with col1:
        st.metric("👤 Total Registered Users", get_total_registered_users())
    with col2:
        st.metric("📅 Logins Today (IST)", get_logins_today())

    # Removed API key usage section (no longer tracked)
    # Activity log
    st.markdown("<h3 style='color:#00BFFF;'>📋 Admin Activity Log</h3>", unsafe_allow_html=True)
    logs = get_all_user_logs()
    if logs:
        st.dataframe(
            {
                "Username": [log[0] for log in logs],
                "Action": [log[1] for log in logs],
                "Timestamp": [log[2] for log in logs]
            },
            use_container_width=True
        )
    else:
        st.info("No logs found yet.")

    st.divider()
    st.subheader("📦 Database Backup & Download")

    if os.path.exists(DB_PATH):
        with open(DB_PATH, "rb") as f:
            st.download_button(
                "⬇️ Download resume_data.db",
                data=f,
                file_name="resume_data_backup.db",
                mime="application/octet-stream"
            )
    else:
        st.warning("⚠️ No database file found yet.")
# Always-visible tabs
tab_labels = [
    "📊 Dashboard",
    "🧾 Resume Builder",
    "💼 Job Search",
    "📚 Course Recommendation"
]

# Add Admin tab only for admin user
if st.session_state.username == "admin":
    tab_labels.append("📁 Admin DB View")

# Create tabs dynamically
tabs = st.tabs(tab_labels)

# Unpack first four (always exist)
tab1, tab2, tab3, tab4 = tabs[:4]

# Handle optional admin tab
tab5 = tabs[4] if len(tabs) > 4 else None
with tab1:
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Orbitron', sans-serif;
        background-color: #0b0c10;
        color: #c5c6c7;
        scroll-behavior: smooth;
    }

    /* ---------- SCROLLBAR ---------- */
    ::-webkit-scrollbar { width: 8px; }
    ::-webkit-scrollbar-track { background: #1f2833; }
    ::-webkit-scrollbar-thumb { background: #00ffff; border-radius: 4px; }

    /* ---------- BANNER ---------- */
    .banner-container {
        width: 100%;
        height: 80px;
        background: linear-gradient(90deg, #000428, #004e92);
        border-bottom: 2px solid cyan;
        overflow: hidden;
        display: flex;
        align-items: center;
        justify-content: flex-start;
        position: relative;
        margin-bottom: 20px;
        border-radius: 12px;
        backdrop-filter: blur(14px);
    }
    .pulse-bar {
        position: absolute;
        display: flex;
        align-items: center;
        font-size: 22px;
        font-weight: bold;
        color: #00ffff;
        white-space: nowrap;
        animation: glideIn 12s linear infinite;
        text-shadow: 0 0 10px #00ffff;
    }
    .pulse-bar .bar {
        width: 10px;
        height: 30px;
        margin-right: 10px;
        background: #00ffff;
        box-shadow: 0 0 8px cyan;
        animation: pulse 1s ease-in-out infinite;
    }
    @keyframes glideIn {
        0% { left: -50%; opacity: 0; }
        10% { opacity: 1; }
        90% { opacity: 1; }
        100% { left: 110%; opacity: 0; }
    }
    @keyframes pulse {
        0%, 100% { height: 20px; background-color: #00ffff; }
        50% { height: 40px; background-color: #ff00ff; }
    }

    /* ---------- HEADER ---------- */
    .header {
        font-size: 28px;
        font-weight: bold;
        text-align: center;
        text-transform: uppercase;
        letter-spacing: 2px;
        padding: 20px 30px;  /* ✅ More spacing inside the bar */
        color: #00ffff;
        text-shadow: 0px 0px 10px #00ffff;
        position: relative;
        overflow: hidden;
        border-radius: 14px;
        background: rgba(10,20,40,0.35);
        backdrop-filter: blur(14px);
        border: 1px solid rgba(0,200,255,0.5);
        box-shadow: 0 0 12px rgba(0,200,255,0.25);
    }
    .header::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.18) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .header:hover::before { left: 100%; top: 100%; }

    /* ---------- SHIMMER (COMMON) ---------- */
    .shimmer::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .shimmer:hover::before { left: 100%; top: 100%; }

    /* ---------- FILE UPLOADER ---------- */
    .stFileUploader > div > div {
        border: 1px solid rgba(0,200,255,0.5);
        border-radius: 14px;
        background: rgba(10,20,40,0.35);
        backdrop-filter: blur(14px);
        color: #cce6ff;
        box-shadow: 0 0 12px rgba(0,200,255,0.3);
        position: relative;
        overflow: hidden;
    }
    .stFileUploader > div > div::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stFileUploader > div > div:hover::before { left: 100%; top: 100%; }

    /* ---------- BUTTONS ---------- */
    .stButton > button {
        position: relative;
        overflow: hidden;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        color: #e6f7ff;
        border-radius: 14px;
        padding: 10px 20px;
        font-size: 16px;
        font-weight: 500;
        text-transform: uppercase;
        backdrop-filter: blur(16px);
        box-shadow: 0 0 12px rgba(0,200,255,0.35),
                    inset 0 0 20px rgba(0,200,255,0.05);
        transition: all 0.3s ease-in-out;
    }
    .stButton > button::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stButton > button:hover::before { left: 100%; top: 100%; }

    /* ---------- INPUTS ---------- */
    .stTextInput > div > input,
    .stTextArea > div > textarea {
        position: relative;
        overflow: hidden;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        border-radius: 14px;
        color: #e6f7ff;
        padding: 10px;
        backdrop-filter: blur(16px);
        box-shadow: 0 0 12px rgba(0,200,255,0.3),
                    inset 0 0 15px rgba(0,200,255,0.05);
        transition: all 0.3s ease-in-out;
    }

    /* ---------- CHAT MESSAGES ---------- */
    .stChatMessage {
        position: relative;
        overflow: hidden;
        font-size: 18px;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.5);
        border-radius: 14px;
        padding: 14px;
        color: #e6f7ff;
        text-shadow: 0 0 6px rgba(0,200,255,0.7);
        box-shadow: 0 0 12px rgba(0,200,255,0.3),
                    inset 0 0 15px rgba(0,200,255,0.05);
    }
    .stChatMessage::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stChatMessage:hover::before { left: 100%; top: 100%; }

    /* ---------- METRICS ---------- */
    .stMetric {
        position: relative;
        overflow: hidden;
        background-color: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        border-radius: 14px;
        padding: 15px;
        box-shadow: 0 0 12px rgba(0,200,255,0.35),
                    inset 0 0 20px rgba(0,200,255,0.05);
        text-align: center;
    }
    .stMetric::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stMetric:hover::before { left: 100%; top: 100%; }

    /* ---------- MOBILE ---------- */
    @media (max-width: 768px) {
        .pulse-bar { font-size: 16px; }
        .header { font-size: 20px; }
    }
    </style>

    <!-- Banner -->
    <div class="banner-container">
        <div class="pulse-bar">
            <div class="bar"></div>
            <div>HIRELYZER - Elevate Your Resume Analysis</div>
        </div>
    </div>

    <!-- Header -->
    <div class="header">💼 HIRELYZER - AI BASED ETHICAL RESUME ANALYZER</div>
    """, unsafe_allow_html=True)

# Load environment variables
load_dotenv()

# Detect Device
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
torch.backends.cudnn.benchmark = True
working_dir = os.path.dirname(os.path.abspath(__file__))

# ------------------- Lazy Initialization -------------------
@st.cache_resource(show_spinner=False)
def get_easyocr_reader():
    import easyocr
    return easyocr.Reader(["en"], gpu=torch.cuda.is_available())

@st.cache_data(show_spinner=False)
def ensure_nltk():
    import nltk
    nltk.download('wordnet', quiet=True)
    return WordNetLemmatizer()

lemmatizer = ensure_nltk()
reader = get_easyocr_reader()

def generate_docx(text, filename="bias_free_resume.docx"):
    doc = Document()
    doc.add_heading('Bias-Free Resume', 0)
    doc.add_paragraph(text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Extract text from PDF
def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text_list = [page.get_text("text") for page in doc if page.get_text("text").strip()]
        doc.close()
        return text_list if text_list else extract_text_from_images(file_path)
    except Exception as e:
        st.error(f"⚠ Error extracting text: {e}")
        return []

def extract_text_from_images(pdf_path):
    try:
        images = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=5)
        return ["\n".join(reader.readtext(np.array(img), detail=0)) for img in images]
    except Exception as e:
        st.error(f"⚠ Error extracting from image: {e}")
        return []

def safe_extract_text(uploaded_file):
    """
    Safely extracts text from uploaded file.
    Prevents app crash if file is not a resume or unreadable.
    """
    try:
        # Save uploaded file to a temp location
        temp_path = f"/tmp/{uploaded_file.name}"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Try PDF text extraction
        text_list = extract_text_from_pdf(temp_path)

        # If nothing readable found
        if not text_list or all(len(t.strip()) == 0 for t in text_list):
            st.warning("⚠️ This file doesn't look like a resume or contains no readable text.")
            return None

        return "\n".join(text_list)

    except Exception as e:
        st.error(f"⚠️ Could not process this file: {e}")
        return None

# Detect bias in resume
# Predefined gender-coded word lists
gender_words = {
    "masculine": [
        "active", "aggressive", "ambitious", "analytical", "assertive", "autonomous", "boast", "bold",
        "challenging", "competitive", "confident", "courageous", "decisive", "determined", "dominant", "driven",
        "dynamic", "forceful", "independent", "individualistic", "intellectual", "lead", "leader", "objective",
        "outspoken", "persistent", "principled", "proactive", "resilient", "self-reliant", "self-sufficient",
        "strong", "superior", "tenacious","guru","tech guru","technical guru", "visionary", "manpower", "strongman", "command",
        "assert", "headstrong", "rockstar", "superstar", "go-getter", "trailblazer", "results-driven",
        "fast-paced", "driven", "determination", "competitive spirit"
    ],
    
    "feminine": [
        "affectionate", "agreeable", "attentive", "collaborative", "committed", "compassionate", "considerate",
        "cooperative", "dependable", "dependent", "emotional", "empathetic", "enthusiastic", "friendly", "gentle",
        "honest", "inclusive", "interpersonal", "kind", "loyal", "modest", "nurturing", "pleasant", "polite",
        "sensitive", "supportive", "sympathetic", "tactful", "tender", "trustworthy", "understanding", "warm",
        "yield", "adaptable", "communal", "helpful", "dedicated", "respectful", "nurture", "sociable",
        "relationship-oriented", "team player", "dependable", "people-oriented", "empathetic listener",
        "gentle communicator", "open-minded"
    ]
}

def detect_bias(text):
    # Split into sentences using simple delimiters
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())

    masc_set, fem_set = set(), set()
    masculine_found, feminine_found = [] , []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    for sent in sentences:
        sent_text = sent.strip()
        sent_lower = sent_text.lower()
        matched_spans = []

        def is_overlapping(start, end):
            return any(start < e and end > s for s, e in matched_spans)

        # 🔵 Highlight masculine words in blue
        for word in masculine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in masc_set:
                        masc_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:blue;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        masculine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

        # 🔴 Highlight feminine words in red
        for word in feminine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in fem_set:
                        fem_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:red;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        feminine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

    masc = len(masculine_found)
    fem = len(feminine_found)
    total = masc + fem
    bias_score = min(total / 20, 1.0) if total > 0 else 0.0

    return round(bias_score, 2), masc, fem, masculine_found, feminine_found

replacement_mapping = {
    "masculine": {
        "active": "engaged",
        "aggressive": "proactive",
        "ambitious": "motivated",
        "analytical": "detail-oriented",
        "assertive": "confident",
        "autonomous": "self-directed",
        "boast": "highlight",
        "bold": "confident",
        "challenging": "demanding",
        "competitive": "goal-oriented",
        "confident": "self-assured",
        "courageous": "bold",
        "decisive": "action-oriented",
        "determined": "focused",
        "dominant": "influential",
        "driven": "committed",
        "dynamic": "adaptable",
        "forceful": "persuasive",
        "guru":"technical expert",
        "independent": "self-sufficient",
        "individualistic": "self-motivated",
        "intellectual": "knowledgeable",
        "lead": "guide",
        "leader": "team lead",
        "objective": "unbiased",
        "outspoken": "expressive",
        "persistent": "resilient",
        "principled": "ethical",
        "proactive": "initiative-taking",
        "resilient": "adaptable",
        "self-reliant": "resourceful",
        "self-sufficient": "capable",
        "strong": "capable",
        "superior": "exceptional",
        "tenacious": "determined",
        "technical guru": "technical expert",
        "visionary": "forward-thinking",
        "manpower": "workforce",
        "strongman": "resilient individual",
        "command": "direct",
        "assert": "state confidently",
        "headstrong": "determined",
        "rockstar": "top performer",
        "superstar": "outstanding contributor",
        "go-getter": "initiative-taker",
        "trailblazer": "innovator",
        "results-driven": "outcome-focused",
        "fast-paced": "dynamic",
        "determination": "commitment",
        "competitive spirit": "goal-oriented mindset"
    },
    
    "feminine": {
        "affectionate": "approachable",
        "agreeable": "cooperative",
        "attentive": "observant",
        "collaborative": "team-oriented",
        "collaborate": "team-oriented",
        "collaborated": "worked together",
        "committed": "dedicated",
        "compassionate": "caring",
        "considerate": "thoughtful",
        "cooperative": "supportive",
        "dependable": "reliable",
        "dependent": "team-oriented",
        "emotional": "passionate",
        "empathetic": "understanding",
        "enthusiastic": "positive",
        "gentle": "respectful",
        "honest": "trustworthy",
        "inclusive": "open-minded",
        "interpersonal": "people-focused",
        "kind": "respectful",
        "loyal": "dedicated",
        "modest": "humble",
        "nurturing": "supportive",
        "pleasant": "positive",
        "polite": "professional",
        "sensitive": "attentive",
        "supportive": "encouraging",
        "sympathetic": "understanding",
        "tactful": "diplomatic",
        "tender": "considerate",
        "trustworthy": "reliable",
        "understanding": "empathetic",
        "warm": "welcoming",
        "yield": "adaptable",
        "adaptable": "flexible",
        "communal": "team-centered",
        "helpful": "supportive",
        "dedicated": "committed",
        "respectful": "considerate",
        "nurture": "develop",
        "sociable": "friendly",
        "relationship-oriented": "team-focused",
        "team player": "collaborative member",
        "people-oriented": "person-focused",
        "empathetic listener": "active listener",
        "gentle communicator": "considerate communicator",
        "open-minded": "inclusive"
    }
}

def rewrite_text_with_llm(text, replacement_mapping, user_location):
    """
    Uses LLM to rewrite a resume with bias-free language, while preserving
    the original content length. Enhances grammar, structure, and clarity.
    Ensures structured formatting and includes relevant links and job suggestions.
    """

    # Create a clear mapping in bullet format
    formatted_mapping = "\n".join(
        [f'- "{key}" → "{value}"' for key, value in replacement_mapping.items()]
    )

    # Prompt for LLM
    prompt = f"""
You are an expert resume editor and career advisor.

Your tasks:

1. ✨ Rewrite the resume text below with these rules:
   - Replace any biased or gender-coded language using the exact matches from the replacement mapping.
   - Do NOT reduce the length of any section — preserve the original **number of words per section**.
   - Improve grammar, tone, sentence clarity, and flow without shortening or removing any content.
   - Do NOT change or remove names, tools, technologies, certifications, or project details.

2. 🧾 Structure the resume using these sections **if present** in the original, keeping the original text size:
   - 🏷️ **Name**
   - 📞 **Contact Information**
   - 📍 **Location**
   - 📧 **Email**
   - 🔗 **LinkedIn** → If missing, insert: 🔗 Please paste your LinkedIn URL here.
   - 🌐 **Portfolio** → If missing, insert: 🌐 Please paste your GitHub or portfolio link here.
   - ✍️ **Professional Summary**
   - 💼 **Work Experience**
   - 🧑‍💼 **Internships**
   - 🛠️ **Skills**
   - 🤝 **Soft Skills**
   - 🎓 **Certifications**
   - 🏫 **Education**
   - 📂 **Projects**
   - 🌟 **Interests**

   - Use bullet points (•) inside each section for clarity.
   - Maintain new lines after each points properly.
   - Keep all hyperlinks intact and show them in full where applicable (e.g., LinkedIn, GitHub, project links).
   - Do not invent or assume any information not present in the original.

3. 📌 Strictly apply this **replacement mapping** (match exact phrases only — avoid altering keywords or terminology):
{formatted_mapping}

4. 💼 Suggest **5 relevant job titles** suited for this candidate based in **{user_location}**. For each:
   - Provide a detailed  reason for relevance.
   - Attach a direct LinkedIn job search URL.

---

### 📄 Original Resume Text
\"\"\"{text}\"\"\"

---

### ✅ Bias-Free Rewritten Resume (Fully Structured, Same Length)

---

### 🎯 Suggested Job Titles with Reasoning and LinkedIn Search Links

1. **[Job Title 1]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%201&location={user_location})

2. **[Job Title 2]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%202&location={user_location})

3. **[Job Title 3]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%203&location={user_location})

4. **[Job Title 4]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%204&location={user_location})

5. **[Job Title 5]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%205&location={user_location})
"""

    # Call the LLM of your choice
    response = call_llm(prompt, session=st.session_state)
    return response

def rewrite_and_highlight(text, replacement_mapping, user_location):
    highlighted_text = text
    masculine_count, feminine_count = 0, 0
    detected_masculine_words, detected_feminine_words = [], []
    matched_spans = []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    def span_overlaps(start, end):
        return any(s < end and e > start for s, e in matched_spans)

    # Highlight and count masculine words
    for word in masculine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:blue;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            masculine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:blue;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_masculine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # Highlight and count feminine words
    for word in feminine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:red;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            feminine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:red;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_feminine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # Rewrite text with neutral terms
    rewritten_text = rewrite_text_with_llm(
        text,
        replacement_mapping["masculine"] | replacement_mapping["feminine"],
        user_location
    )

    return highlighted_text, rewritten_text, masculine_count, feminine_count, detected_masculine_words, detected_feminine_words

# ✅ Enhanced Grammar evaluation using LLM with suggestions
def get_grammar_score_with_llm(text, max_score=5):
    grammar_prompt = f"""
You are a grammar and tone evaluator AI. Analyze the following resume text and:

1. Give a grammar score out of {max_score} based on grammar quality, sentence structure, clarity, and tone.
2. Return a 1-sentence summary of the grammar and tone.
3. Provide 3 to 5 **specific improvement suggestions** (bullet points) for enhancing grammar, clarity, tone, or structure.

**Scoring Guidelines for Balance:**
- {max_score}: Exceptional - Professional, error-free, excellent flow
- {max_score-1}: Very Good - Minor issues, mostly professional
- {max_score-2}: Good - Some grammar issues but readable and professional
- {max_score-3}: Fair - Noticeable issues but understandable
- {max_score-4}: Poor - Multiple errors affecting readability
- 0-1: Very Poor - Significant grammar problems

Return response in the exact format below:

Score: <number>
Feedback: <summary>
Suggestions:
- <suggestion 1>
- <suggestion 2>
...

---
{text}
---
"""

    response = call_llm(grammar_prompt, session=st.session_state).strip()
    score_match = re.search(r"Score:\s*(\d+)", response)
    feedback_match = re.search(r"Feedback:\s*(.+)", response)
    suggestions = re.findall(r"- (.+)", response)

    score = int(score_match.group(1)) if score_match else max(3, max_score-2)  # More generous default
    feedback = feedback_match.group(1).strip() if feedback_match else "Grammar appears adequate for professional communication."
    return score, feedback, suggestions

# ✅ Main ATS Evaluation Function
def ats_percentage_score(
    resume_text,
    job_description,
    job_title="Unknown",
    logic_profile_score=None,
    edu_weight=20,
    exp_weight=35,
    skills_weight=30,
    lang_weight=5,
    keyword_weight=10
):
    import datetime

    # ✅ Grammar evaluation
    grammar_score, grammar_feedback, grammar_suggestions = get_grammar_score_with_llm(
        resume_text, max_score=lang_weight
    )

    # ✅ Domain similarity detection using LLM
    resume_domain = db_manager.detect_domain_llm(
        "Unknown", 
        resume_text, 
        session=st.session_state  # ✅ pass the Groq API key from session
    )
    job_domain = db_manager.detect_domain_llm(
        job_title, 
        job_description, 
        session=st.session_state  # ✅ pass the Groq API key from session
    )
    similarity_score = get_domain_similarity(resume_domain, job_domain)

    # ✅ Balanced domain penalty
    MAX_DOMAIN_PENALTY = 15
    domain_penalty = round((1 - similarity_score) * MAX_DOMAIN_PENALTY)

    # ✅ Optional profile score note
    logic_score_note = (
        f"\n\nOptional Note: The system also calculated a logic-based profile score of {logic_profile_score}/100 "
        f"based on resume length, experience, and skills."
        if logic_profile_score else ""
    )

    # ✅ FIXED: Stable education scoring with 2025 cutoff
    current_year = datetime.datetime.now().year
    current_month = datetime.datetime.now().month
    
    # ✅ FIXED: Education completion detection with 2025 cutoff
    def determine_education_status(education_text, end_year_str):
        """
        Determine if education is completed or ongoing based on 2025 cutoff and keywords.
        Returns 'completed' or 'ongoing'.
        """
        try:
            end_year = int(end_year_str.strip())
        except (ValueError, AttributeError):
            # If we can't parse the year, default to ongoing
            return "ongoing"
        
        # Apply 2025 cutoff rule (HARDCODED - NOT dynamic)
        if end_year < 2025:
            education_status = "completed"
        elif end_year == 2025:
            education_status = "completed"
        else:  # end_year > 2025
            education_status = "ongoing"
        
        # Check for explicit keywords that might override numeric rules
        education_lower = education_text.lower()
        ongoing_keywords = ["pursuing", "present", "ongoing", "currently enrolled", "in progress"]
        completed_keywords = ["graduated", "completed", "finished"]
        
        # Override rule: If end year < 2025, always completed regardless of text
        if end_year < 2025:
            return "completed"
        
        # For years >= 2025, check keywords
        if end_year < 2025:
            return "completed"
        
        # For years >= 2025, check keywords
        if any(keyword in education_lower for keyword in ongoing_keywords):
            education_status = "ongoing"
        elif any(keyword in education_lower for keyword in completed_keywords):
            education_status = "completed"
        
        return education_status
    
    # ✅ UPDATED: Stable education scoring with priority degrees minimum
    prompt = f"""
You are a professional ATS evaluator specializing in **technical roles** (AI/ML, Blockchain, Cloud, Data, Software, Cybersecurity). 
Your role is to provide **balanced, objective scoring** that reflects industry standards and recognizes candidate potential while maintaining professional standards.

🎯 **BALANCED SCORING GUIDELINES - Tech-Focused (AI/ML/Blockchain/Software/Data):**

**Education Scoring Framework ({edu_weight} points max):**

⚡ **PRIORITY RULE - Minimum Points for Relevant Degrees:**
If candidate is **currently pursuing OR has completed** any of these degrees:
- BSc CS / BSc Computer Science
- BSc Mathematics / BSc Maths
- MSc CS / MSc Computer Science
- MSc Mathematics / MSc Maths
- MCA (Master of Computer Applications)
- BE CS / BTech CS / BTech IT

→ **ASSIGN MINIMUM {int(edu_weight * 0.75)} points** out of {edu_weight} max points
→ **DO NOT penalize** for ongoing status - pursuing counts equally as completed
→ If completed with strong academic performance, allow scoring up to {int(edu_weight * 0.9)}-{edu_weight} points

**CRITICAL DATE PARSING RULES:**
- If end year < 2025 → ✅ ALWAYS Completed (HARDCODED CUTOFF)
- If end year == 2025 → ✅ Completed
- If end year > 2025 → 🔄 Ongoing  

**EXPLICIT STATUS INDICATORS (override year logic for years >= 2025):**
- Words like "pursuing", "currently enrolled", "in progress" → 🔄 Ongoing  
- Words like "Graduated", "Completed", "Finished" → ✅ Completed
- **OVERRIDE RULE**: If end year < 2025, it is ✅ Completed no matter what the text says (2025 IS THE CUTOFF YEAR).

**SCORING IMPACT:**
- ✅ Completed relevant education → Full scoring potential (up to max points)
- 🔄 Ongoing relevant education → **MINIMUM {int(edu_weight * 0.75)} points for priority degrees listed above**
- Education score is based ONLY on degree relevance and completion status
- DO NOT add points for certifications/projects in education - these belong in skills/experience sections

**Stable Education Scoring Framework (Independent of Job Description):**
- {int(edu_weight * 0.90)}-{edu_weight}: Outstanding (completed highly relevant degree with excellent academic performance)
- {int(edu_weight * 0.75)}-{int(edu_weight * 0.85)}: Excellent (priority degrees listed above - completed or ongoing)
- {int(edu_weight * 0.60)}-{int(edu_weight * 0.70)}: Very Good (related technical/quantitative degree)
- {int(edu_weight * 0.45)}-{int(edu_weight * 0.55)}: Good (somewhat related education with transferable knowledge)
- {int(edu_weight * 0.30)}-{int(edu_weight * 0.40)}: Fair (different degree but shows analytical/technical foundation)
- {int(edu_weight * 0.15)}-{int(edu_weight * 0.25)}: Basic (unrelated degree)
- 0-{int(edu_weight * 0.10)}: Insufficient (no degree information or incomplete details)


**Experience Scoring Framework ({exp_weight} points max):**
- {int(exp_weight * 0.91)}-{exp_weight}: Exceptional (exceeds requirements + perfect fit + leadership + outstanding results)
- {int(exp_weight * 0.80)}-{int(exp_weight * 0.89)}: Excellent (meets/exceeds years + strong domain fit + leadership + clear results)
- {int(exp_weight * 0.69)}-{int(exp_weight * 0.77)}: Very Good (adequate years + good domain fit + solid responsibilities + some results)
- {int(exp_weight * 0.57)}-{int(exp_weight * 0.66)}: Good (reasonable years + relevant experience + decent responsibilities)
- {int(exp_weight * 0.43)}-{int(exp_weight * 0.54)}: Fair (some gaps in years OR domain but shows potential)
- {int(exp_weight * 0.29)}-{int(exp_weight * 0.40)}: Basic (limited experience but relevant skills/potential shown)
- {int(exp_weight * 0.14)}-{int(exp_weight * 0.26)}: Entry Level (minimal experience but shows promise)
- 0-{int(exp_weight * 0.11)}: Insufficient (major gaps with no transferable skills)

**Skills Scoring Framework ({skills_weight} points max):**
- {int(skills_weight * 0.93)}-{skills_weight}: Outstanding (90%+ required skills + expert proficiency + recent usage)
- {int(skills_weight * 0.80)}-{int(skills_weight * 0.90)}: Excellent (80%+ required skills + advanced proficiency)
- {int(skills_weight * 0.67)}-{int(skills_weight * 0.77)}: Very Good (70%+ required skills + good proficiency)
- {int(skills_weight * 0.53)}-{int(skills_weight * 0.63)}: Good (60%+ required skills + adequate proficiency)
- {int(skills_weight * 0.40)}-{int(skills_weight * 0.50)}: Fair (50%+ required skills + basic proficiency OR strong learning ability)
- {int(skills_weight * 0.27)}-{int(skills_weight * 0.37)}: Basic (40%+ skills OR strong foundational skills with growth potential)
- {int(skills_weight * 0.13)}-{int(skills_weight * 0.23)}: Limited (30%+ skills but shows willingness to learn)
- 0-{int(skills_weight * 0.10)}: Insufficient (<30% skills with no evidence of learning ability)

**Keyword Scoring Framework ({keyword_weight} points max):**
- {int(keyword_weight * 0.90)}-{keyword_weight}: Excellent optimization (85%+ critical terms + industry language)
- {int(keyword_weight * 0.80)}: Very Good (75%+ critical terms + good industry awareness)
- {int(keyword_weight * 0.60)}-{int(keyword_weight * 0.70)}: Good (65%+ critical terms + adequate industry knowledge)
- {int(keyword_weight * 0.40)}-{int(keyword_weight * 0.50)}: Fair (50%+ critical terms + some industry understanding)
- {int(keyword_weight * 0.20)}-{int(keyword_weight * 0.30)}: Basic (35%+ critical terms + basic awareness)
- {int(keyword_weight * 0.10)}: Limited (20%+ critical terms)
- 0: Poor (<20% critical terms)

**EVALUATION INSTRUCTIONS (Tech-Focused):**
- Always credit **projects, GitHub repos, hackathons, Kaggle competitions, blockchain DApps, cloud deployments, AI model training, open-source contributions**.
- Emphasize **cutting-edge skills**: LLMs, Generative AI, Web3, Smart Contracts, DeFi, Cloud-Native tools, MLOps, Vector DBs.
- Highlight both **industry experience** and **hands-on learning** (projects, MOOCs, certifications).
- Be encouraging but factual: focus on **growth potential + adaptability**.

**EVALUATION INSTRUCTIONS - BE ENCOURAGING BUT HONEST:**

Follow this exact structure and be **specific with evidence while highlighting strengths**:

### 🏷️ Candidate Name
<Extract full name clearly - check resume header, contact section, or first few lines>

### 🏫 Education Analysis
**Score:** <0–{edu_weight}> / {edu_weight}

**Scoring Rationale:**
- Degree Level & Relevance: <Check if degree qualifies for minimum 15 points rule - BSc/MSc CS, BSc/MSc Maths, MCA, BE/BTech CS/IT>
- Completion Status: <Apply 2025 cutoff rule and keyword overrides>
- Academic Foundation: <Assess degree relevance to technical roles>
- **Score Justification:** <Apply minimum 15 points if relevant degree detected; pursuing status not penalized; score based only on degree relevance>


### 💼 Experience Analysis  
**Score:** <0–{exp_weight}> / {exp_weight}

**Experience Breakdown:**
- Total Years: <X years - consider quality over quantity>
- Role Progression: <Look for growth, even if not linear>
- Domain Relevance: <Consider transferable skills from related fields>
- Leadership Evidence: <Include informal leadership, mentoring, project ownership>
- Quantified Achievements: <Value any metrics, even small improvements>
- Technology/Tools Usage: <Credit learning new tools, adaptability>
- Transferable Skills: <Highlight skills that apply across domains>
- **Score Justification:** <Emphasize growth potential and adaptability>

### 🛠 Skills Analysis
**Score:** <0–{skills_weight}> / {skills_weight}

**Skills Assessment:**
- Technical Skills Present: <List with evidence, include learning in progress>
- Soft Skills Demonstrated: <Value communication, teamwork, problem-solving>
- Domain-Specific Expertise: <Consider related domain knowledge>
- Skill Currency: <Value recent learning and adaptation>
- Learning Ability: <Evidence of picking up new skills>

**Skills Gaps (Opportunities for Growth):**
- <Skill 1 - frame as development opportunity>
- <Skill 2 - suggest how existing skills could transfer>  
- <Skill 3 - note if easily learnable>
- <Skill 4 - additional growth areas>
- <Skill 5 - more opportunities if applicable>

**Score Justification:** <Focus on existing strengths + learning potential>

### 🗣 Language Quality Analysis
**Score:** {grammar_score} / {lang_weight}
**Grammar & Professional Tone:** {grammar_feedback}
**Assessment:** <Be constructive - focus on communication effectiveness>

### 🔑 Keyword Analysis
**Score:** <0–{keyword_weight}> / {keyword_weight}

**Keyword Assessment:**
- Industry Terminology: <Credit related industry knowledge>
- Role-Specific Terms: <Look for equivalent terms, not just exact matches>
- Technical Vocabulary: <Value understanding even if different tools>

**Keyword Enhancement Opportunities:**
- <Keyword 1 from job description>
- <Keyword 2 from job description>
- <Keyword 3 from job description>
- <Keyword 4 from job description>
- <Keyword 5 from job description>
- <Keyword 6 from job description>
- <Keyword 7 from job description>
- <Keyword 8 from job description>

**INSTRUCTION**: Extract ALL important keywords, technical terms, industry jargon, tool names, certification names, and role-specific terminology from the job description that are missing from the resume. Include variations and synonyms.

**Score Justification:** <Credit understanding of concepts even if terminology differs>

### ✅ Final Assessment

**Overall Evaluation:**
<4-6 sentences covering:>
- Primary strengths and unique value proposition
- Growth areas framed as development opportunities
- Cultural/team fit indicators and soft skills
- Clear recommendation with constructive reasoning

**Development Areas:** <Frame gaps as growth opportunities, not failures>
**Key Strengths:** <Highlight what makes this candidate valuable>
**Recommendation:** <Be specific about interview potential and role fit>

---

**IMPORTANT REMINDERS FOR BALANCED EVALUATION:**
- Look for potential, not just perfect matches
- Value diverse backgrounds and transferable skills
- Consider the candidate's career stage and growth trajectory
- Credit all forms of learning and skill development
- Be constructive in feedback - focus on opportunities
- Recognize that great employees come from varied backgrounds
- LIST ALL missing skills and keywords comprehensively (aim for 5-8 items each if gaps exist)
- Be thorough in identifying development opportunities from the job description
- **CRITICAL**: Analyze the ENTIRE job description systematically - go through each requirement, skill, and qualification mentioned
- **KEYWORD EXTRACTION**: Identify ALL technical terms, tools, frameworks, methodologies, certifications mentioned in job description
- **SKILL MAPPING**: Compare each job requirement against resume content - if not found, list it as missing
- **CONTEXT UNDERSTANDING**: Consider synonyms and related terms (e.g., "JavaScript" and "JS", "Machine Learning" and "ML")
- **PRIORITY RANKING**: Focus on must-have vs nice-to-have requirements from job description
- **EXPERIENCE MATCHING**: Look for similar roles, projects, or responsibilities even if not exact title matches
- **EDUCATION PRIORITY**: Apply minimum 15 points rule for BSc/MSc CS, BSc/MSc Maths, MCA, BE/BTech CS/IT degrees
Context for Evaluation:
- Current Date: {datetime.datetime.now().strftime('%B %Y')} (Year: {current_year}, Month: {current_month})
- Grammar Score: {grammar_score} / {lang_weight}
- Grammar Feedback: {grammar_feedback}  
- Resume Domain: {resume_domain}
- Job Domain: {job_domain}
- Domain Mismatch Penalty: {domain_penalty} points (similarity: {similarity_score:.2f})

---

📄 **Job Description:**
{job_description}

📄 **Resume Text:**
{resume_text}

{logic_score_note}
"""
   
   
    ats_result = call_llm(prompt, session=st.session_state).strip()

    def extract_section(pattern, text, default="N/A"):
        match = re.search(pattern, text, re.DOTALL)
        return match.group(1).strip() if match else default

    def extract_score(pattern, text, default=0):
        match = re.search(pattern, text)
        return int(match.group(1)) if match else default

    # Extract key sections
    candidate_name = extract_section(r"### 🏷️ Candidate Name(.*?)###", ats_result, "Not Found")
    edu_analysis = extract_section(r"### 🏫 Education Analysis(.*?)###", ats_result)
    exp_analysis = extract_section(r"### 💼 Experience Analysis(.*?)###", ats_result)
    skills_analysis = extract_section(r"### 🛠 Skills Analysis(.*?)###", ats_result)
    lang_analysis = extract_section(r"### 🗣 Language Quality Analysis(.*?)###", ats_result)
    keyword_analysis = extract_section(r"### 🔑 Keyword Analysis(.*?)###", ats_result)
    final_thoughts = extract_section(r"### ✅ Final Assessment(.*)", ats_result)

    # Extract scores with improved patterns (LLM now scores directly using sidebar weights)
    edu_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", edu_analysis)
    exp_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", exp_analysis)
    skills_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", skills_analysis)
    keyword_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", keyword_analysis)
    lang_score = grammar_score  # Grammar score already uses lang_weight

    # ✅ Apply minimum thresholds to avoid overly harsh penalties
    edu_score = max(edu_score, int(edu_weight * 0.15))  # Minimum 15% of weight
    exp_score = max(exp_score, int(exp_weight * 0.15))  # Minimum 15% of weight
    skills_score = max(skills_score, int(skills_weight * 0.15))  # Minimum 15% of weight
    keyword_score = max(keyword_score, int(keyword_weight * 0.10))  # Minimum 10% of weight

    # Extract missing items with better parsing - now called "opportunities"
    missing_keywords_section = extract_section(r"\*\*Keyword Enhancement Opportunities:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    missing_skills_section = extract_section(r"\*\*Skills Gaps \(Opportunities for Growth\):\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    
    # Fallback to old patterns if new ones don't match
    if not missing_keywords_section.strip():
        missing_keywords_section = extract_section(r"\*\*Missing Critical Keywords:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    if not missing_skills_section.strip():
        missing_skills_section = extract_section(r"\*\*Missing Critical Skills:\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    
    # Improved extraction - handle multiple formats and get all items
    def extract_list_items(text):
        if not text.strip():
            return "None identified"
        
        # Find all bullet points with various formats
        items = []
        lines = text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Remove various bullet point formats
            cleaned_line = re.sub(r'^[-•*]\s*', '', line)  # Remove -, •, * bullets
            cleaned_line = re.sub(r'^\d+\.\s*', '', cleaned_line)  # Remove numbered lists
            cleaned_line = cleaned_line.strip()
            
            if cleaned_line and len(cleaned_line) > 2:  # Avoid empty or very short items
                items.append(cleaned_line)
        
        return ', '.join(items) if items else "None identified"
    
    missing_keywords = extract_list_items(missing_keywords_section)
    missing_skills = extract_list_items(missing_skills_section)

    # ✅ IMPROVED: More balanced total score calculation
    total_score = edu_score + exp_score + skills_score + lang_score + keyword_score
    
    # Apply domain penalty more gently
    total_score = max(total_score - domain_penalty, int(total_score * 0.7))  # Never go below 70% of pre-penalty score
    
    # ✅ IMPROVED: More generous score caps and bonus for well-rounded candidates
    total_score = min(total_score, 100)
    total_score = max(total_score, 15)  # Minimum score of 15 to avoid completely crushing candidates

    # ✅ IMPROVED: More encouraging score formatting with better thresholds
    formatted_score = (
        "🌟 Exceptional Match" if total_score >= 85 else  # Lowered from 90
        "✅ Strong Match" if total_score >= 70 else       # Lowered from 75
        "🟡 Good Potential" if total_score >= 55 else    # Lowered from 60
        "⚠️ Fair Match" if total_score >= 40 else        # Lowered from 45
        "🔄 Needs Development" if total_score >= 25 else # New category
        "❌ Poor Match"
    )

    # ✅ Format suggestions nicely
    suggestions_html = ""
    if grammar_suggestions:
        suggestions_html = "<ul>" + "".join([f"<li>{s}</li>" for s in grammar_suggestions]) + "</ul>"

    updated_lang_analysis = f"""
{lang_analysis}
<br><b>LLM Feedback Summary:</b> {grammar_feedback}
<br><b>Improvement Suggestions:</b> {suggestions_html}
"""

    # Enhanced final thoughts with domain analysis
    final_thoughts += f"""

**📊 Technical Assessment Details:**
- Domain Similarity Score: {similarity_score:.2f}/1.0  
- Domain Penalty Applied: {domain_penalty}/{MAX_DOMAIN_PENALTY} points
- Resume Domain: {resume_domain}
- Target Job Domain: {job_domain}

**💡 Balanced Scoring Notes:**
- Minimum score thresholds applied to prevent overly harsh penalties
- Transferable skills and learning potential considered
- Growth opportunities highlighted rather than just gaps identified
- **Date Logic Applied**: Year-only ranges properly classified as completed/ongoing based on current date context
"""

    return ats_result, {
        "Candidate Name": candidate_name,
        "Education Score": edu_score,
        "Experience Score": exp_score,
        "Skills Score": skills_score,
        "Language Score": lang_score,
        "Keyword Score": keyword_score,
        "ATS Match %": total_score,
        "Formatted Score": formatted_score,
        "Education Analysis": edu_analysis,
        "Experience Analysis": exp_analysis,
        "Skills Analysis": skills_analysis,
        "Language Analysis": updated_lang_analysis,
        "Keyword Analysis": keyword_analysis,
        "Final Thoughts": final_thoughts,
        "Missing Keywords": missing_keywords,
        "Missing Skills": missing_skills,
        "Resume Domain": resume_domain,
        "Job Domain": job_domain,
        "Domain Penalty": domain_penalty,
        "Domain Similarity Score": similarity_score
    }

# Setup Vector DB
def setup_vectorstore(documents):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    if DEVICE == "cuda":
        embeddings.model = embeddings.model.to(torch.device("cuda"))
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    doc_chunks = text_splitter.split_text("\n".join(documents))
    return FAISS.from_texts(doc_chunks, embeddings)

# Create Conversational Chain
def create_chain(vectorstore):
    # 🔁 Get a rotated admin key
    keys = load_groq_api_keys()
    index = st.session_state.get("key_index", 0)
    groq_api_key = keys[index % len(keys)]
    st.session_state["key_index"] = index + 1

    # ✅ Create the ChatGroq object
    llm = ChatGroq(model="llama-3.3-70b-versatile", temperature=0, groq_api_key=groq_api_key)

    # ✅ Build the chain
    chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        return_source_documents=True
    )
    return chain

# Chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# ---------------- Sidebar Layout with Inline Images ----------------
st.sidebar.markdown("### 🏷️ Job Information")

# ---------------- Job Information Dropdown ----------------
with st.sidebar.expander("![Job](https://img.icons8.com/ios-filled/20/briefcase.png) Enter Job Details", expanded=False):
    job_title = st.text_input(
        "![Job](https://img.icons8.com/ios-filled/20/briefcase.png) Job Title"
    )

    user_location = st.text_input(
        "![Location](https://img.icons8.com/ios-filled/20/marker.png) Preferred Job Location (City, Country)"
    )

    job_description = st.text_area(
        "![Description](https://img.icons8.com/ios-filled/20/document.png) Paste Job Description",
        height=200
    )

    if job_description.strip() == "":
        st.warning("Please enter a job description to evaluate the resumes.")

# ---------------- Advanced Weights Dropdown ----------------
with st.sidebar.expander("![Settings](https://img.icons8.com/ios-filled/20/settings.png) Customize ATS Scoring Weights", expanded=False):
    edu_weight = st.slider("![Education](https://img.icons8.com/ios-filled/20/graduation-cap.png) Education Weight", 0, 50, 20)
    exp_weight = st.slider("![Experience](https://img.icons8.com/ios-filled/20/portfolio.png) Experience Weight", 0, 50, 35)
    skills_weight = st.slider("![Skills](https://img.icons8.com/ios-filled/20/gear.png) Skills Match Weight", 0, 50, 30)
    lang_weight = st.slider("![Language](https://img.icons8.com/ios-filled/20/language.png) Language Quality Weight", 0, 10, 5)
    keyword_weight = st.slider("![Keyword](https://img.icons8.com/ios-filled/20/key.png) Keyword Match Weight", 0, 20, 10)

    total_weight = edu_weight + exp_weight + skills_weight + lang_weight + keyword_weight

    # ---------------- Inline SVG Validation ----------------
    if total_weight != 100:
        st.markdown(
            f"""
            <div style="display:flex;align-items:center;gap:6px;
                        border:1px solid #fca5a5;
                        background:#fee2e2;
                        padding:8px;
                        border-radius:6px;">
                <svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" fill="red" viewBox="0 0 24 24">
                    <path d="M12 2C6.48 2 2 6.48 2 12s4.48 10 10 10
                             10-4.48 10-10S17.52 2 12 2zm0 15
                             c-.83 0-1.5.67-1.5 1.5S11.17 20
                             12 20s1.5-.67 1.5-1.5S12.83 17
                             12 17zm1-4V7h-2v6h2z"/>
                </svg>
                <span style="color:#b91c1c;font-weight:500;">
                    Total = {total_weight}. Please make it exactly 100.
                </span>
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            f"""
            <div style="display:flex;align-items:center;gap:6px;
                        border:1px solid #86efac;
                        background:#dcfce7;
                        padding:8px;
                        border-radius:6px;">
                <svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" fill="green" viewBox="0 0 24 24">
                    <path d="M9 16.2l-3.5-3.5-1.4 1.4L9
                             19 20.3 7.7l-1.4-1.4z"/>
                </svg>
                <span style="color:#166534;font-weight:500;">
                    Total weight = 100
                </span>
            </div>
            """,
            unsafe_allow_html=True
        )

with tab1:
    # 🎨 CSS for sliding success message
    st.markdown("""
    <style>
    .slide-message {
      position: relative;
      overflow: hidden;
      margin: 10px 0;
      padding: 10px 15px;
      border-radius: 10px;
      font-weight: bold;
      display: flex;
      align-items: center;
      gap: 8px;
      animation: slideIn 0.8s ease forwards;
    }
    .slide-message svg {
      width: 18px;
      height: 18px;
      flex-shrink: 0;
    }
    .success-msg { background: rgba(0,255,127,0.12); border-left: 5px solid #00FF7F; color:#00FF7F; }
    .error-msg   { background: rgba(255,99,71,0.12);  border-left: 5px solid #FF6347; color:#FF6347; }
    .warn-msg    { background: rgba(255,215,0,0.12); border-left: 5px solid #FFD700; color:#FFD700; }

    @keyframes slideIn {
      0%   { transform: translateX(100%); opacity: 0; }
      100% { transform: translateX(0); opacity: 1; }
    }
    </style>
    """, unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "📄 Upload PDF Resumes",
        type=["pdf"],
        accept_multiple_files=True,
        help="Upload one or more resumes in PDF format (max 200MB each)."
    )

    if uploaded_files:
        for uploaded_file in uploaded_files:
            with st.container():
                st.subheader(f"📄 Original Resume Preview: {uploaded_file.name}")

                try:
                    # ✅ Show PDF preview safely
                    pdf_viewer(
                        uploaded_file.read(),
                        key=f"pdf_viewer_{uploaded_file.name}"
                    )

                    # Reset pointer so file can be read again later
                    uploaded_file.seek(0)

                    # ✅ Extract text safely
                    resume_text = safe_extract_text(uploaded_file)

                    if resume_text:
                        st.markdown(f"""
                        <div class='slide-message success-msg'>
                            <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                              stroke-width="2" viewBox="0 0 24 24"><path d="M5 13l4 4L19 7"/></svg>
                            ✅ Successfully processed <b>{uploaded_file.name}</b>
                        </div>
                        """, unsafe_allow_html=True)
                        # 🔹 Continue with ATS scoring, bias detection, etc. here
                    else:
                        st.markdown(f"""
                        <div class='slide-message warn-msg'>
                            ⚠️ <b>{uploaded_file.name}</b> does not contain valid resume text.
                        </div>
                        """, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f"""
                    <div class='slide-message error-msg'>
                        ❌ Could not display or process <b>{uploaded_file.name}</b>: {e}
                    </div>
                    """, unsafe_allow_html=True)

# ✅ Initialize state
# Initialize session state
if "resume_data" not in st.session_state:
    st.session_state.resume_data = []

if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

resume_data = st.session_state.resume_data

# ✏️ Resume Evaluation Logic
if uploaded_files and job_description:
    all_text = []

    for uploaded_file in uploaded_files:
        if uploaded_file.name in st.session_state.processed_files:
            continue

        # ✅ Improved optimized scanner animation with better performance
        scanner_placeholder = st.empty()

        # ✅ IMPROVED: More efficient CSS animations with GPU acceleration
        OPTIMIZED_SCANNER_HTML = f"""
        <style>
        .scanner-overlay {{
            position: fixed;
            top: 0; left: 0;
            width: 100vw; height: 100vh;
            background: linear-gradient(135deg, #0b0c10 0%, #1a1c29 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            z-index: 9999;
            will-change: transform, opacity;
        }}
        
        .scanner-doc {{
            width: 280px;
            height: 340px;
            background: linear-gradient(145deg, #f8f9fa, #e9ecef);
            border-radius: 16px;
            position: relative;
            overflow: hidden;
            box-shadow: 0 20px 40px rgba(0, 191, 255, 0.3);
            transform: translateZ(0);
            will-change: transform;
            animation: docFloat 3s ease-in-out infinite alternate;
        }}
        
        @keyframes docFloat {{
            0% {{ transform: translateY(0px) scale(1); }}
            100% {{ transform: translateY(-8px) scale(1.02); }}
        }}
        
        .doc-header {{
            padding: 20px;
            text-align: center;
            border-bottom: 2px solid #e9ecef;
        }}
        
        .doc-avatar {{
            width: 50px;
            height: 50px;
            background: linear-gradient(135deg, #667eea, #764ba2);
            border-radius: 50%;
            margin: 0 auto 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 20px;
            color: white;
        }}
        
        .doc-title {{
            font-size: 16px;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 5px;
            font-family: 'Segoe UI', sans-serif;
        }}
        
        .doc-content {{
            padding: 15px;
            font-size: 12px;
            color: #6c757d;
            line-height: 1.4;
        }}
        
        .scan-line {{
            position: absolute;
            top: 0; left: 0;
            width: 100%; height: 4px;
            background: linear-gradient(90deg, transparent, rgba(0,191,255,0.8), transparent);
            animation: scanMove 2.5s ease-in-out infinite;
            box-shadow: 0 0 20px rgba(0,191,255,0.6);
            transform: translateZ(0);
            will-change: transform;
        }}
        
        @keyframes scanMove {{
            0% {{ top: 0; opacity: 1; }}
            50% {{ opacity: 0.8; }}
            100% {{ top: 340px; opacity: 1; }}
        }}
        
        .scanner-text {{
            margin-top: 30px;
            font-family: 'Orbitron', 'Segoe UI', sans-serif;
            font-weight: 600;
            font-size: 18px;
            color: #00bfff;
            text-shadow: 0 0 10px rgba(0,191,255,0.5);
            animation: textPulse 2s ease-in-out infinite;
        }}
        
        @keyframes textPulse {{
            0%, 100% {{ opacity: 1; transform: scale(1); }}
            50% {{ opacity: 0.8; transform: scale(1.05); }}
        }}
        
        .progress-bar {{
            width: 200px;
            height: 4px;
            background: rgba(255,255,255,0.2);
            border-radius: 2px;
            margin-top: 20px;
            overflow: hidden;
        }}
        
        .progress-fill {{
            height: 100%;
            background: linear-gradient(90deg, #00bfff, #1e90ff);
            border-radius: 2px;
            animation: progressFill 3s ease-in-out infinite;
            transform: translateX(-100%);
        }}
        
        @keyframes progressFill {{
            0% {{ transform: translateX(-100%); }}
            100% {{ transform: translateX(0); }}
        }}
        
        /* Mobile optimizations */
        @media (max-width: 768px) {{
            .scanner-doc {{ width: 240px; height: 300px; }}
            .scanner-text {{ font-size: 16px; }}
        }}
        </style>
        
        <div class="scanner-overlay">
            <div class="scanner-doc">
                <div class="scan-line"></div>
                <div class="doc-header">
                    <div class="doc-avatar">👤</div>
                    <div class="doc-title">{job_title}</div>
                </div>
                <div class="doc-content">
                    • Analyzing candidate profile...<br>
                    • Extracting key skills...<br>
                    • Matching with job requirements...<br>
                    • Calculating ATS compatibility...<br>
                    • Checking for bias patterns...
                </div>
            </div>
            <div class="scanner-text">Scanning Resume...</div>
            <div class="progress-bar">
                <div class="progress-fill"></div>
            </div>
        </div>
        """
        
        scanner_placeholder.markdown(OPTIMIZED_SCANNER_HTML, unsafe_allow_html=True)

        # ✅ Save uploaded file
        file_path = os.path.join(working_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # ✅ Reduced delay for better UX
        time.sleep(4)

        # ✅ Extract text from PDF
        text = extract_text_from_pdf(file_path)
        if not text:
            st.warning(f"⚠️ Could not extract text from {uploaded_file.name}. Skipping.")
            scanner_placeholder.empty()
            continue

        all_text.append(" ".join(text))
        full_text = " ".join(text)

        # ✅ Bias detection
        bias_score, masc_count, fem_count, detected_masc, detected_fem = detect_bias(full_text)

        # ✅ Rewrite and highlight gender-biased words
        highlighted_text, rewritten_text, _, _, _, _ = rewrite_and_highlight(
            full_text, replacement_mapping, user_location
        )

        # ✅ LLM-based ATS Evaluation
        ats_result, ats_scores = ats_percentage_score(
            resume_text=full_text,
            job_description=job_description,
            logic_profile_score=None,
            edu_weight=edu_weight,
            exp_weight=exp_weight,
            skills_weight=skills_weight,
            lang_weight=lang_weight,
            keyword_weight=keyword_weight
        )

        # ✅ Extract structured ATS values
        candidate_name = ats_scores.get("Candidate Name", "Not Found")
        ats_score = ats_scores.get("ATS Match %", 0)
        edu_score = ats_scores.get("Education Score", 0)
        exp_score = ats_scores.get("Experience Score", 0)
        skills_score = ats_scores.get("Skills Score", 0)
        lang_score = ats_scores.get("Language Score", 0)
        keyword_score = ats_scores.get("Keyword Score", 0)
        formatted_score = ats_scores.get("Formatted Score", "N/A")
        fit_summary = ats_scores.get("Final Thoughts", "N/A")
        language_analysis_full = ats_scores.get("Language Analysis", "N/A")

        missing_keywords_raw = ats_scores.get("Missing Keywords", "N/A")
        missing_skills_raw = ats_scores.get("Missing Skills", "N/A")
        missing_keywords = [kw.strip() for kw in missing_keywords_raw.split(",") if kw.strip()] if missing_keywords_raw != "N/A" else []
        missing_skills = [sk.strip() for sk in missing_skills_raw.split(",") if sk.strip()] if missing_skills_raw != "N/A" else []

        domain = db_manager.detect_domain_llm(
            job_title,
            job_description,
            session=st.session_state  # ✅ pass the Groq API key from session
        )

        bias_flag = "🔴 High Bias" if bias_score > 0.6 else "🟢 Fair"
        ats_flag = "⚠️ Low ATS" if ats_score < 50 else "✅ Good ATS"

        # ✅ Store everything in session state
        st.session_state.resume_data.append({
            "Resume Name": uploaded_file.name,
            "Candidate Name": candidate_name,
            "ATS Report": ats_result,
            "ATS Match %": ats_score,
            "Formatted Score": formatted_score,
            "Education Score": edu_score,
            "Experience Score": exp_score,
            "Skills Score": skills_score,
            "Language Score": lang_score,
            "Keyword Score": keyword_score,
            "Education Analysis": ats_scores.get("Education Analysis", ""),
            "Experience Analysis": ats_scores.get("Experience Analysis", ""),
            "Skills Analysis": ats_scores.get("Skills Analysis", ""),
            "Language Analysis": language_analysis_full,
            "Keyword Analysis": ats_scores.get("Keyword Analysis", ""),
            "Final Thoughts": fit_summary,
            "Missing Keywords": missing_keywords,
            "Missing Skills": missing_skills,
            "Bias Score (0 = Fair, 1 = Biased)": bias_score,
            "Bias Status": bias_flag,
            "Masculine Words": masc_count,
            "Feminine Words": fem_count,
            "Detected Masculine Words": detected_masc,
            "Detected Feminine Words": detected_fem,
            "Text Preview": full_text[:300] + "...",
            "Highlighted Text": highlighted_text,
            "Rewritten Text": rewritten_text,
            "Domain": domain
        })

        insert_candidate(
            (
                uploaded_file.name,
                candidate_name,
                ats_score,
                edu_score,
                exp_score,
                skills_score,
                lang_score,
                keyword_score,
                bias_score
            ),
            job_title=job_title,
            job_description=job_description
        )

        st.session_state.processed_files.add(uploaded_file.name)

        # ✅ IMPROVED: Smoother success animation with better transitions
        SUCCESS_HTML = """
        <style>
        .success-overlay {
            position: fixed;
            top: 0; left: 0;
            width: 100vw; height: 100vh;
            background: linear-gradient(135deg, #0b0c10 0%, #1a1c29 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            z-index: 9999;
            animation: fadeIn 0.5s ease-out;
        }
        
        @keyframes fadeIn {
            0% { opacity: 0; }
            100% { opacity: 1; }
        }
        
        .success-circle {
            width: 140px;
            height: 140px;
            border: 3px solid #00bfff;
            border-radius: 50%;
            position: relative;
            display: flex;
            align-items: center;
            justify-content: center;
            background: radial-gradient(circle, rgba(0,191,255,0.1) 0%, rgba(0,191,255,0.05) 50%, transparent 100%);
            animation: successPulse 2s ease-in-out infinite;
        }
        
        @keyframes successPulse {
            0%, 100% { 
                transform: scale(1);
                box-shadow: 0 0 20px rgba(0,191,255,0.3);
            }
            50% { 
                transform: scale(1.05);
                box-shadow: 0 0 30px rgba(0,191,255,0.6);
            }
        }
        
        .success-checkmark {
            font-size: 48px;
            color: #00ff7f;
            animation: checkmarkPop 0.8s ease-out;
        }
        
        @keyframes checkmarkPop {
            0% { transform: scale(0) rotate(-45deg); opacity: 0; }
            50% { transform: scale(1.2) rotate(-10deg); opacity: 0.8; }
            100% { transform: scale(1) rotate(0deg); opacity: 1; }
        }
        
        .success-text {
            margin-top: 25px;
            font-family: 'Orbitron', 'Segoe UI', sans-serif;
            font-size: 20px;
            font-weight: 600;
            color: #00bfff;
            text-shadow: 0 0 10px rgba(0,191,255,0.5);
            animation: textSlideUp 0.8s ease-out 0.3s both;
        }
        
        @keyframes textSlideUp {
            0% { transform: translateY(20px); opacity: 0; }
            100% { transform: translateY(0); opacity: 1; }
        }
        
        .success-subtitle {
            margin-top: 10px;
            font-size: 14px;
            color: #8e9aaf;
            animation: textSlideUp 0.8s ease-out 0.5s both;
        }
        </style>
        
        <div class="success-overlay">
            <div class="success-circle">
                <div class="success-checkmark">✓</div>
            </div>
            <div class="success-text">Scan Complete!</div>
            <div class="success-subtitle">Resume analysis ready</div>
        </div>
        """
        
        # Clear scanner and show success animation
        scanner_placeholder.empty()
        success_placeholder = st.empty()
        success_placeholder.markdown(SUCCESS_HTML, unsafe_allow_html=True)

        # ⏳ Shorter delay for better UX, then clear and rerun
        time.sleep(3)
        success_placeholder.empty()
        st.rerun()

    # ✅ Optional vectorstore setup
    if all_text:
        st.session_state.vectorstore = setup_vectorstore(all_text)
        st.session_state.chain = create_chain(st.session_state.vectorstore)

# 🔄 Developer Reset Button
with tab1:
    if st.button("🔄 Refresh view"):
        st.session_state.processed_files.clear()
        st.session_state.resume_data.clear()

        # Temporary placeholder for sliding success message
        msg_placeholder = st.empty()
        msg_placeholder.markdown("""
        <div class='slide-message success-msg'>
            <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
              stroke-width="2" viewBox="0 0 24 24"><path d="M5 13l4 4L19 7"/></svg>
            ✅ Cleared uploaded resume history. You can re-upload now.
        </div>
        """, unsafe_allow_html=True)

        # Wait 3 seconds then clear message
        time.sleep(3)
        msg_placeholder.empty()

def generate_resume_report_html(resume):
    candidate_name = resume.get('Candidate Name', 'Not Found')
    resume_name = resume.get('Resume Name', 'Unknown')
    rewritten_text = resume.get('Rewritten Text', '').replace("\n", "<br/>")

    masculine_words_list = resume.get("Detected Masculine Words", [])
    masculine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in masculine_words_list
    ) if masculine_words_list else "<i>None detected.</i>"

    feminine_words_list = resume.get("Detected Feminine Words", [])
    feminine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in feminine_words_list
    ) if feminine_words_list else "<i>None detected.</i>"

    ats_report_html = resume.get("ATS Report", "").replace("\n", "<br/>")

    def style_analysis(analysis, fallback="N/A"):
        if not analysis or analysis == "N/A":
            return f"<p><i>{fallback}</i></p>"

        if "**Score:**" in analysis:
            parts = analysis.split("**Score:**")
            rest = parts[1].split("**", 1)
            score_text = rest[0].strip()
            remaining = rest[1].strip() if len(rest) > 1 else ""
            return f"<p><b>Score:</b> {score_text}</p><p>{remaining}</p>"
        else:
            return f"<p>{analysis}</p>"

    edu_analysis = style_analysis(resume.get("Education Analysis", "").replace("\n", "<br/>"))
    exp_analysis = style_analysis(resume.get("Experience Analysis", "").replace("\n", "<br/>"))
    skills_analysis = style_analysis(resume.get("Skills Analysis", "").replace("\n", "<br/>"))
    keyword_analysis = style_analysis(resume.get("Keyword Analysis", "").replace("\n", "<br/>"))
    final_thoughts = resume.get("Final Thoughts", "N/A").replace("\n", "<br/>")

    lang_analysis_raw = resume.get("Language Analysis", "").replace("\n", "<br/>")
    lang_analysis = f"<div>{lang_analysis_raw}</div>" if lang_analysis_raw else "<p><i>No language analysis available.</i></p>"

    ats_match = resume.get('ATS Match %', 'N/A')
    edu_score = resume.get('Education Score', 'N/A')
    exp_score = resume.get('Experience Score', 'N/A')
    skills_score = resume.get('Skills Score', 'N/A')
    lang_score = resume.get('Language Score', 'N/A')
    keyword_score = resume.get('Keyword Score', 'N/A')
    masculine_count = len(masculine_words_list)
    feminine_count = len(feminine_words_list)
    bias_score = resume.get('Bias Score (0 = Fair, 1 = Biased)', 'N/A')

    return f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Helvetica, sans-serif;
                font-size: 12pt;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.4em;
            }}
            li {{
                margin-bottom: 5px;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 12px;
                border-left: 4px solid #666;
            }}
            .box {{
                padding: 10px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;
            }}
        </style>
    </head>
    <body>

    <h1>Resume Analysis Report</h1>

    <h2>Candidate: {candidate_name}</h2>
    <p><b>Resume File:</b> {resume_name}</p>

    <h2>ATS Evaluation</h2>
    <table>
        <tr><td><b>ATS Match</b></td><td>{ats_match}%</td></tr>
        <tr><td><b>Education</b></td><td>{edu_score}</td></tr>
        <tr><td><b>Experience</b></td><td>{exp_score}</td></tr>
        <tr><td><b>Skills</b></td><td>{skills_score}</td></tr>
        <tr><td><b>Language</b></td><td>{lang_score}</td></tr>
        <tr><td><b>Keyword</b></td><td>{keyword_score}</td></tr>
    </table>

    <div class="section-title">ATS Report</div>
    <div class="box">{ats_report_html}</div>

    <div class="section-title">Education Analysis</div>
    <div class="box">{edu_analysis}</div>

    <div class="section-title">Experience Analysis</div>
    <div class="box">{exp_analysis}</div>

    <div class="section-title">Skills Analysis</div>
    <div class="box">{skills_analysis}</div>

    <div class="section-title">Language Analysis</div>
    <div class="box">{lang_analysis}</div>

    <div class="section-title">Keyword Analysis</div>
    <div class="box">{keyword_analysis}</div>

    <div class="section-title">Final Thoughts</div>
    <div class="box">{final_thoughts}</div>

    <h2>Gender Bias Analysis</h2>
    <table>
        <tr><td><b>Masculine Words</b></td><td>{masculine_count}</td></tr>
        <tr><td><b>Feminine Words</b></td><td>{feminine_count}</td></tr>
        <tr><td><b>Bias Score (0 = Fair, 1 = Biased)</b></td><td>{bias_score}</td></tr>
    </table>

    <div class="section-title">Masculine Words Detected</div>
    <div class="box">{masculine_words}</div>

    <div class="section-title">Feminine Words Detected</div>
    <div class="box">{feminine_words}</div>

    <h2>Rewritten Bias-Free Resume</h2>
    <div class="box">{rewritten_text}</div>

    </body>
    </html>
    """

# === TAB 1: Dashboard ===
with tab1:
    resume_data = st.session_state.get("resume_data", [])

    if resume_data:
        # ✅ Calculate total counts safely
        total_masc = sum(len(r.get("Detected Masculine Words", [])) for r in resume_data)
        total_fem = sum(len(r.get("Detected Feminine Words", [])) for r in resume_data)
        avg_bias = round(np.mean([r.get("Bias Score (0 = Fair, 1 = Biased)", 0) for r in resume_data]), 2)
        total_resumes = len(resume_data)

        st.markdown("### 📊 Summary Statistics")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📄 Resumes Uploaded", total_resumes)
        with col2:
            st.metric("🔎 Avg. Bias Score", avg_bias)
        with col3:
            st.metric("🔵 Total Masculine Words", total_masc)
        with col4:
            st.metric("🔴 Total Feminine Words", total_fem)

        st.markdown("### 🗂️ Resumes Overview")
        df = pd.DataFrame(resume_data)

        # ✅ Add calculated count columns safely
        df["Masculine Words Count"] = df["Detected Masculine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)
        df["Feminine Words Count"] = df["Detected Feminine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)

        overview_cols = [
            "Resume Name", "Candidate Name", "ATS Match %", "Education Score",
            "Experience Score", "Skills Score", "Language Score", "Keyword Score",
            "Bias Score (0 = Fair, 1 = Biased)", "Masculine Words Count", "Feminine Words Count"
        ]

        st.dataframe(df[overview_cols], use_container_width=True)

        st.markdown("### 📊 Visual Analysis")
        chart_tab1, chart_tab2 = st.tabs(["📉 Bias Score Chart", "⚖ Gender-Coded Words"])
        with chart_tab1:
            st.subheader("Bias Score Comparison Across Resumes")
            st.bar_chart(df.set_index("Resume Name")[["Bias Score (0 = Fair, 1 = Biased)"]])
        with chart_tab2:
            st.subheader("Masculine vs Feminine Word Usage")
            fig, ax = plt.subplots(figsize=(10, 5))
            index = np.arange(len(df))
            bar_width = 0.35
            ax.bar(index, df["Masculine Words Count"], bar_width, label="Masculine", color="#3498db")
            ax.bar(index + bar_width, df["Feminine Words Count"], bar_width, label="Feminine", color="#e74c3c")
            ax.set_xlabel("Resumes", fontsize=12)
            ax.set_ylabel("Word Count", fontsize=12)
            ax.set_title("Gender-Coded Word Usage per Resume", fontsize=14)
            ax.set_xticks(index + bar_width / 2)
            ax.set_xticklabels(df["Resume Name"], rotation=45, ha='right')
            ax.legend()
            st.pyplot(fig)

        st.markdown("### 📝 Detailed Resume Reports")
        for resume in resume_data:
            candidate_name = resume.get("Candidate Name", "Not Found")
            resume_name = resume.get("Resume Name", "Unknown")
            missing_keywords = resume.get("Missing Keywords", [])
            missing_skills = resume.get("Missing Skills", [])

            with st.expander(f"📄 {resume_name} | {candidate_name}"):
                st.markdown(f"### 📊 ATS Evaluation for: **{candidate_name}**")
                score_col1, score_col2, score_col3 = st.columns(3)
                with score_col1:
                    st.metric("📈 Overall Match", f"{resume.get('ATS Match %', 'N/A')}%")
                with score_col2:
                    st.metric("🏆 Formatted Score", resume.get("Formatted Score", "N/A"))
                with score_col3:
                    st.metric("🧠 Language Quality", f"{resume.get('Language Score', 'N/A')} / {lang_weight}")

                col_a, col_b, col_c, col_d = st.columns(4)
                with col_a:
                    st.metric("🎓 Education Score", f"{resume.get('Education Score', 'N/A')} / {edu_weight}")
                with col_b:
                    st.metric("💼 Experience Score", f"{resume.get('Experience Score', 'N/A')} / {exp_weight}")
                with col_c:
                    st.metric("🛠 Skills Score", f"{resume.get('Skills Score', 'N/A')} / {skills_weight}")
                with col_d:
                    st.metric("🔍 Keyword Score", f"{resume.get('Keyword Score', 'N/A')} / {keyword_weight}")

                # Fit summary
                st.markdown("### 📝 Fit Summary")
                st.write(resume.get('Final Thoughts', 'N/A'))

                # ATS Report
                if resume.get("ATS Report"):
                    st.markdown("### 📋 ATS Evaluation Report")
                    st.markdown(resume["ATS Report"], unsafe_allow_html=True)

                # ATS Chart
                st.markdown("### 📊 ATS Score Breakdown Chart")
                ats_df = pd.DataFrame({
                    'Component': ['Education', 'Experience', 'Skills', 'Language', 'Keywords'],
                    'Score': [
                        resume.get("Education Score", 0),
                        resume.get("Experience Score", 0),
                        resume.get("Skills Score", 0),
                        resume.get("Language Score", 0),
                        resume.get("Keyword Score", 0)
                    ]
                })
                ats_chart = alt.Chart(ats_df).mark_bar().encode(
                    x=alt.X('Component', sort=None),
                    y=alt.Y('Score', scale=alt.Scale(domain=[0, 50])),
                    color='Component',
                    tooltip=['Component', 'Score']
                ).properties(
                    title="ATS Evaluation Breakdown",
                    width=600,
                    height=300
                )
                st.altair_chart(ats_chart, use_container_width=True)

                # 🔷 Detailed ATS Analysis Cards
                st.markdown("### 🔍 Detailed ATS Section Analyses")
                for section_title, key in [
                    ("🏫 Education Analysis", "Education Analysis"),
                    ("💼 Experience Analysis", "Experience Analysis"),
                    ("🛠 Skills Analysis", "Skills Analysis"),
                    ("🗣 Language Quality Analysis", "Language Analysis"),
                    ("🔑 Keyword Analysis", "Keyword Analysis"),
                    ("✅ Final Thoughts", "Final Thoughts")
                ]:
                    analysis_content = resume.get(key, "N/A")
                    if "**Score:**" in analysis_content:
                        parts = analysis_content.split("**Score:**")
                        rest = parts[1].split("**", 1)
                        score_text = rest[0].strip()
                        remaining = rest[1].strip() if len(rest) > 1 else ""
                        formatted_score = f"<div style='background:#4c1d95;color:white;padding:8px;border-radius:6px;margin-bottom:5px;'><b>Score:</b> {score_text}</div>"
                        analysis_html = formatted_score + f"<p>{remaining}</p>"
                    else:
                        analysis_html = f"<p>{analysis_content}</p>"

                    st.markdown(f"""
<div style="background:#5b3cc4; color:white; padding:10px; border-radius:6px;">
  <h3>{section_title}</h3>
</div>
<div style="background:#2d2d3a; color:white; padding:10px; border-radius:6px;">
{analysis_html}
</div>
""", unsafe_allow_html=True)

                st.divider()

                detail_tab1, detail_tab2 = st.tabs(["🔎 Bias Analysis", "✅ Rewritten Resume"])

                with detail_tab1:
                    st.markdown("#### Bias-Highlighted Original Text")
                    st.markdown(resume["Highlighted Text"], unsafe_allow_html=True)

                    st.markdown("### 📌 Gender-Coded Word Counts:")
                    bias_col1, bias_col2 = st.columns(2)

                    with bias_col1:
                        st.metric("🔵 Masculine Words", len(resume["Detected Masculine Words"]))
                        if resume["Detected Masculine Words"]:
                            st.markdown("### 📚 Detected Masculine Words with Context:")
                            for item in resume["Detected Masculine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.write(f"🔵 **{word}**: {sentence}", unsafe_allow_html=True)
                        else:
                            st.info("No masculine words detected.")

                    with bias_col2:
                        st.metric("🔴 Feminine Words", len(resume["Detected Feminine Words"]))
                        if resume["Detected Feminine Words"]:
                            st.markdown("### 📚 Detected Feminine Words with Context:")
                            for item in resume["Detected Feminine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.write(f"🔴 **{word}**: {sentence}", unsafe_allow_html=True)
                        else:
                            st.info("No feminine words detected.")

                with detail_tab2:
                    st.markdown("#### ✨ Bias-Free Rewritten Resume")
                    st.write(resume["Rewritten Text"])
                    docx_file = generate_docx(resume["Rewritten Text"])
                    st.download_button(
                        label="📥 Download Bias-Free Resume (.docx)",
                        data=docx_file,
                        file_name=f"{resume['Resume Name'].split('.')[0]}_bias_free.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"download_docx_{resume['Resume Name']}"
                    )
                    html_report = generate_resume_report_html(resume)
                    
                    pdf_file = html_to_pdf_bytes(html_report)
                    st.download_button(
                    label="📄 Download Full Analysis Report (.pdf)",
                    data=pdf_file,
                    file_name=f"{resume['Resume Name'].split('.')[0]}_report.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"download_pdf_{resume['Resume Name']}"
                    )               

    else:           
        st.warning("⚠️ Please upload resumes to view dashboard analytics.")

# ---------------- Sidebar (ONLY in Tab 2) ----------------
from xhtml2pdf import pisa
from io import BytesIO

def html_to_pdf_bytes(html_string):
    styled_html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: 400mm 297mm;  /* Original custom large page size */
                margin-top: 10mm;
                margin-bottom: 10mm;
                margin-left: 10mm;
                margin-right: 10mm;
            }}
            body {{
                font-size: 14pt;
                font-family: "Segoe UI", "Helvetica", sans-serif;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2, h3 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 10px;
            }}
            .box {{
                padding: 8px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;  /* More elegant than full border */
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.5em;
            }}
            li {{
                margin-bottom: 5px;
            }}
        </style>
    </head>
    <body>
        {html_string}
    </body>
    </html>
    """

    pdf_io = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_io)
    pdf_io.seek(0)
    return pdf_io

def render_template_default(session_state, profile_img_html=""):
    """Default professional template - keeps the exact same design as before"""
    
    # Enhanced SKILLS with professional, muted colors
    skills_html = "".join(
        f"""
        <div style='display:inline-block; 
                    background: linear-gradient(135deg, #e2e8f0 0%, #cbd5e1 100%);
                    color: #334155; 
                    padding: 10px 18px; 
                    margin: 8px 8px 8px 0; 
                    border-radius: 25px; 
                    font-size: 14px; 
                    font-weight: 600;
                    box-shadow: 0 2px 8px rgba(148, 163, 184, 0.2);
                    transition: all 0.3s ease;
                    text-shadow: none;
                    border: 1px solid rgba(148, 163, 184, 0.3);'>
            {s.strip()}
        </div>
        """
        for s in session_state['skills'].split(',')
        if s.strip()
    )

    # Enhanced LANGUAGES with soft, professional design
    languages_html = "".join(
        f"""
        <div style='display:inline-block; 
                    background: linear-gradient(135deg, #f1f5f9 0%, #e2e8f0 100%);
                    color: #475569; 
                    padding: 10px 18px; 
                    margin: 8px 8px 8px 0; 
                    border-radius: 25px; 
                    font-size: 14px; 
                    font-weight: 600;
                    box-shadow: 0 2px 8px rgba(100, 116, 139, 0.15);
                    transition: all 0.3s ease;
                    text-shadow: none;
                    border: 1px solid rgba(148, 163, 184, 0.3);'>
            {lang.strip()}
        </div>
        """
        for lang in session_state['languages'].split(',')
        if lang.strip()
    )

    # Enhanced INTERESTS with subtle colors
    interests_html = "".join(
        f"""
        <div style='display:inline-block; 
                    background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
                    color: #0f172a; 
                    padding: 10px 18px; 
                    margin: 8px 8px 8px 0; 
                    border-radius: 25px; 
                    font-size: 14px; 
                    font-weight: 600;
                    box-shadow: 0 2px 8px rgba(14, 165, 233, 0.1);
                    transition: all 0.3s ease;
                    text-shadow: none;
                    border: 1px solid rgba(186, 230, 253, 0.5);'>
            {interest.strip()}
        </div>
        """
        for interest in session_state['interests'].split(',')
        if interest.strip()
    )

    # Enhanced SOFT SKILLS with warm but professional styling
    Softskills_html = "".join(
        f"""
        <div style='display:inline-block; 
                    background: linear-gradient(135deg, #fefce8 0%, #fef3c7 100%);
                    color: #451a03; 
                    padding: 10px 20px; 
                    margin: 8px 8px 8px 0; 
                    border-radius: 25px; 
                    font-size: 14px; 
                    font-family: "Segoe UI", sans-serif; 
                    font-weight: 600;
                    box-shadow: 0 2px 8px rgba(217, 119, 6, 0.1);
                    transition: all 0.3s ease;
                    border: 1px solid rgba(254, 215, 170, 0.6);'>
            {skill.strip().title()}
        </div>
        """
        for skill in session_state['Softskills'].split(',')
        if skill.strip()
    )

    # Enhanced EXPERIENCE with professional, subtle design
    experience_html = ""
    for exp in session_state.experience_entries:
        if exp["company"] or exp["title"]:
            # Handle paragraphs and single line breaks
            description_lines = [line.strip() for line in exp["description"].strip().split("\n\n")]
            description_html = "".join(
                f"<div style='margin-bottom: 10px; line-height: 1.6;'>{line.replace(chr(10), '<br>')}</div>"
                for line in description_lines if line
            )

            experience_html += f"""
            <div style='
                margin-bottom: 24px;
                padding: 20px;
                border-radius: 12px;
                background: linear-gradient(145deg, #fafafa 0%, #f4f4f5 100%);
                box-shadow: 
                    0 4px 12px rgba(0, 0, 0, 0.05),
                    0 1px 3px rgba(0, 0, 0, 0.1);
                font-family: "Inter", "Segoe UI", sans-serif;
                color: #374151;
                line-height: 1.6;
                border: 1px solid rgba(229, 231, 235, 0.8);
                position: relative;
                overflow: hidden;
            '>
                <!-- Subtle accent bar -->
                <div style='
                    position: absolute;
                    top: 0;
                    left: 0;
                    right: 0;
                    height: 3px;
                    background: linear-gradient(90deg, #6b7280, #9ca3af);
                '></div>
                
                <!-- Header Card -->
                <div style='
                    background: rgba(255, 255, 255, 0.8);
                    border-radius: 8px;
                    padding: 14px 18px;
                    margin-bottom: 12px;
                    border: 1px solid rgba(229, 231, 235, 0.6);
                '>
                    <div style='
                        display: flex;
                        justify-content: space-between;
                        align-items: center;
                        font-weight: 700;
                        font-size: 18px;
                        margin-bottom: 6px;
                        color: #1f2937;
                        width: 100%;
                    '>
                        <div style='display: flex; align-items: center;'>
                            <div style='
                                width: 6px; 
                                height: 6px; 
                                background: #6b7280;
                                border-radius: 50%; 
                                margin-right: 12px;
                            '></div>
                            <span>{exp['company']}</span>
                        </div>
                        <div style='
                            display: inline-flex;
                            align-items: center;
                            gap: 6px;
                            background: linear-gradient(135deg, #f9fafb, #f3f4f6);
                            color: #374151;
                            padding: 5px 14px;
                            border-radius: 16px;
                            font-size: 14px;
                            font-weight: 600;
                            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                            border: 1px solid rgba(209, 213, 219, 0.5);
                        '>
                            <svg xmlns="http://www.w3.org/2000/svg" width="14" height="14" fill="currentColor" viewBox="0 0 16 16">
                                <path d="M3.5 0a.5.5 0 0 1 .5.5V1h8V.5a.5.5 0 0 1 1 0V1h1a2 2 0 0 1 2 2v11a2 2 0 0 1-2 2H2a2 2 0 0 1-2-2V3a2 2 0 0 1 2-2h1V.5a.5.5 0 0 1 .5-.5zM1 4v10a1 1 0 0 0 1 1h12a1 1 0 0 0 1-1V4H1z"/>
                            </svg>
                            <span>{exp['duration']}</span>
                        </div>
                    </div>

                    <div style='
                        display: flex;
                        align-items: center;
                        font-size: 16px;
                        font-weight: 600;
                        color: #4b5563;
                    '>
                        <div style='
                            width: 4px; 
                            height: 4px; 
                            background: #6b7280;
                            border-radius: 50%; 
                            margin-right: 10px;
                        '></div>
                        <span>{exp['title']}</span>
                    </div>
                </div>

                <!-- Description -->
                <div style='
                    font-size: 15px;
                    font-weight: 500;
                    color: #374151;
                    line-height: 1.7;
                    padding-left: 8px;
                '>
                    <div style='
                        border-left: 2px solid #d1d5db;
                        padding-left: 16px;
                        margin-left: 8px;
                    '>
                        {description_html}
                    </div>
                </div>
            </div>
            """

    # Convert experience to list if multiple lines
    # Escape HTML and convert line breaks
    summary_html = session_state['summary'].replace('\n', '<br>')

    # Enhanced EDUCATION with professional styling
    education_html = ""
    for edu in session_state.education_entries:
        if edu.get("institution") or edu.get("details"):
            degree_text = ""
            if edu.get("degree"):
                degree_val = edu["degree"]
                if isinstance(degree_val, list):
                    degree_val = ", ".join(degree_val)
                degree_text = f"""
                <div style='
                    display: flex; 
                    align-items: center; 
                    font-size: 15px; 
                    color: #374151; 
                    margin-bottom: 8px;
                    font-weight: 600;
                '>
                    <div style='
                        width: 4px; 
                        height: 4px; 
                        background: #6b7280;
                        border-radius: 50%; 
                        margin-right: 10px;
                    '></div>
                    <b>{degree_val}</b>
                </div>
                """

            # Education Card
            education_html += f"""
            <div style='
                margin-bottom: 26px;
                padding: 22px 26px;
                border-radius: 12px;
                background: linear-gradient(145deg, #f9fafb 0%, #f3f4f6 100%);
                box-shadow: 
                    0 4px 12px rgba(0, 0, 0, 0.06),
                    0 1px 3px rgba(0, 0, 0, 0.08);
                font-family: "Inter", "Segoe UI", sans-serif;
                color: #1f2937;
                line-height: 1.6;
                border: 1px solid #e5e7eb;
                position: relative;
                overflow: hidden;
            '>
                <!-- Subtle accent bar -->
                <div style='
                    position: absolute;
                    top: 0;
                    left: 0;
                    right: 0;
                    height: 3px;
                    background: linear-gradient(90deg, #6b7280, #9ca3af);
                '></div>

                <div style='
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    font-size: 18px;
                    font-weight: 700;
                    margin-bottom: 12px;
                    width: 100%;
                    color: #111827;
                '>
                    <div style='display: flex; align-items: center;'>
                        <div style='
                            width: 6px; 
                            height: 6px; 
                            background: #6b7280;
                            border-radius: 50%; 
                            margin-right: 12px;
                        '></div>
                        <span>{edu.get('institution', '')}</span>
                    </div>
                    <div style='
                        display: flex;
                        align-items: center;
                        gap: 6px;
                        background: rgba(255, 255, 255, 0.7);
                        color: #374151;
                        padding: 6px 16px;
                        border-radius: 16px;
                        font-weight: 600;
                        font-size: 14px;
                        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                        border: 1px solid #d1d5db;
                    '>
                        <!-- Inline SVG Calendar Icon -->
                        <svg xmlns="http://www.w3.org/2000/svg" 
                            fill="none" viewBox="0 0 24 24" 
                            stroke="currentColor" width="16" height="16">
                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                                d="M8 7V3m8 4V3m-9 8h10M5 21h14a2 2 0 002-2V7a2 
                                2 0 00-2-2H5a2 2 0 00-2 2v12a2 2 0 002 2z" />
                        </svg>
                        {edu.get('year', '')}
                    </div>
                </div>
                {degree_text}
                <div style='
                    font-size: 14px; 
                    font-style: italic;
                    color: #374151;
                    line-height: 1.6;
                    padding-left: 18px;
                    border-left: 2px solid #9ca3af;
                '>
                    {edu.get('details', '')}
                </div>
            </div>
            """

    # Enhanced PROJECTS with professional card design
    projects_html = ""
    for proj in session_state.project_entries:
        if proj.get("title") or proj.get("description"):
            tech_val = proj.get("tech")
            if isinstance(tech_val, list):
                tech_val = ", ".join(tech_val)
            tech_text = f"""
            <div style='
                display: flex; 
                align-items: center; 
                font-size: 14px; 
                color: #374151; 
                margin-bottom: 12px;
                font-weight: 600;
                background: rgba(255, 255, 255, 0.7);
                padding: 8px 16px;
                border-radius: 8px;
                border: 1px solid rgba(229, 231, 235, 0.6);
            '>
                <div style='
                    width: 4px; 
                    height: 4px; 
                    background: #6b7280;
                    border-radius: 50%; 
                    margin-right: 10px;
                '></div>
                <b>Technologies:</b>&nbsp;&nbsp;{tech_val if tech_val else ''}
            </div>
            """ if tech_val else ""

            description_items = ""
            if proj.get("description"):
                description_lines = [line.strip() for line in proj["description"].splitlines() if line.strip()]
                description_items = "".join(f"<li style='margin-bottom: 6px; line-height: 1.6;'>{line}</li>" for line in description_lines)

            projects_html += f"""
            <div style='
                margin-bottom: 30px;
                padding: 26px;
                border-radius: 12px;
                background: linear-gradient(145deg, #f8fafc 0%, #f1f5f9 100%);
                box-shadow: 
                    0 4px 12px rgba(100, 116, 139, 0.1),
                    0 1px 3px rgba(0, 0, 0, 0.1);
                font-family: "Inter", "Segoe UI", sans-serif;
                color: #334155;
                line-height: 1.7;
                border: 1px solid rgba(203, 213, 225, 0.5);
                position: relative;
                overflow: hidden;
            '>
                <!-- Subtle accent bar -->
                <div style='
                    position: absolute;
                    top: 0;
                    left: 0;
                    right: 0;
                    height: 3px;
                    background: linear-gradient(90deg, #64748b, #94a3b8);
                '></div>

                <div style='
                    font-size: 19px;
                    font-weight: 700;
                    margin-bottom: 16px;
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    color: #1e293b;
                    width: 100%;
                '>
                    <div style='display: flex; align-items: center;'>
                        <div style='
                            width: 6px; 
                            height: 6px; 
                            background: #64748b;
                            border-radius: 50%; 
                            margin-right: 12px;
                        '></div>
                        <span>{proj.get('title', '')}</span>
                    </div>
                    <div style='
                        display: flex;
                        align-items: center;
                        gap: 6px;
                        background: linear-gradient(135deg, #f1f5f9, #e2e8f0);
                        color: #334155;
                        padding: 8px 18px;
                        border-radius: 16px;
                        font-weight: 600;
                        font-size: 14px;
                        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                        border: 1px solid rgba(203, 213, 225, 0.6);
                    '>
                        <!-- Inline SVG Clock Icon -->
                        <svg xmlns="http://www.w3.org/2000/svg" 
                            fill="none" viewBox="0 0 24 24" 
                            stroke="currentColor" width="16" height="16">
                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                                d="M12 8v4l3 3m6-3a9 9 0 11-18 0 
                                   9 9 0 0118 0z" />
                        </svg>
                        {proj.get('duration', '')}
                    </div>
                </div>
                {tech_text}
                <div style='
                    font-size: 15px; 
                    color: #334155;
                    background: rgba(255, 255, 255, 0.6);
                    padding: 18px;
                    border-radius: 8px;
                    border: 1px solid rgba(229, 231, 235, 0.6);
                '>
                    <div style='
                        font-weight: 600; 
                        margin-bottom: 12px;
                        color: #1e293b;
                        display: flex;
                        align-items: center;
                    '>
                        <div style='
                            width: 4px; 
                            height: 4px; 
                            background: #64748b;
                            border-radius: 50%; 
                            margin-right: 10px;
                        '></div>
                        Description:
                    </div>
                    <ul style='
                        margin-top: 8px; 
                        padding-left: 24px; 
                        color: #334155;
                        list-style-type: none;
                    '>
                        {description_items}
                    </ul>
                </div>
            </div>
            """

    # Enhanced PROJECT LINKS with professional styling
    project_links_html = ""
    if session_state.project_links:
        project_links_html = """
        <div style='margin-bottom: 20px;'>
            <h4 class='section-title' style='
                color: #374151;
                font-size: 20px;
                margin-bottom: 8px;
                display: flex;
                align-items: center;
                padding-bottom: 4px;
            '>
                <div style='
                    width: 6px; 
                    height: 6px; 
                    background: #6b7280;
                    border-radius: 50%; 
                    margin-right: 12px;
                '></div>
                Project Links
            </h4>
        </div>
        """ + "".join(
            f"""
            <div style='
                background: linear-gradient(135deg, #f9fafb 0%, #f3f4f6 100%);
                padding: 14px 20px;
                border-radius: 8px;
                margin-bottom: 12px;
                border: 1px solid rgba(209, 213, 219, 0.6);
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
            '>
                <div style='
                    width: 4px; 
                    height: 4px; 
                    background: #6b7280;
                    border-radius: 50%; 
                    display: inline-block;
                    margin-right: 12px;
                    vertical-align: middle;
                '></div>
                <a href="{link}" style='
                    color: #374151; 
                    font-weight: 600; 
                    text-decoration: none;
                    font-size: 15px;
                '>🔗 Project {i+1}</a>
            </div>
            """
            for i, link in enumerate(session_state.project_links)
        )

    # Enhanced CERTIFICATES with professional design
    certificate_links_html = ""
    if session_state.certificate_links:
        certificate_links_html = """
        <h4 class='section-title' style='
            color: #374151;
            font-size: 20px;
            margin-bottom: 16px;
            display: flex;
            align-items: center;
        '>
            <div style='
                width: 6px; 
                height: 6px; 
                background: #6b7280;
                border-radius: 50%; 
                margin-right: 12px;
            '></div>
            Certificates
        </h4>
        """
        for cert in session_state.certificate_links:
            if cert["name"] and cert["link"]:
                description = cert.get('description', '').replace('\n', '<br>')
                name = cert['name']
                link = cert['link']
                duration = cert.get('duration', '')

                card_html = f"""
                <div style='
                    background: linear-gradient(145deg, #f9fafb 0%, #f3f4f6 100%);
                    padding: 24px 28px;
                    border-radius: 12px;
                    margin-bottom: 26px;
                    box-shadow: 
                        0 4px 12px rgba(107, 114, 128, 0.08),
                        0 1px 3px rgba(0, 0, 0, 0.08);
                    font-family: "Inter", "Segoe UI", sans-serif;
                    color: #374151;
                    position: relative;
                    line-height: 1.7;
                    border: 1px solid rgba(209, 213, 219, 0.6);
                    overflow: hidden;
                '>
                    <!-- Accent bar -->
                    <div style='
                        position: absolute;
                        top: 0;
                        left: 0;
                        right: 0;
                        height: 3px;
                        background: linear-gradient(90deg, #6b7280, #9ca3af);
                    '></div>

                    <!-- Duration Badge -->
                    <div style='
                        position: absolute;
                        top: 20px;
                        right: 28px;
                        font-size: 13px;
                        font-weight: 600;
                        color: #374151;
                        background: linear-gradient(135deg, #ffffff, #f9fafb);
                        padding: 8px 14px;
                        border-radius: 16px;
                        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.08);
                        border: 1px solid rgba(209, 213, 219, 0.6);
                        display: flex;
                        align-items: center;
                        gap: 6px;
                    '>
                        <!-- Inline SVG clock icon -->
                        <svg xmlns="http://www.w3.org/2000/svg" 
                            fill="none" viewBox="0 0 24 24" 
                            stroke="currentColor" width="14" height="14">
                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                                d="M12 6v6l4 2m6-2a10 10 0 11-20 0 10 10 0 0120 0z"/>
                        </svg>
                        {duration}
                    </div>

                    <!-- Certificate Title -->
                    <div style='
                        font-size: 18px;
                        font-weight: 700;
                        color: #111827;
                        margin-bottom: 12px;
                        margin-right: 120px;
                        display: flex;
                        align-items: center;
                    '>
                        <div style='
                            width: 6px; 
                            height: 6px; 
                            background: #6b7280;
                            border-radius: 50%; 
                            margin-right: 12px;
                        '></div>
                        <a href="{link}" target="_blank" style='
                            color: #111827;
                            text-decoration: none;
                            transition: color 0.3s ease;
                        '>{name}</a>
                    </div>

                    <!-- Description -->
                    <div style='
                        font-size: 15px;
                        color: #374151;
                        background: rgba(255, 255, 255, 0.8);
                        padding: 16px;
                        border-radius: 8px;
                        border: 1px solid rgba(209, 213, 219, 0.6);
                        line-height: 1.6;
                    '>
                        <div style='
                            display: flex;
                            align-items: flex-start;
                            margin-bottom: 8px;
                        '>
                            <div style='
                                width: 4px; 
                                height: 4px; 
                                background: #6b7280;
                                border-radius: 50%; 
                                margin-right: 12px;
                                margin-top: 8px;
                                flex-shrink: 0;
                            '></div>
                            <div>{description}</div>
                        </div>
                    </div>
                </div>
                """
                certificate_links_html += card_html

    # Main HTML content - exactly as before
    html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{session_state['name']} - Professional Resume</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">
    <style>
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            line-height: 1.6;
            color: #1a202c;
            background: #ffffff;
            min-height: 100vh;
        }}
        
        .resume-container {{
            width: 100%;
            min-height: 100vh;
            background: #ffffff;
        }}
        
        .resume-container::before {{
            content: '';
            display: block;
            height: 4px;
            background: linear-gradient(90deg, #6b7280, #9ca3af);
        }}
        
        .header-section {{
            background: #f8fafc;
            padding: 40px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            border-bottom: 1px solid #e2e8f0;
        }}
        
        .name-title {{
            flex: 1;
        }}
        
        .name-title h1 {{
            font-size: 42px;
            font-weight: 800;
            color: #1a202c;
            margin-bottom: 8px;
        }}
        
        .name-title h2 {{
            font-size: 24px;
            font-weight: 600;
            color: #4a5568;
            margin: 0;
        }}
        
        .profile-image {{
            flex-shrink: 0;
            margin-left: 40px;
        }}
        
        .main-content {{
            display: flex;
            min-height: 800px;
        }}
        
        .sidebar {{
            width: 350px;
            background: #f7fafc;
            padding: 40px 30px;
            border-right: 1px solid #e2e8f0;
        }}
        
        .main-section {{
            flex: 1;
            padding: 40px;
            background: #ffffff;
        }}
        
        .contact-info {{
            margin-bottom: 40px;
        }}
        
        .contact-item {{
            display: flex;
            align-items: center;
            margin-bottom: 12px;
            padding: 8px 0;
        }}
        
        .contact-icon {{
            width: 20px;
            height: 20px;
            margin-right: 15px;
            opacity: 0.8;
        }}
        
        .contact-item span, .contact-item a {{
            font-size: 14px;
            color: #4a5568;
            text-decoration: none;
            font-weight: 500;
        }}
        
        .contact-item a:hover {{
            color: #6b7280;
            transition: color 0.3s ease;
        }}
        
        .section-title {{
            font-size: 22px;
            font-weight: 700;
            color: #2d3748;
            margin: 35px 0 15px 0;
        }}
        
        .section-content {{
            margin-bottom: 30px;
        }}
        
        .summary-text {{
            font-size: 16px;
            line-height: 1.8;
            color: #4a5568;
            background: #f8fafc;
            padding: 25px;
            border-radius: 8px;
            border-left: 3px solid #9ca3af;
        }}
        
        @media (max-width: 768px) {{
            .main-content {{
                flex-direction: column;
            }}
            
            .sidebar {{
                width: 100%;
            }}
            
            .header-section {{
                flex-direction: column;
                text-align: center;
            }}
            
            .profile-image {{
                margin: 20px 0 0 0;
            }}
            
            .name-title h1 {{
                font-size: 32px;
            }}
        }}
        
        @media (max-width: 480px) {{
            .header-section, .sidebar, .main-section {{
                padding: 20px;
            }}
        }}
    </style>
</head>
<body>
    <div class="resume-container">
        <div class="header-section">
            <div class="name-title">
                <h1>{session_state['name']}</h1>
                <h2>{session_state['job_title']}</h2>
            </div>
            <div class="profile-image">
                {profile_img_html}
            </div>
        </div>

        <div class="main-content">
            <div class="sidebar">
                <div class="contact-info">
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path fill-rule="evenodd" d="M5.05 4.05a7 7 0 119.9 9.9L10 18.9l-4.95-4.95a7 7 0 010-9.9zM10 11a2 2 0 100-4 2 2 0 000 4z" clip-rule="evenodd"></path>
                        </svg>
                        <span>{session_state['location']}</span>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path d="M2 3a1 1 0 011-1h2.153a1 1 0 01.986.836l.74 4.435a1 1 0 01-.54 1.06l-1.548.773a11.037 11.037 0 006.105 6.105l.774-1.548a1 1 0 011.059-.54l4.435.74a1 1 0 01.836.986V17a1 1 0 01-1 1h-2C7.82 18 2 12.18 2 5V3z"></path>
                        </svg>
                        <span>{session_state['phone']}</span>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path d="M2.003 5.884L10 9.882l7.997-3.998A2 2 0 0016 4H4a2 2 0 00-1.997 1.884z"></path>
                            <path d="M18 8.118l-8 4-8-4V14a2 2 0 002 2h12a2 2 0 002-2V8.118z"></path>
                        </svg>
                        <a href="mailto:{session_state['email']}">{session_state['email']}</a>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 24 24">
                            <path d="M20.447 20.452h-3.554v-5.569c0-1.328-.027-3.037-1.852-3.037-1.853 0-2.136 1.445-2.136 2.939v5.667H9.351V9h3.414v1.561h.046c.477-.9 1.637-1.85 3.37-1.85 3.601 0 4.267 2.37 4.267 5.455v6.286zM5.337 7.433a2.062 2.062 0 01-2.063-2.065 2.064 2.064 0 112.063 2.065zm1.782 13.019H3.555V9h3.564v11.452zM22.225 0H1.771C.792 0 0 .774 0 1.729v20.542C0 23.227.792 24 1.771 24h20.451C23.2 24 24 23.227 24 22.271V1.729C24 .774 23.2 0 22.222 0h.003z"/>
                        </svg>
                        <a href="{session_state['linkedin']}" target="_blank">LinkedIn</a>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path fill-rule="evenodd" d="M4.083 9h1.946c.089-1.546.383-2.97.837-4.118A6.004 6.004 0 004.083 9zM10 2a8 8 0 100 16 8 8 0 000-16zm0 2c-.076 0-.232.032-.465.262-.238.234-.497.623-.737 1.182-.389.907-.673 2.142-.766 3.556h3.936c-.093-1.414-.377-2.649-.766-3.556-.24-.56-.5-.948-.737-1.182C10.232 4.032 10.076 4 10 4zm3.971 5c-.089-1.546-.383-2.97-.837-4.118A6.004 6.004 0 0115.917 9h-1.946zm-2.003 2H8.032c.093 1.414.377 2.649.766 3.556.24.56.5.948.737 1.182.233.23.389.262.465.262.076 0 .232-.032.465-.262.238-.234.498-.623.737-1.182.389-.907.673-2.142.766-3.556zm1.166 4.118c.454-1.147.748-2.572.837-4.118h1.946a6.004 6.004 0 01-2.783 4.118zm-6.268 0C6.412 13.97 6.118 12.546 6.03 11H4.083a6.004 6.004 0 002.783 4.118z" clip-rule="evenodd"></path>
                        </svg>
                        <a href="{session_state['portfolio']}" target="_blank">Portfolio</a>
                    </div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Skills</h3>
                    <div>{skills_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Languages</h3>
                    <div>{languages_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Interests</h3>
                    <div>{interests_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Soft Skills</h3>
                    <div>{Softskills_html}</div>
                </div>
            </div>

            <div class="main-section">
                <div class="section-content">
                    <h3 class="section-title">Professional Summary</h3>
                    <div class="summary-text">{summary_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Work Experience</h3>
                    {experience_html}
                </div>

                <div class="section-content">
                    <h3 class="section-title">Education</h3>
                    {education_html}
                </div>

                <div class="section-content">
                    <h3 class="section-title">Projects</h3>
                    {projects_html}
                </div>

                <div class="section-content">
                    {project_links_html}
                </div>

                <div class="section-content">
                    {certificate_links_html}
                </div>
            </div>
        </div>
    </div>
</body>
</html>
"""
    
    return html_content

def render_template_modern(session_state, profile_img_html=""):
    """Modern minimal template with clean design, pill-style tags for enhanced visual appeal"""
    
    # Process lists into pill tags instead of plain lists
    skills_list = [s.strip() for s in session_state['skills'].split(',') if s.strip()]
    languages_list = [l.strip() for l in session_state['languages'].split(',') if l.strip()]
    interests_list = [i.strip() for i in session_state['interests'].split(',') if i.strip()]
    softskills_list = [s.strip() for s in session_state['Softskills'].split(',') if s.strip()]
    
    # Create unified pill-style tags for all sections
    skills_pills = "".join([
        f"""<span style="
            display: inline-block;
            background: linear-gradient(135deg, #e0f2fe 0%, #bae6fd 100%);
            color: #0c4a6e;
            padding: 8px 16px;
            margin: 4px 6px 4px 0;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
            box-shadow: 0 2px 4px rgba(8, 145, 178, 0.1);
            border: 1px solid rgba(14, 165, 233, 0.2);
        ">{skill}</span>""" for skill in skills_list
    ])
    
    # Create unified pill-style tags for languages
    languages_pills = "".join([
        f"""<span style="
            display: inline-block;
            background: linear-gradient(135deg, #e0f2fe 0%, #bae6fd 100%);
            color: #0c4a6e;
            padding: 8px 16px;
            margin: 4px 6px 4px 0;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
            box-shadow: 0 2px 4px rgba(8, 145, 178, 0.1);
            border: 1px solid rgba(14, 165, 233, 0.2);
        ">{lang}</span>""" for lang in languages_list
    ])
    
    # Create unified pill-style tags for interests
    interests_pills = "".join([
        f"""<span style="
            display: inline-block;
            background: linear-gradient(135deg, #e0f2fe 0%, #bae6fd 100%);
            color: #0c4a6e;
            padding: 8px 16px;
            margin: 4px 6px 4px 0;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
            box-shadow: 0 2px 4px rgba(8, 145, 178, 0.1);
            border: 1px solid rgba(14, 165, 233, 0.2);
        ">{interest}</span>""" for interest in interests_list
    ])
    
    # Create unified pill-style tags for soft skills
    softskills_pills = "".join([
        f"""<span style="
            display: inline-block;
            background: linear-gradient(135deg, #e0f2fe 0%, #bae6fd 100%);
            color: #0c4a6e;
            padding: 8px 16px;
            margin: 4px 6px 4px 0;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
            box-shadow: 0 2px 4px rgba(8, 145, 178, 0.1);
            border: 1px solid rgba(14, 165, 233, 0.2);
        ">{skill}</span>""" for skill in softskills_list
    ])
    
    html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{session_state['name']} - Modern Resume</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Inter', 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #374151;
            background: #ffffff;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 50px;
            padding: 40px 0;
            position: relative;
        }}
        
        .header::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 50%;
            transform: translateX(-50%);
            width: 80px;
            height: 3px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 2px;
        }}
        
        .profile-image-container {{
            margin-bottom: 25px;
        }}
        
        .profile-image-container img {{
            width: 160px;
            height: 160px;
            border-radius: 50%;
            object-fit: cover;
            object-position: center;
            border: 4px solid #3b82f6;
            box-shadow: 0 8px 32px rgba(59, 130, 246, 0.3), 0 0 0 8px rgba(59, 130, 246, 0.1);
            display: block;
            margin: 0 auto;
        }}
        
        .header h1 {{
            font-size: 2.75rem;
            font-weight: 700;
            color: #1f2937;
            margin-bottom: 12px;
            letter-spacing: -0.025em;
        }}
        
        .header h2 {{
            font-size: 1.35rem;
            font-weight: 500;
            color: #6b7280;
            margin-bottom: 25px;
        }}
        
        .contact-info {{
            display: flex;
            justify-content: center;
            flex-wrap: wrap;
            gap: 25px;
            font-size: 0.95rem;
            color: #4b5563;
        }}
        
        .contact-info a {{
            color: #3b82f6;
            text-decoration: none;
            font-weight: 500;
            transition: color 0.2s ease;
        }}
        
        .contact-info a:hover {{
            color: #1d4ed8;
        }}
        
        .section {{
            margin-bottom: 40px;
        }}
        
        .section h3 {{
            font-size: 1.5rem;
            font-weight: 700;
            color: #1f2937;
            margin-bottom: 20px;
            position: relative;
            padding-bottom: 10px;
            text-align: center;
        }}
        
        .section h3::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 50%;
            transform: translateX(-50%);
            width: 50px;
            height: 2px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 1px;
        }}
        
        .project-links {{
            text-align: center;
        }}
        
        .summary {{
            font-size: 1.1rem;
            line-height: 1.8;
            color: #4b5563;
            background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
            padding: 30px;
            border-radius: 12px;
            border: 1px solid #e2e8f0;
            position: relative;
        }}
        
        .summary::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 4px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 12px 12px 0 0;
        }}
        
        .experience-item, .education-item, .project-item {{
            margin-bottom: 30px;
            padding: 25px;
            background: linear-gradient(135deg, #fafbfc 0%, #f4f6f8 100%);
            border-radius: 12px;
            border: 1px solid #e5e7eb;
            position: relative;
        }}
        
        .experience-item::before, .education-item::before, .project-item::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 3px;
            background: linear-gradient(90deg, #6b7280, #9ca3af);
            border-radius: 12px 12px 0 0;
        }}
        
        .item-header {{
            display: flex;
            justify-content: space-between;
            align-items: baseline;
            margin-bottom: 8px;
            flex-wrap: wrap;
            gap: 10px;
        }}
        
        .item-title {{
            font-weight: 700;
            color: #1f2937;
            font-size: 1.2rem;
        }}
        
        .item-duration {{
            color: #6b7280;
            font-size: 0.95rem;
            font-weight: 600;
            background: linear-gradient(135deg, #f9fafb, #f3f4f6);
            padding: 6px 14px;
            border-radius: 16px;
            border: 1px solid #d1d5db;
        }}
        
        .item-subtitle {{
            color: #3b82f6;
            font-size: 1.05rem;
            margin-bottom: 12px;
            font-weight: 600;
        }}
        
        .item-description {{
            color: #4b5563;
            line-height: 1.7;
            font-size: 1rem;
        }}
        
        .pills-container {{
            display: flex;
            flex-wrap: wrap;
            justify-content: center;
            gap: 8px;
            margin-top: 10px;
        }}
        
        .links a {{
            display: inline-block;
            color: #3b82f6;
            text-decoration: none;
            margin-right: 20px;
            margin-bottom: 8px;
            font-weight: 500;
            padding: 8px 16px;
            background: linear-gradient(135deg, #eff6ff, #dbeafe);
            border-radius: 8px;
            border: 1px solid #bfdbfe;
            transition: all 0.2s ease;
        }}
        
        .links a:hover {{
            background: linear-gradient(135deg, #dbeafe, #bfdbfe);
            transform: translateY(-1px);
        }}
        
        @media (max-width: 768px) {{
            body {{
                padding: 20px 15px;
            }}
            
            .contact-info {{
                flex-direction: column;
                align-items: center;
                gap: 8px;
            }}
            
            .item-header {{
                flex-direction: column;
                align-items: flex-start;
                gap: 8px;
            }}
            
            .header h1 {{
                font-size: 2.2rem;
            }}
            
            .experience-item, .education-item, .project-item {{
                padding: 20px;
            }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="profile-image-container">
            {profile_img_html}
        </div>
        <h1>{session_state['name']}</h1>
        <h2>{session_state['job_title']}</h2>
        <div class="contact-info">
            <span>📍 {session_state['location']}</span>
            <span>📞 {session_state['phone']}</span>
            <a href="mailto:{session_state['email']}">✉️ {session_state['email']}</a>
            <a href="{session_state['linkedin']}" target="_blank">🔗 LinkedIn</a>
            <a href="{session_state['portfolio']}" target="_blank">🌐 Portfolio</a>
        </div>
    </div>

    <div class="section">
        <h3>Professional Summary</h3>
        <div class="summary">{session_state['summary'].replace(chr(10), '<br>')}</div>
    </div>

    <div class="section">
        <h3>Work Experience</h3>
        {"".join([f'''
        <div class="experience-item">
            <div class="item-header">
                <div class="item-title">{exp.get('title', '')}</div>
                <div class="item-duration">{exp.get('duration', '')}</div>
            </div>
            <div class="item-subtitle">{exp.get('company', '')}</div>
            <div class="item-description">{exp.get('description', '').replace(chr(10), '<br>')}</div>
        </div>
        ''' for exp in session_state.experience_entries if exp.get('company') or exp.get('title')])}
    </div>

    <div class="section">
        <h3>Education</h3>
        {"".join([f'''
        <div class="education-item">
            <div class="item-header">
                <div class="item-title">{edu.get('degree', '')}</div>
                <div class="item-duration">{edu.get('year', '')}</div>
            </div>
            <div class="item-subtitle">{edu.get('institution', '')}</div>
            <div class="item-description">{edu.get('details', '')}</div>
        </div>
        ''' for edu in session_state.education_entries if edu.get('institution') or edu.get('degree')])}
    </div>

    <div class="section">
        <h3>Projects</h3>
        {"".join([f'''
        <div class="project-item">
            <div class="item-header">
                <div class="item-title">{proj.get('title', '')}</div>
                <div class="item-duration">{proj.get('duration', '')}</div>
            </div>
            <div class="item-subtitle">Technologies: {proj.get('tech', '')}</div>
            <div class="item-description">{proj.get('description', '').replace(chr(10), '<br>')}</div>
        </div>
        ''' for proj in session_state.project_entries if proj.get('title')])}
    </div>

    <div class="section">
        <h3>Technical Skills</h3>
        <div class="pills-container">
            {skills_pills}
        </div>
    </div>

    <div class="section">
        <h3>Languages</h3>
        <div class="pills-container">
            {languages_pills}
        </div>
    </div>

    <div class="section">
        <h3>Professional Interests</h3>
        <div class="pills-container">
            {interests_pills}
        </div>
    </div>

    <div class="section">
        <h3>Core Competencies</h3>
        <div class="pills-container">
            {softskills_pills}
        </div>
    </div>

    {f'''
    <div class="section">
        <h3>Project Links</h3>
        <div class="links project-links">
            {"".join([f'<a href="{link}" target="_blank">🔗 Project {i+1}</a>' for i, link in enumerate(session_state.project_links)])}
        </div>
    </div>
    ''' if session_state.project_links else ''}

    {f'''
    <div class="section">
        <h3>Professional Certifications</h3>
        {"".join([f'''
        <div class="project-item">
            <div class="item-header">
                <div class="item-title"><a href="{cert['link']}" target="_blank" style="color: #1f2937; text-decoration: none;">{cert['name']}</a></div>
                <div class="item-duration">{cert.get('duration', '')}</div>
            </div>
            <div class="item-description">{cert.get('description', '')}</div>
        </div>
        ''' for cert in session_state.certificate_links if cert.get('name')])}
    </div>
    ''' if any(cert.get('name') for cert in session_state.certificate_links) else ''}

</body>
</html>
"""
    
    return html_content

def render_template_sidebar(session_state, profile_img_html=""):
    """Enhanced elegant sidebar template with improved styling, pill tags, and better visual hierarchy"""
    
    # Process lists for pill-style tags
    skills_list = [s.strip() for s in session_state['skills'].split(',') if s.strip()]
    languages_list = [l.strip() for l in session_state['languages'].split(',') if l.strip()]
    interests_list = [i.strip() for i in session_state['interests'].split(',') if i.strip()]
    softskills_list = [s.strip() for s in session_state['Softskills'].split(',') if s.strip()]
    
    # Create pill-style tags for sidebar sections
    skills_pills = "".join([
        f"""<div style="
            display: inline-block;
            background: rgba(56, 189, 248, 0.15);
            color: #e0f2fe;
            padding: 8px 16px;
            margin: 5px 8px 5px 0;
            border-radius: 18px;
            font-size: 0.85rem;
            font-weight: 600;
            border: 1px solid rgba(56, 189, 248, 0.3);
            box-shadow: 0 2px 4px rgba(56, 189, 248, 0.1);
        ">{skill}</div>""" for skill in skills_list
    ])
    
    languages_pills = "".join([
        f"""<div style="
            display: inline-block;
            background: rgba(34, 197, 94, 0.15);
            color: #dcfce7;
            padding: 8px 16px;
            margin: 5px 8px 5px 0;
            border-radius: 18px;
            font-size: 0.85rem;
            font-weight: 600;
            border: 1px solid rgba(34, 197, 94, 0.3);
            box-shadow: 0 2px 4px rgba(34, 197, 94, 0.1);
        ">{lang}</div>""" for lang in languages_list
    ])
    
    interests_pills = "".join([
        f"""<div style="
            display: inline-block;
            background: rgba(245, 158, 11, 0.15);
            color: #fef3c7;
            padding: 8px 16px;
            margin: 5px 8px 5px 0;
            border-radius: 18px;
            font-size: 0.85rem;
            font-weight: 600;
            border: 1px solid rgba(245, 158, 11, 0.3);
            box-shadow: 0 2px 4px rgba(245, 158, 11, 0.1);
        ">{interest}</div>""" for interest in interests_list
    ])
    
    softskills_pills = "".join([
        f"""<div style="
            display: inline-block;
            background: rgba(168, 85, 247, 0.15);
            color: #f3e8ff;
            padding: 8px 16px;
            margin: 5px 8px 5px 0;
            border-radius: 18px;
            font-size: 0.85rem;
            font-weight: 600;
            border: 1px solid rgba(168, 85, 247, 0.3);
            box-shadow: 0 2px 4px rgba(168, 85, 247, 0.1);
        ">{skill}</div>""" for skill in softskills_list
    ])
    
    # Enhanced profile image styling
    enhanced_profile_img = ""
    if profile_img_html:
        # Extract the img tag and enhance it
        import re
        img_match = re.search(r'<img[^>]*>', profile_img_html)
        if img_match:
            enhanced_profile_img = img_match.group(0).replace(
                'style="',
                'style="width: 160px; height: 160px; border-radius: 50%; object-fit: cover; object-position: center; border: 4px solid #38bdf8; box-shadow: 0 8px 32px rgba(56, 189, 248, 0.3), 0 0 0 8px rgba(56, 189, 248, 0.1); margin-bottom: 20px; '
            )
    
    html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{session_state['name']} - Elegant Resume</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Inter', 'Segoe UI', sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f8fafc;
        }}
        
        .resume-container {{
            width: 100%;
            display: flex;
            min-height: 100vh;
            background: white;
            box-shadow: 0 0 30px rgba(0,0,0,0.1);
        }}
        
        .sidebar {{
            width: 350px;
            background: linear-gradient(180deg, #1e293b 0%, #334155 100%);
            color: white;
            padding: 40px 30px;
            position: relative;
        }}
        
        .sidebar::before {{
            content: '';
            position: absolute;
            top: 0;
            right: 0;
            width: 4px;
            height: 100%;
            background: linear-gradient(180deg, #38bdf8, #06b6d4);
        }}
        
        .main-content {{
            flex: 1;
            padding: 40px 50px;
            background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
        }}
        
        .profile-section {{
            text-align: center;
            margin-bottom: 45px;
            position: relative;
        }}
        
        .profile-section::after {{
            content: '';
            position: absolute;
            bottom: -20px;
            left: 50%;
            transform: translateX(-50%);
            width: 60px;
            height: 3px;
            background: linear-gradient(90deg, #38bdf8, #06b6d4);
            border-radius: 2px;
        }}
        
        .profile-section h1 {{
            font-size: 1.95rem;
            margin-bottom: 12px;
            color: #f8fafc;
            font-weight: 700;
            letter-spacing: -0.025em;
        }}
        
        .profile-section h2 {{
            font-size: 1.1rem;
            color: #cbd5e1;
            margin-bottom: 25px;
            font-weight: 500;
        }}
        
        .contact-section {{
            margin-bottom: 40px;
        }}
        
        .contact-item {{
            display: flex;
            align-items: center;
            margin-bottom: 18px;
            padding: 12px;
            background: rgba(56, 189, 248, 0.1);
            border-radius: 10px;
            border: 1px solid rgba(56, 189, 248, 0.2);
            transition: all 0.3s ease;
        }}
        
        .contact-item:hover {{
            background: rgba(56, 189, 248, 0.15);
            transform: translateX(5px);
        }}
        
        .contact-icon {{
            margin-right: 15px;
            font-size: 1.1rem;
            color: #38bdf8;
            width: 20px;
            text-align: center;
        }}
        
        .contact-item span, .contact-item a {{
            color: #e2e8f0;
            text-decoration: none;
            font-weight: 500;
            font-size: 0.9rem;
            word-break: break-word;
            overflow-wrap: anywhere;
            max-width: 100%;
            display: inline-block;
        }}
        
        .contact-item a:hover {{
            color: #38bdf8;
            transition: color 0.3s ease;
        }}
        
        .sidebar-section {{
            margin-bottom: 40px;
        }}
        
        .sidebar-section h3 {{
            font-size: 1.2rem;
            margin-bottom: 20px;
            color: #38bdf8;
            text-transform: uppercase;
            letter-spacing: 1px;
            font-weight: 700;
            position: relative;
            padding-bottom: 10px;
        }}
        
        .sidebar-section h3::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 0;
            width: 40px;
            height: 2px;
            background: linear-gradient(90deg, #38bdf8, #06b6d4);
            border-radius: 1px;
        }}
        
        .main-section {{
            margin-bottom: 40px;
        }}
        
        .main-section h3 {{
            font-size: 1.65rem;
            color: #1e293b;
            margin-bottom: 25px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            font-weight: 700;
            position: relative;
            padding-bottom: 15px;
        }}
        
        .main-section h3::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 0;
            width: 60px;
            height: 3px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 2px;
        }}
        
        .summary {{
            font-size: 1.1rem;
            line-height: 1.8;
            color: #4b5563;
            background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
            padding: 30px;
            border-radius: 15px;
            border: 1px solid #bae6fd;
            position: relative;
            box-shadow: 0 4px 6px rgba(59, 130, 246, 0.05);
        }}
        
        .summary::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 4px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 15px 15px 0 0;
        }}
        
        .content-item {{
            margin-bottom: 30px;
            padding: 30px;
            background: linear-gradient(135deg, #ffffff 0%, #f9fafb 100%);
            border-radius: 15px;
            border: 1px solid #e5e7eb;
            position: relative;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        }}
        
        .content-item::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 4px;
            background: linear-gradient(90deg, #6b7280, #9ca3af);
            border-radius: 15px 15px 0 0;
        }}
        
        .content-item:last-child {{
            margin-bottom: 0;
        }}
        
        .item-header {{
            display: flex;
            justify-content: space-between;
            align-items: baseline;
            margin-bottom: 12px;
            flex-wrap: wrap;
            gap: 15px;
        }}
        
        .item-title {{
            font-size: 1.25rem;
            font-weight: 700;
            color: #1e293b;
        }}
        
        .item-duration {{
            color: #6b7280;
            font-size: 0.95rem;
            font-weight: 600;
            background: linear-gradient(135deg, #f1f5f9, #e2e8f0);
            padding: 8px 16px;
            border-radius: 20px;
            border: 1px solid #cbd5e1;
        }}
        
        .item-company {{
            color: #3b82f6;
            font-size: 1.1rem;
            margin-bottom: 15px;
            font-weight: 700;
        }}
        
        .item-description {{
            color: #4b5563;
            line-height: 1.7;
            font-size: 1rem;
        }}
        
        .project-tech {{
            background: linear-gradient(135deg, #dbeafe, #bfdbfe);
            color: #1e40af;
            padding: 10px 18px;
            border-radius: 10px;
            font-size: 0.9rem;
            margin-bottom: 15px;
            display: inline-block;
            font-weight: 600;
            border: 1px solid #93c5fd;
        }}
        
        @media (max-width: 768px) {{
            .resume-container {{
                flex-direction: column;
            }}
            
            .sidebar {{
                width: 100%;
            }}
            
            .item-header {{
                flex-direction: column;
                align-items: flex-start;
                gap: 8px;
            }}
            
            .main-content {{
                padding: 30px 25px;
            }}
            
            .sidebar {{
                padding: 30px 25px;
            }}
        }}
    </style>
</head>
<body>
    <div class="resume-container">
        <div class="sidebar">
            <div class="profile-section">
                {enhanced_profile_img}
                <h1>{session_state['name']}</h1>
                <h2>{session_state['job_title']}</h2>
            </div>
            
            <div class="contact-section">
                <div class="contact-item">
                    <div class="contact-icon">📍</div>
                    <span>{session_state['location']}</span>
                </div>
                <div class="contact-item">
                    <div class="contact-icon">📞</div>
                    <span>{session_state['phone']}</span>
                </div>
                <div class="contact-item">
                    <div class="contact-icon">✉️</div>
                    <a href="mailto:{session_state['email']}">{session_state['email']}</a>
                </div>
                <div class="contact-item">
                    <div class="contact-icon">🔗</div>
                    <a href="{session_state['linkedin']}" target="_blank">LinkedIn Profile</a>
                </div>
                <div class="contact-item">
                    <div class="contact-icon">🌐</div>
                    <a href="{session_state['portfolio']}" target="_blank">Portfolio Website</a>
                </div>
            </div>
            
            <div class="sidebar-section">
                <h3>Technical Skills</h3>
                <div>{skills_pills}</div>
            </div>
            
            <div class="sidebar-section">
                <h3>Languages</h3>
                <div>{languages_pills}</div>
            </div>
            
            <div class="sidebar-section">
                <h3>Interests</h3>
                <div>{interests_pills}</div>
            </div>
            
            <div class="sidebar-section">
                <h3>Core Competencies</h3>
                <div>{softskills_pills}</div>
            </div>
        </div>
        
        <div class="main-content">
            <div class="main-section">
                <h3>Professional Summary</h3>
                <div class="summary">{session_state['summary'].replace(chr(10), '<br>')}</div>
            </div>
            
            <div class="main-section">
                <h3>Professional Experience</h3>
                {"".join([f'''
                <div class="content-item">
                    <div class="item-header">
                        <div class="item-title">{exp.get('title', '')}</div>
                        <div class="item-duration">{exp.get('duration', '')}</div>
                    </div>
                    <div class="item-company">{exp.get('company', '')}</div>
                    <div class="item-description">{exp.get('description', '').replace(chr(10), '<br>')}</div>
                </div>
                ''' for exp in session_state.experience_entries if exp.get('company') or exp.get('title')])}
            </div>
            
            <div class="main-section">
                <h3>Education & Qualifications</h3>
                {"".join([f'''
                <div class="content-item">
                    <div class="item-header">
                        <div class="item-title">{edu.get('degree', '')}</div>
                        <div class="item-duration">{edu.get('year', '')}</div>
                    </div>
                    <div class="item-company">{edu.get('institution', '')}</div>
                    <div class="item-description">{edu.get('details', '')}</div>
                </div>
                ''' for edu in session_state.education_entries if edu.get('institution') or edu.get('degree')])}
            </div>
            
            <div class="main-section">
                <h3>Key Projects</h3>
                {"".join([f'''
                <div class="content-item">
                    <div class="item-header">
                        <div class="item-title">{proj.get('title', '')}</div>
                        <div class="item-duration">{proj.get('duration', '')}</div>
                    </div>
                    <div class="project-tech">Technologies: {proj.get('tech', '')}</div>
                    <div class="item-description">{proj.get('description', '').replace(chr(10), '<br>')}</div>
                </div>
                ''' for proj in session_state.project_entries if proj.get('title')])}
            </div>
            
            {f'''
            <div class="main-section">
                <h3>Project Portfolio</h3>
                {"".join([f'''<div class="content-item" style="padding: 20px;"><a href="{link}" target="_blank" style="color: #3b82f6; text-decoration: none; font-weight: 600; font-size: 1.1rem;">🔗 Project Repository {i+1}</a></div>''' for i, link in enumerate(session_state.project_links)])}
            </div>
            ''' if session_state.project_links else ''}
            
            {f'''
            <div class="main-section">
                <h3>Professional Certifications</h3>
                {"".join([f'''
                <div class="content-item">
                    <div class="item-header">
                        <div class="item-title"><a href="{cert['link']}" target="_blank" style="color: #1e293b; text-decoration: none;">{cert['name']}</a></div>
                        <div class="item-duration">{cert.get('duration', '')}</div>
                    </div>
                    <div class="item-description">{cert.get('description', '')}</div>
                </div>
                ''' for cert in session_state.certificate_links if cert.get('name')])}
            </div>
            ''' if any(cert.get('name') for cert in session_state.certificate_links) else ''}
        </div>
    </div>
</body>
</html>
"""
    
    return html_content

def generate_cover_letter_from_resume_builder():
    import streamlit as st
    from datetime import datetime
    import re
    from llm_manager import call_llm  # Ensure you import this

    name = st.session_state.get("name", "")
    job_title = st.session_state.get("job_title", "")
    summary = st.session_state.get("summary", "")
    skills = st.session_state.get("skills", "")
    location = st.session_state.get("location", "")
    today_date = datetime.today().strftime("%B %d, %Y")

    # ✅ Input boxes for contact info
    company = st.text_input("🏢 Target Company", placeholder="e.g., Google")
    linkedin = st.text_input("🔗 LinkedIn URL", placeholder="e.g., https://linkedin.com/in/username")
    email = st.text_input("📧 Email", placeholder="e.g., you@example.com")
    mobile = st.text_input("📞 Mobile Number", placeholder="e.g., +91 9876543210")

    # ✅ Button to prevent relooping
    if st.button("✉️ Generate Cover Letter"):
        # ✅ Validate input before generating
        if not all([name, job_title, summary, skills, company, linkedin, email, mobile]):
            st.warning("⚠️ Please fill in all fields including LinkedIn, email, and mobile.")
            return

        prompt = f"""
You are a professional cover letter writer.

Write a formal and compelling cover letter using the information below. 
Format it as a real letter with:
1. Date
2. Recipient heading
3. Proper salutation
4. Three short paragraphs
5. Professional closing

Ensure you **only include the company name once** in the header or salutation, 
and avoid repeating it redundantly in the body.

### Heading Info:
{today_date}
Hiring Manager, {company}, {location}

### Candidate Info:
- Name: {name}
- Job Title: {job_title}
- Summary: {summary}
- Skills: {skills}
- Location: {location}

### Instructions:
- Do not use HTML tags. 
- Return plain text only.
"""

        # ✅ Call LLM
        cover_letter = call_llm(prompt, session=st.session_state).strip()

        # ✅ Store plain text
        st.session_state["cover_letter"] = cover_letter

        # ✅ Build HTML wrapper for preview (safe)
        cover_letter_html = f"""
        <div style="font-family: Georgia, serif; font-size: 13pt; line-height: 1.6; 
                    color: #000; background: #fff; padding: 25px; 
                    border-radius: 8px; box-shadow: 0px 2px 6px rgba(0,0,0,0.1); 
                    max-width: 800px; margin: auto;">
            <div style="text-align:center; margin-bottom:15px;">
                <div style="font-size:18pt; font-weight:bold; color:#003366;">{name}</div>
                <div style="font-size:14pt; color:#555;">{job_title}</div>
                <div style="font-size:10pt; margin-top:5px;">
                    <a href="{linkedin}" style="color:#003366;">{linkedin}</a><br/>
                    📧 {email} | 📞 {mobile}
                </div>
            </div>
            <hr/>
            <pre style="white-space: pre-wrap; font-family: Georgia, serif; font-size: 12pt; color:#000;">
{cover_letter}
            </pre>
        </div>
        """

        st.session_state["cover_letter_html"] = cover_letter_html

        # ✅ Show nicely in Streamlit
        st.markdown(cover_letter_html, unsafe_allow_html=True)

# Import necessary modules first
import streamlit as st

# Tab setup (assuming this is within a tab2 context)
with tab2:
    st.session_state.active_tab = "Resume Builder"

    # ---------- Title with Blue Glassmorphism + Shine ----------
    st.markdown("""
    <style>
    .glass-title {
        background: rgba(10, 20, 40, 0.5);
        border-radius: 20px;
        padding: 20px;
        backdrop-filter: blur(14px);
        box-shadow: 0 8px 32px rgba(0, 200, 255, 0.25);
        border: 1px solid rgba(0, 200, 255, 0.3);
        text-align: center;
        position: relative;
        overflow: hidden;
    }
    .glass-title h2 {
        color: #4da6ff;
        margin: 0;
        text-shadow: 0 0 12px rgba(0,200,255,0.7);
        font-weight: 600;
    }
    .glass-title::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.18) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .glass-title:hover::before {
        left: 100%;
        top: 100%;
    }
    </style>

    <div class="glass-title">
        <h2>🧾 Advanced Resume Builder</h2>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<hr style='border-top: 2px solid rgba(0,200,255,0.4);'>", unsafe_allow_html=True)

    # ---------- Global Styles (Glassmorphism + Glow + Shine) ----------
    st.markdown("""
        <style>
        /* File uploader */
        .uploadedFile { 
            background: rgba(10, 20, 40, 0.6) !important;
            border: 1px solid rgba(0,200,255,0.5) !important;
            border-radius: 14px !important;
            color: #cce6ff !important;
            box-shadow: 0 0 12px rgba(0,200,255,0.3) !important;
        }

        /* Sidebar expander style */
        .streamlit-expanderHeader {
            background: rgba(10, 20, 40, 0.45);
            border-radius: 12px;
            color: #4da6ff !important;
            font-weight: bold;
            backdrop-filter: blur(12px);
            box-shadow: 0 4px 12px rgba(0,200,255,0.25);
            transition: all 0.3s ease-in-out;
        }
        .streamlit-expanderHeader:hover {
            background: rgba(0, 200, 255, 0.12);
            box-shadow: 0 0 16px rgba(0,200,255,0.4);
        }
        .streamlit-expanderContent {
            background: rgba(10, 20, 40, 0.45);
            border-radius: 10px;
            padding: 8px;
            color: #e6f7ff;
        }

        /* Selectbox */
        div[data-baseweb="select"] {
            background: rgba(10, 20, 40, 0.35);
            border: 1px solid rgba(0, 200, 255, 0.6);
            border-radius: 12px;
            color: #e6f7ff;
            backdrop-filter: blur(14px);
            box-shadow: 0 0 10px rgba(0,200,255,0.3);
        }

        /* Buttons with Shine Effect */
        div.stButton > button {
            position: relative;
            background: rgba(10, 20, 40, 0.35);
            border: 1px solid rgba(0, 200, 255, 0.6);
            color: #e6f7ff;
            border-radius: 14px;
            padding: 10px 20px;
            font-size: 15px;
            font-weight: 500;
            backdrop-filter: blur(16px);
            box-shadow: 0 0 12px rgba(0, 200, 255, 0.35),
                        inset 0 0 20px rgba(0, 200, 255, 0.05);
            overflow: hidden;
            transition: all 0.3s ease-in-out;
        }
        div.stButton > button::before {
            content: "";
            position: absolute;
            top: -50%;
            left: -50%;
            width: 200%;
            height: 200%;
            background: linear-gradient(
                120deg,
                rgba(255,255,255,0.15) 0%,
                rgba(255,255,255,0.05) 40%,
                transparent 60%
            );
            transform: rotate(25deg);
            transition: all 0.6s;
        }
        div.stButton > button:hover::before {
            left: 100%;
            top: 100%;
        }
        div.stButton > button:hover {
            background: rgba(0, 200, 255, 0.12);
            box-shadow: 0 0 20px rgba(0, 200, 255, 0.65),
                        inset 0 0 25px rgba(0, 200, 255, 0.15);
            transform: translateY(-2px);
        }
        div.stButton > button:active {
            transform: scale(0.95);
            box-shadow: 0 0 10px rgba(0, 200, 255, 0.45);
        }
        </style>
    """, unsafe_allow_html=True)

    # 🎨 Template Selection
    st.markdown("### 🎨 Choose Resume Template")
    selected_template = st.selectbox(
        "🎨 Choose Resume Template",
        ["Default (Professional)", "Modern Minimal", "Elegant Sidebar"],
        key="template_selector"
    )

    # 📸 Upload profile photo
    uploaded_image = st.file_uploader("Upload a Profile Image", type=["png", "jpg", "jpeg"], key="profile_img_upload")
    profile_img_html = ""

    if uploaded_image:
        import base64
        encoded_image = base64.b64encode(uploaded_image.read()).decode()
        st.session_state["encoded_profile_image"] = encoded_image

        profile_img_html = f"""
        <div style="display: flex; justify-content: flex-end; margin-top: 20px;">
            <img src="data:image/png;base64,{encoded_image}" alt="Profile Photo"
                 style="
                    width: 140px;
                    height: 140px;
                    border-radius: 50%;
                    object-fit: cover;
                    object-position: center;
                    border: 4px solid rgba(255,255,255,0.6);
                    box-shadow:
                        0 0 0 3px #4da6ff,
                        0 8px 25px rgba(77, 166, 255, 0.3),
                        0 4px 15px rgba(0, 0, 0, 0.15);
                    transition: transform 0.3s ease-in-out;
                "
                onmouseover="this.style.transform='scale(1.07)'"
                onmouseout="this.style.transform='scale(1)'"
             />
        </div>
        """
        st.markdown(profile_img_html, unsafe_allow_html=True)
    else:
        st.info("📸 Please upload a clear, front-facing profile photo (square or vertical preferred).")

    # ---------------- Session State Defaults ----------------
    fields = ["name", "email", "phone", "linkedin", "location", "portfolio", "summary",
              "skills", "languages", "interests", "Softskills", "job_title"]
    for f in fields:
        st.session_state.setdefault(f, "")

    st.session_state.setdefault("experience_entries", [{"title": "", "company": "", "duration": "", "description": ""}])
    st.session_state.setdefault("education_entries", [{"degree": "", "institution": "", "year": "", "details": ""}])
    st.session_state.setdefault("project_entries", [{"title": "", "tech": "", "duration": "", "description": ""}])
    st.session_state.setdefault("project_links", [])
    st.session_state.setdefault("certificate_links", [{"name": "", "link": "", "duration": "", "description": ""}])

    # ---------------- Sidebar (ONLY in Tab 2) ----------------
    with st.sidebar:
        st.markdown("### ✨ Manage Resume Sections")

        if "edit_mode" not in st.session_state:
            st.session_state.edit_mode = "Add"

        mode = st.selectbox("Mode", ["Add", "Delete"], index=0, key="mode_dropdown")
        st.session_state.edit_mode = mode
        st.markdown("---")

        # 💼 Experience
        with st.expander("💼 Experience"):
            if st.button(f"{'➕ Add' if mode=='Add' else '❌ Delete'} Experience", key="exp_btn"):
                if mode == "Add":
                    st.session_state.experience_entries.append(
                        {"title": "", "company": "", "duration": "", "description": ""}
                    )
                elif mode == "Delete" and len(st.session_state.experience_entries) > 1:
                    st.session_state.experience_entries.pop()

        # 🎓 Education
        with st.expander("🎓 Education"):
            if st.button(f"{'➕ Add' if mode=='Add' else '❌ Delete'} Education", key="edu_btn"):
                if mode == "Add":
                    st.session_state.education_entries.append(
                        {"degree": "", "institution": "", "year": "", "details": ""}
                    )
                elif mode == "Delete" and len(st.session_state.education_entries) > 1:
                    st.session_state.education_entries.pop()

        # 🛠 Projects
        with st.expander("🛠 Projects"):
            if st.button(f"{'➕ Add' if mode=='Add' else '❌ Delete'} Project", key="proj_btn"):
                if mode == "Add":
                    st.session_state.project_entries.append(
                        {"title": "", "tech": "", "duration": "", "description": ""}
                    )
                elif mode == "Delete" and len(st.session_state.project_entries) > 1:
                    st.session_state.project_entries.pop()

        # 📜 Certificates
        with st.expander("📜 Certificates"):
            if st.button(f"{'➕ Add' if mode=='Add' else '❌ Delete'} Certificate", key="cert_btn"):
                if mode == "Add":
                    st.session_state.certificate_links.append(
                        {"name": "", "link": "", "duration": "", "description": ""}
                    )
                elif mode == "Delete" and len(st.session_state.certificate_links) > 1:
                    st.session_state.certificate_links.pop()

    # ---------------- Resume Form ----------------
    with st.form("resume_form", clear_on_submit=False):
        st.markdown("### 👤 <u>Personal Information</u>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.name = st.text_input("👤 Full Name", value=st.session_state.name, key="name_input")
            st.session_state.phone = st.text_input("📞 Phone Number", value=st.session_state.phone, key="phone_input")
            st.session_state.location = st.text_input("📍 Location", value=st.session_state.location, key="loc_input")
        with col2:
            st.session_state.email = st.text_input("📧 Email", value=st.session_state.email, key="email_input")
            st.session_state.linkedin = st.text_input("🔗 LinkedIn", value=st.session_state.linkedin, key="ln_input")
            st.session_state.portfolio = st.text_input("🌐 Portfolio", value=st.session_state.portfolio, key="port_input")
            st.session_state.job_title = st.text_input("💼 Job Title", value=st.session_state.job_title, key="job_input")

        st.markdown("### 📝 <u>Professional Summary</u>", unsafe_allow_html=True)
        st.session_state.summary = st.text_area("Summary", value=st.session_state.summary, key="summary_input")

        st.markdown("### 💼 <u>Skills, Languages, Interests & Soft Skills</u>", unsafe_allow_html=True)
        st.session_state.skills = st.text_area("Skills (comma-separated)", value=st.session_state.skills, key="skills_input")
        st.session_state.languages = st.text_area("Languages (comma-separated)", value=st.session_state.languages, key="lang_input")
        st.session_state.interests = st.text_area("Interests (comma-separated)", value=st.session_state.interests, key="int_input")
        st.session_state.Softskills = st.text_area("Softskills (comma-separated)", value=st.session_state.Softskills, key="soft_input")

        st.markdown("### 🧱 <u>Work Experience</u>", unsafe_allow_html=True)
        for idx, exp in enumerate(st.session_state.experience_entries):
            with st.expander(f"Experience #{idx+1}", expanded=True):
                exp["title"] = st.text_input("Job Title", value=exp.get("title", ""), key=f"title_{idx}_{len(st.session_state.experience_entries)}")
                exp["company"] = st.text_input("Company", value=exp.get("company", ""), key=f"company_{idx}_{len(st.session_state.experience_entries)}")
                exp["duration"] = st.text_input("Duration", value=exp.get("duration", ""), key=f"duration_{idx}_{len(st.session_state.experience_entries)}")
                exp["description"] = st.text_area("Description", value=exp.get("description", ""), key=f"description_{idx}_{len(st.session_state.experience_entries)}")

        st.markdown("### 🎓 <u>Education</u>", unsafe_allow_html=True)
        for idx, edu in enumerate(st.session_state.education_entries):
            with st.expander(f"Education #{idx+1}", expanded=True):
                edu["degree"] = st.text_input("Degree", value=edu.get("degree", ""), key=f"degree_{idx}_{len(st.session_state.education_entries)}")
                edu["institution"] = st.text_input("Institution", value=edu.get("institution", ""), key=f"institution_{idx}_{len(st.session_state.education_entries)}")
                edu["year"] = st.text_input("Year", value=edu.get("year", ""), key=f"edu_year_{idx}_{len(st.session_state.education_entries)}")
                edu["details"] = st.text_area("Details", value=edu.get("details", ""), key=f"edu_details_{idx}_{len(st.session_state.education_entries)}")

        st.markdown("### 🛠 <u>Projects</u>", unsafe_allow_html=True)
        for idx, proj in enumerate(st.session_state.project_entries):
            with st.expander(f"Project #{idx+1}", expanded=True):
                proj["title"] = st.text_input("Project Title", value=proj.get("title", ""), key=f"proj_title_{idx}_{len(st.session_state.project_entries)}")
                proj["tech"] = st.text_input("Tech Stack", value=proj.get("tech", ""), key=f"proj_tech_{idx}_{len(st.session_state.project_entries)}")
                proj["duration"] = st.text_input("Duration", value=proj.get("duration", ""), key=f"proj_duration_{idx}_{len(st.session_state.project_entries)}")
                proj["description"] = st.text_area("Description", value=proj.get("description", ""), key=f"proj_desc_{idx}_{len(st.session_state.project_entries)}")

        st.markdown("### 🔗 Project Links")
        project_links_input = st.text_area("Enter one project link per line:", value="\n".join(st.session_state.project_links), key="proj_links_input")
        if project_links_input:
            st.session_state.project_links = [link.strip() for link in project_links_input.splitlines() if link.strip()]

        st.markdown("### 🧾 <u>Certificates</u>", unsafe_allow_html=True)
        for idx, cert in enumerate(st.session_state.certificate_links):
            with st.expander(f"Certificate #{idx+1}", expanded=True):
                cert["name"] = st.text_input("Certificate Name", value=cert.get("name", ""), key=f"cert_name_{idx}_{len(st.session_state.certificate_links)}")
                cert["link"] = st.text_input("Certificate Link", value=cert.get("link", ""), key=f"cert_link_{idx}_{len(st.session_state.certificate_links)}")
                cert["duration"] = st.text_input("Duration", value=cert.get("duration", ""), key=f"cert_duration_{idx}_{len(st.session_state.certificate_links)}")
                cert["description"] = st.text_area("Description", value=cert.get("description", ""), key=f"cert_description_{idx}_{len(st.session_state.certificate_links)}")

        submitted = st.form_submit_button("📑 Generate Resume")

        if submitted:
            st.success("✅ Resume Generated Successfully! Scroll down to preview or download.")

        st.markdown("""
        <style>
            .heading-large {
                font-size: 36px;
                font-weight: bold;
                color: #336699;
            }
            .subheading-large {
                font-size: 30px;
                font-weight: bold;
                color: #336699;
            }
            .tab-section {
                margin-top: 20px;
            }
        </style>
        """, unsafe_allow_html=True)

        # --- Visual Resume Preview Section ---
        st.markdown("## 🧾 <span style='color:#336699;'>Resume Preview</span>", unsafe_allow_html=True)
        st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

        left, right = st.columns([1, 2])

        with left:
            st.markdown(f"""
                <h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>
                <h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>

                <p style='font-size:14px;'>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                </p>
            """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Skills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for skill in [s.strip() for s in st.session_state["skills"].split(",") if s.strip()]:
                st.markdown(f"<div style='margin-left:10px;'>• {skill}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Languages</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for lang in [l.strip() for l in st.session_state["languages"].split(",") if l.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {lang}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Interests</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for interest in [i.strip() for i in st.session_state["interests"].split(",") if i.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {interest}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Softskills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for Softskills  in [i.strip() for i in st.session_state["Softskills"].split(",") if i.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {Softskills}</div>", unsafe_allow_html=True)   

        with right:
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            summary_text = st.session_state['summary'].replace('\n', '<br>')
            st.markdown(f"<p style='font-size:17px;'>{summary_text}</p>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for exp in st.session_state.experience_entries:
                if exp["company"] or exp["title"]:
                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between;'>
                            <b>🏢 {exp['company']}</b><span style='color:gray;'>📆  {exp['duration']}</span>
                        </div>
                        <div style='font-size:14px;'>💼 <i>{exp['title']}</i></div>
                        <div style='font-size:17px;'>📝 {exp['description']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for edu in st.session_state.education_entries:
                if edu["institution"] or edu["degree"]:
                    st.markdown(f"""
                    <div style='margin-bottom: 15px; padding: 10px 15px;color: white; border-radius: 8px;'>
                        <div style='display: flex; justify-content: space-between; font-size: 16px; font-weight: bold;'>
                            <span>🏫 {edu['institution']}</span>
                            <span style='color: gray;'>📅 {edu['year']}</span>
                        </div>
                        <div style='font-size: 14px; margin-top: 5px;'>🎓 <i>{edu['degree']}</i></div>
                        <div style='font-size: 14px;'>📄 {edu['details']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for proj in st.session_state.project_entries:
                st.markdown(f"""
                <div style='margin-bottom:15px; padding: 10px;'>
                <strong style='font-size:16px;'>{proj['title']}</strong><br>
                <span style='font-size:14px; word-wrap:break-word; overflow-wrap:break-word; white-space:normal;'>
                   🛠️ <strong>Tech Stack:</strong> {proj['tech']}
             </span><br>
            <span style='font-size:14px;'>⏳ <strong>Duration:</strong> {proj['duration']}</span><br>
            <span style='font-size:17px;'>📝 <strong>Description:</strong> {proj['description']}</span>
            </div>
            """, unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

            if st.session_state.certificate_links:
                st.markdown("<h4 style='color:#336699;'>Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                
                for cert in st.session_state.certificate_links:
                    if cert["name"] and cert["link"]:
                        st.markdown(f"""
                        <div style='display:flex; justify-content:space-between;'>
                            <a href="{cert['link']}" target="_blank"><b>📄 {cert['name']}</b></a><span style='color:gray;'>{cert['duration']}</span>
                        </div>
                        <div style='margin-bottom:10px; font-size:14px;'>{cert['description']}</div>
                        """, unsafe_allow_html=True)

import re

with tab2:
    st.markdown("## ✨ <span style='color:#336699;'>Enhanced AI Resume Preview</span>", unsafe_allow_html=True)
    st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

    col1, spacer, col2 = st.columns([1, 0.2, 1])

    with col1:
        if st.button("🔁 Clear Preview"):
            st.session_state.pop("ai_output", None)

    with col2:
        if st.button("🚀 Generate AI Resume Preview"):
            # Normalize and ensure at least 2 experience entries
            experience_entries = st.session_state.get('experience_entries', [])
            normalized_experience_entries = []
            for entry in experience_entries:
                if isinstance(entry, dict):
                    title = entry.get("title", "")
                    desc = entry.get("description", "")
                    formatted = f"{title}\n{desc}".strip()
                else:
                    formatted = entry.strip()
                normalized_experience_entries.append(formatted)
            while len(normalized_experience_entries) < 2:
                normalized_experience_entries.append("Placeholder Experience")

            # Normalize and ensure at least 2 project entries
            project_entries = st.session_state.get('project_entries', [])
            normalized_project_entries = []
            for entry in project_entries:
                if isinstance(entry, dict):
                    title = entry.get("title", "")
                    desc = entry.get("description", "")
                    formatted = f"{title}\n{desc}".strip()
                else:
                    formatted = entry.strip()
                normalized_project_entries.append(formatted)
            while len(normalized_project_entries) < 2:
                normalized_project_entries.append("Placeholder Project")

            enhance_prompt = f"""
            You are a professional, unbiased resume optimization specialist. Create an ATS-friendly resume that uses inclusive, neutral language and avoids gendered or culturally biased terms.

            ROLE-SPECIFIC ENHANCEMENT for: "{st.session_state['job_title']}"
            Enhance all sections to align precisely with this role's industry standards, required competencies, and professional expectations.

            LANGUAGE GUIDELINES:
            - Use neutral, professional terminology
            - Avoid gendered language (e.g., "rockstar," "ninja," "guru")
            - Focus on skills, achievements, and measurable outcomes
            - Use inclusive action verbs: developed, implemented, optimized, collaborated, analyzed, designed, managed, executed

            FORMATTING REQUIREMENTS (CRITICAL - Follow exactly):
            Each section must start with the exact label followed by a colon and content on the next line.

            SECTION ENHANCEMENT RULES:

            1. SUMMARY: Write 3-4 bullet points highlighting role-specific expertise, quantifiable achievements, and core competencies. Use strong action verbs and avoid subjective adjectives.

            2. EXPERIENCE: Structure as lettered entries (A., B., C.) with:
               - Company Name (Duration) format
               - Role-specific responsibilities as bullet points
               - Focus on achievements, not just duties
               - Include metrics where possible

            3. PROJECTS: Structure as lettered entries (A., B., C.) with:
               - Project Title
               - Tech Stack: (role-relevant technologies only)
               - Duration: (timeframe)
               - Description: 4-5 bullet points covering implementation, challenges solved, technologies used, and measurable impact

            4. SKILLS: List 6-8 current, industry-standard technical skills relevant to the role
            5. SOFTSKILLS: List 6-8 professional competencies using neutral language
            6. LANGUAGES: List only spoken/written languages
            7. INTERESTS: List 3-6 professional interests aligned with the role
            8. CERTIFICATES: List 3-6 real, industry-recognized certifications with provider and duration

            DOMAIN-SPECIFIC REQUIREMENTS:
            - For Technical Roles: Focus on programming languages, frameworks, tools, methodologies
            - For Security Roles: Emphasize security tools, compliance standards, threat analysis
            - For Data Roles: Highlight analytics tools, statistical methods, visualization platforms
            - For Management Roles: Stress leadership frameworks, process improvement, team development

            OUTPUT FORMAT (EXACT STRUCTURE REQUIRED):

            Summary:
            • [Achievement-focused bullet point with quantifiable impact]
            • [Core competency statement with role-specific expertise]
            • [Professional strength with measurable outcome]

            Experience:
            A. [Company Name] ([Duration])
               • [Role Title]
               • [Specific responsibility with measurable outcome]
               • [Achievement or project contribution]
               • [Process improvement or efficiency gain]

            B. [Company Name] ([Duration])
               • [Role Title]
               • [Specific responsibility with measurable outcome]
               • [Achievement or project contribution]

            Projects:
            A. [Project Title]
               • Tech Stack: [Relevant technologies only]
               • Duration: [Start – End timeframe]
               • Description:
                 - [Specific implementation or feature developed]
                 - [Technology used and its application context]
                 - [Performance improvement or problem solved with metrics]
                 - [Collaborative achievement or technical innovation]
                 - [Additional impact or learning outcome]

            B. [Project Title]
               • Tech Stack: [Relevant technologies only]
               • Duration: [Start – End timeframe]
               • Description:
                 - [Specific implementation details]
                 - [Technical challenges addressed]
                 - [Measurable results or improvements]
                 - [Skills demonstrated or technologies mastered]

            Skills:
            [Skill 1], [Skill 2], [Skill 3], [Skill 4], [Skill 5], [Skill 6]

            SoftSkills:
            [Professional Competency 1], [Professional Competency 2], [Professional Competency 3], [Professional Competency 4], [Professional Competency 5], [Professional Competency 6]

            Languages:
            [Language 1], [Language 2], [Language 3]

            Interests:
            [Professional Interest 1], [Professional Interest 2], [Professional Interest 3], [Professional Interest 4]

            Certificates:
            [Certificate Name] – [Provider] ([Duration/Level])
            [Certificate Name] – [Provider] ([Duration/Level])
            [Certificate Name] – [Provider] ([Duration/Level])

            ENHANCEMENT SOURCE DATA:
            Transform and enhance the following user inputs while maintaining accuracy and relevance:

            Summary:
            {st.session_state['summary']}

            Experience:
            {normalized_experience_entries}

            Projects:
            {normalized_project_entries}

            Skills:
            {st.session_state['skills']}

            SoftSkills:
            {st.session_state['Softskills']}

            Languages:
            {st.session_state['languages']}

            Interests:
            {st.session_state['interests']}

            Certificates:
            {[cert['name'] for cert in st.session_state['certificate_links'] if cert['name']]}
            """

            with st.spinner("🧠 Thinking..."):
                ai_output = call_llm(enhance_prompt, session=st.session_state)
                st.session_state["ai_output"] = ai_output

    # ------------------------- PARSE + RENDER -------------------------
    if "ai_output" in st.session_state:
        ai_output = st.session_state["ai_output"]

        def extract_section(label, output, default=""):
            match = re.search(rf"{label}:\s*(.*?)(?=\n\w+:|\Z)", output, re.DOTALL)
            return match.group(1).strip() if match else default

        summary_enhanced = extract_section("Summary", ai_output, st.session_state['summary'])
        experience_raw = extract_section("Experience", ai_output)
        experience_blocks = re.split(r"\n(?=[A-Z]\. )", experience_raw.strip())
        projects_raw = extract_section("Projects", ai_output)
        projects_blocks = re.split(r"\n(?=[A-Z]\. )", projects_raw.strip())
        skills_list = extract_section("Skills", ai_output, st.session_state['skills'])
        softskills_list = extract_section("SoftSkills", ai_output, st.session_state['Softskills'])
        languages_list = extract_section("Languages", ai_output, st.session_state['languages'])
        interests_list = extract_section("Interests", ai_output, st.session_state['interests'])
        certificates_list = extract_section("Certificates", ai_output)

        # ------------------------- UI RENDER -------------------------
        left, right = st.columns([1, 2])

        with left:
            st.markdown(f"""
                <h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>
                <h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>
                <p style='font-size:14px;'>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                </p>
            """, unsafe_allow_html=True)

            def render_bullet_section(title, items):
                st.markdown(f"<h4 style='color:#336699;'>{title}</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for item in [i.strip() for i in items.split(",") if i.strip()]:
                    st.markdown(f"<div style='margin-left:10px;'>• {item}</div>", unsafe_allow_html=True)

            render_bullet_section("Skills", skills_list)
            render_bullet_section("Languages", languages_list)
            render_bullet_section("Interests", interests_list)
            render_bullet_section("Soft Skills", softskills_list)

        with right:
            formatted_summary = summary_enhanced.replace('\n• ', '<br>• ').replace('\n', '<br>')
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            st.markdown(f"<p style='font-size:17px;'>{formatted_summary}</p>", unsafe_allow_html=True)

            # Experience
            if experience_blocks:
                st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                experience_titles = [entry.get("title", "").strip().upper() for entry in st.session_state.experience_entries]
                for idx, exp_block in enumerate(experience_blocks):
                    lines = exp_block.strip().split("\n")
                    if not lines:
                        continue
                    heading = lines[0]
                    description_lines = lines[1:]
                    match = re.match(r"[A-Z]\.\s*(.+?)\s*\((.*?)\)", heading)
                    company, duration = (match.group(1).strip(), match.group(2).strip()) if match else (heading, "")
                    role = experience_titles[idx] if idx < len(experience_titles) else ""
                    formatted_exp = "<br>".join(description_lines)

                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between;'>
                            <b>🏢 {company.upper()}</b><span style='color:gray;'>📆 {duration}</span>
                        </div>
                        <div style='font-size:14px;'>💼 <i>{role}</i></div>
                        <div style='font-size:17px;'>📝 {formatted_exp}</div>
                    </div>
                    """, unsafe_allow_html=True)

            # Education
            st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for edu in st.session_state.education_entries:
                st.markdown(f"""
                <div style='margin-bottom:15px; padding:10px 15px; border-radius:8px;'>
                    <div style='display: flex; justify-content: space-between; font-size: 16px; font-weight: bold;'>
                        <span>🏫 {edu['institution']}</span>
                        <span style='color: gray;'>📅 {edu['year']}</span>
                    </div>
                    <div style='font-size: 14px;'>🎓 <i>{edu['degree']}</i></div>
                    <div style='font-size: 14px;'>📄 {edu['details']}</div>
                </div>
                """, unsafe_allow_html=True)

            # Projects
            if projects_blocks:
                st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for idx, proj_block in enumerate(projects_blocks):
                    proj = st.session_state.project_entries[idx] if idx < len(st.session_state.project_entries) else {}
                    title = proj.get("title", "")
                    tech = proj.get("tech", "")
                    duration = proj.get("duration", "")
                    description = proj_block
                    for keyword in [title, f"Tech Stack: {tech}", f"Duration: {duration}"]:
                        if keyword and keyword in description:
                            description = description.replace(keyword, "")
                    formatted_proj = description.strip().replace('\n• ', '<br>• ').replace('\n', '<br>')
                    label = chr(65 + idx)

                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding: 10px;'>
                        <strong style='font-size:16px;'>📌 <span style='color:#444;'>{label}. </span>{title}</strong><br>
                        <span style='font-size:14px;'>🛠️ <strong>Tech Stack:</strong> {tech}</span><br>
                        <span style='font-size:14px;'>⏳ <strong>Duration:</strong> {duration}</span><br>
                        <span style='font-size:17px;'>📄 <strong>Description:</strong></span><br>
                        <div style='margin-top:4px; font-size:15px;'>{formatted_proj}</div>
                    </div>
                    """, unsafe_allow_html=True)

            # Certificates
            if certificates_list:
                st.markdown("<h4 style='color:#336699;'>📜 Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                certs = re.split(r"\n|(?<=\))(?=\s*[A-Z])|(?<=[a-z]\))(?= [A-Z])", certificates_list)
                for cert in [c.strip() for c in certs if c.strip()]:
                    st.markdown(f"<div style='margin-left:10px;'>• {cert}</div>", unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

    # Generate HTML content based on selected template
    if submitted:
        # Determine which template to use
        if selected_template == "Default (Professional)":
            html_content = render_template_default(st.session_state, profile_img_html)
        elif selected_template == "Modern Minimal":
            html_content = render_template_modern(st.session_state, profile_img_html)
        elif selected_template == "Elegant Sidebar":
            html_content = render_template_sidebar(st.session_state, profile_img_html)
        else:
            # Fallback to default
            html_content = render_template_default(st.session_state, profile_img_html)

        # Store the generated content
        st.session_state["generated_html"] = html_content

with tab2:
    # ==========================
    # 📥 Resume Download Header
    # ==========================
    if "generated_html" in st.session_state:
        st.markdown(
            """
            <div style='text-align: center; margin-top: 20px; margin-bottom: 30px;'>
                <h2 style='color: #2f4f6f; font-family: Arial, sans-serif; font-size: 24px;'>
                    📥 Download Your Resume
                </h2>
                <p style="color:#555; font-size:14px;">
                    Choose your preferred format below
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )

        col1, = st.columns(1)

        # HTML Resume Download Button
        with col1:
            html_bytes = st.session_state["generated_html"].encode("utf-8")
            html_file = BytesIO(html_bytes)
            
            st.download_button(
                label="⬇️ Download as Template",
                data=html_file,
                file_name=f"{st.session_state['name'].replace(' ', '_')}_Resume.html",
                mime="text/html",
                key="download_resume_html"
            )

        # PDF Resume Download Button
        pdf_resume_bytes = html_to_pdf_bytes(st.session_state["generated_html"])
        
        # ✅ Extra Help Note
        st.markdown("""
        ✅ After downloading your HTML resume, you can 
        <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
        convert it to PDF using Sejda's free online tool</a>.
        """, unsafe_allow_html=True)

        # ==========================
        # 📩 Cover Letter Expander
        # ==========================
        with st.expander("📩 Generate Cover Letter from This Resume"):
            generate_cover_letter_from_resume_builder()

        # ==========================
        # ✉️ Generated Cover Letter Downloads (NO PREVIEW HERE)
        # ==========================
        if "cover_letter" in st.session_state:
            st.markdown(
                """
                <div style="margin-top: 30px; margin-bottom: 20px;">
                    <h3 style="color: #003366;">✉️ Generated Cover Letter</h3>
                    <p style="color:#555; font-size:14px;">
                        You can download your generated cover letter in multiple formats.
                    </p>
                </div>
                """,
                unsafe_allow_html=True
            )

            # ✅ Use already-rendered HTML from session (don't show again)
            styled_cover_letter = st.session_state.get("cover_letter_html", "")

            # ✅ Generate PDF from styled HTML
            pdf_file = html_to_pdf_bytes(styled_cover_letter)

            # ✅ DOCX Generator (preserves line breaks)
            def create_docx_from_text(text, filename="cover_letter.docx"):
                from docx import Document
                bio = BytesIO()
                doc = Document()
                doc.add_heading("Cover Letter", 0)

                for line in text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
                    else:
                        doc.add_paragraph("")  # preserve empty lines

                doc.save(bio)
                bio.seek(0)
                return bio

            # ==========================
            # 📥 Cover Letter Download Buttons
            # ==========================
            st.markdown("""
            <div style="margin-top: 25px; margin-bottom: 15px;">
                <strong>⬇️ Download Your Cover Letter:</strong>
            </div>
            """, unsafe_allow_html=True)

            col1,col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 Download Cover Letter (.docx)",
                    data=create_docx_from_text(st.session_state["cover_letter"]),
                    file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_coverletter_docx"
                )
            
            with col2:
                st.download_button(
                    label="📥 Download Cover Letter (Template)",
                    data=styled_cover_letter.encode("utf-8"),
                    file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.html",
                    mime="text/html",
                    key="download_coverletter_html"
                )

            # ✅ Helper note
            st.markdown("""
            ✅ If the HTML cover letter doesn't display properly, you can 
            <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
            convert it to PDF using Sejda's free online tool</a>.
            """, unsafe_allow_html=True)

FEATURED_COMPANIES = {
    "tech": [
        {
            "name": "Google",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/2/2f/Google_2015_logo.svg",
            "color": "#4285F4",
            "careers_url": "https://careers.google.com",
            "description": "Leading technology company known for search, cloud, and innovation",
            "categories": ["Software", "AI/ML", "Cloud", "Data Science"]
        },
        {
            "name": "Microsoft",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/4/44/Microsoft_logo.svg",
            "color": "#00A4EF",
            "careers_url": "https://careers.microsoft.com",
            "description": "Global leader in software, cloud, and enterprise solutions",
            "categories": ["Software", "Cloud", "Gaming", "Enterprise"]
        },
        {
            "name": "Amazon",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/a/a9/Amazon_logo.svg",
            "color": "#FF9900",
            "careers_url": "https://www.amazon.jobs",
            "description": "E-commerce and cloud computing giant",
            "categories": ["Software", "Operations", "Cloud", "Retail"]
        },
        {
            "name": "Apple",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/f/fa/Apple_logo_black.svg",
            "color": "#555555",
            "careers_url": "https://www.apple.com/careers",
            "description": "Innovation leader in consumer technology",
            "categories": ["Software", "Hardware", "Design", "AI/ML"]
        },
        {
            "name": "Facebook",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/0/05/Facebook_Logo_%282019%29.png",
            "color": "#1877F2",
            "careers_url": "https://www.metacareers.com/",
            "description": "Social media and technology company",
            "categories": ["Software", "Marketing", "Networking", "AI/ML"]
        },
        {
            "name": "Netflix",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/0/08/Netflix_2015_logo.svg",
            "color": "#E50914",
            "careers_url": "https://explore.jobs.netflix.net/careers",
            "description": "Streaming media company",
            "categories": ["Software", "Marketing", "Design", "Service"],
            "website": "https://jobs.netflix.com/",
            "industry": "Entertainment & Technology"
        }
    ],
    "indian_tech": [
        {
            "name": "TCS",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/f/f6/TCS_New_Logo.svg",
            "color": "#0070C0",
            "careers_url": "https://www.tcs.com/careers",
            "description": "India's largest IT services company",
            "categories": ["IT Services", "Consulting", "Digital"]
        },
        {
            "name": "Infosys",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/5/55/Infosys_logo.svg",
            "color": "#007CC3",
            "careers_url": "https://www.infosys.com/careers",
            "description": "Global leader in digital services and consulting",
            "categories": ["IT Services", "Consulting", "Digital"]
        },
        {
            "name": "Wipro",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/8/80/Wipro_Primary_Logo_Color_RGB.svg",
            "color": "#341F65",
            "careers_url": "https://careers.wipro.com",
            "description": "Leading global information technology company",
            "categories": ["IT Services", "Consulting", "Digital"]
        },
        {
            "name": "HCL",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/5/5e/HCL_Technologies_logo.svg",
            "color": "#0075C9",
            "careers_url": "https://www.hcltech.com/careers",
            "description": "Global technology company",
            "categories": ["IT Services", "Engineering", "Digital"]
        }
    ],
    "global_corps": [
        {
            "name": "IBM",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/5/51/IBM_logo.svg",
            "color": "#1F70C1",
            "careers_url": "https://www.ibm.com/careers",
            "description": "Global leader in technology and consulting",
            "categories": ["Software", "Consulting", "AI/ML", "Cloud"],
            "website": "https://www.ibm.com/careers/",
            "industry": "Technology & Consulting"
        },
        {
            "name": "Accenture",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/8/80/Accenture_Logo.svg",
            "color": "#A100FF",
            "careers_url": "https://www.accenture.com/careers",
            "description": "Global professional services company",
            "categories": ["Consulting", "Technology", "Digital"]
        },
        {
            "name": "Cognizant",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/6/6e/Cognizant_logo_2022.svg",
            "color": "#1299D8",
            "careers_url": "https://careers.cognizant.com",
            "description": "Leading professional services company",
            "categories": ["IT Services", "Consulting", "Digital"]
        }
    ]
}


JOB_MARKET_INSIGHTS = {
    "trending_skills": [
        {"name": "Artificial Intelligence", "growth": "+45%", "icon": "fas fa-brain"},
        {"name": "Cloud Computing", "growth": "+38%", "icon": "fas fa-cloud"},
        {"name": "Data Science", "growth": "+35%", "icon": "fas fa-chart-line"},
        {"name": "Cybersecurity", "growth": "+32%", "icon": "fas fa-shield-alt"},
        {"name": "DevOps", "growth": "+30%", "icon": "fas fa-code-branch"},
        {"name": "Machine Learning", "growth": "+28%", "icon": "fas fa-robot"},
        {"name": "Blockchain", "growth": "+25%", "icon": "fas fa-lock"},
        {"name": "Big Data", "growth": "+23%", "icon": "fas fa-database"},
        {"name": "Internet of Things", "growth": "+21%", "icon": "fas fa-wifi"}
    ],
    "top_locations": [
        {"name": "Bangalore", "jobs": "50,000+", "icon": "fas fa-city"},
        {"name": "Mumbai", "jobs": "35,000+", "icon": "fas fa-city"},
        {"name": "Delhi NCR", "jobs": "30,000+", "icon": "fas fa-city"},
        {"name": "Hyderabad", "jobs": "25,000+", "icon": "fas fa-city"},
        {"name": "Pune", "jobs": "20,000+", "icon": "fas fa-city"},
        {"name": "Chennai", "jobs": "15,000+", "icon": "fas fa-city"},
        {"name": "Noida", "jobs": "10,000+", "icon": "fas fa-city"},
        {"name": "Vadodara", "jobs": "7,000+", "icon": "fas fa-city"},
        {"name": "Ahmedabad", "jobs": "6,000+", "icon": "fas fa-city"},
        {"name": "Remote", "jobs": "3,000+", "icon": "fas fa-globe-americas"},
    ],
    "salary_insights": [
        {"role": "Machine Learning Engineer", "range": "10-35 LPA", "experience": "0-5 years"},
        {"role": "Big Data Engineer", "range": "8-30 LPA", "experience": "0-5 years"},
        {"role": "Software Engineer", "range": "5-25 LPA", "experience": "0-5 years"},
        {"role": "Data Scientist", "range": "8-30 LPA", "experience": "0-5 years"},
        {"role": "DevOps Engineer", "range": "6-28 LPA", "experience": "0-5 years"},
        {"role": "UI/UX Designer", "range": "5-25 LPA", "experience": "0-5 years"},
        {"role": "Full Stack Developer", "range": "8-30 LPA", "experience": "0-5 years"},
        {"role": "C++/C#/Python/Java Developer", "range": "6-26 LPA", "experience": "0-5 years"},
        {"role": "Django Developer", "range": "7-27 LPA", "experience": "0-5 years"},
        {"role": "Cloud Engineer", "range": "6-26 LPA", "experience": "0-5 years"},
        {"role": "Google Cloud/AWS/Azure Engineer", "range": "6-26 LPA", "experience": "0-5 years"},
        {"role": "Salesforce Engineer", "range": "6-26 LPA", "experience": "0-5 years"},
    ]
}

def get_featured_companies(category=None):
    """Get featured companies with original logos, optionally filtered by category"""
    def has_valid_logo(company):
        return "logo_url" in company and company["logo_url"].startswith("https://upload.wikimedia.org/")

    if category and category in FEATURED_COMPANIES:
        return [company for company in FEATURED_COMPANIES[category] if has_valid_logo(company)]

    return [
        company for companies in FEATURED_COMPANIES.values()
        for company in companies if has_valid_logo(company)
    ]


def get_market_insights():
    """Get job market insights"""
    return JOB_MARKET_INSIGHTS

def get_company_info(company_name):
    """Get company information by name"""
    for companies in FEATURED_COMPANIES.values():
        for company in companies:
            if company["name"] == company_name:
                return company
    return None

def get_companies_by_industry(industry):
    """Get list of companies by industry"""
    companies = []
    for companies_list in FEATURED_COMPANIES.values():
        for company in companies_list:
            if "industry" in company and company["industry"] == industry:
                companies.append(company)
    return companies

# Sample job search function
import uuid
import urllib.parse
import sqlite3
import datetime
import streamlit as st
from zoneinfo import ZoneInfo
import requests
import re

# ✅ RapidAPI Configuration (from Streamlit secrets)
RAPID_API_KEY = st.secrets["rapidapi"]["key"]
RAPID_API_HOST = st.secrets["rapidapi"]["host"]

def clean_html(raw_html: str) -> str:
    """Remove HTML tags and comments from API descriptions."""
    if not raw_html:
        return ""
    # Remove comments
    raw_html = re.sub(r"<!--.*?-->", "", raw_html, flags=re.DOTALL)
    # Remove all tags
    return re.sub(r"<.*?>", "", raw_html).strip()

def fetch_live_jobs(job_role, location, job_type=None, remote_only=False, results=10):
    url = f"https://{RAPID_API_HOST}/search"
    querystring = {
        "query": f"{job_role} in {location}",
        "page": "1",
        "num_pages": "1",
        "remote_jobs_only": str(remote_only).lower()
    }

    # 🔹 Map UI dropdown values to RapidAPI accepted filters
    type_map = {
        "Full-time": "FULLTIME",
        "Part-time": "PARTTIME",
        "Contract": "CONTRACTOR",
        "Internship": "INTERN",
        "Temporary": "TEMPORARY",
        "Volunteer": "VOLUNTEER"
    }
    if job_type and job_type in type_map:
        querystring["employment_types"] = type_map[job_type]

    headers = {
        "X-RapidAPI-Key": RAPID_API_KEY,
        "X-RapidAPI-Host": RAPID_API_HOST
    }
    try:
        response = requests.get(url, headers=headers, params=querystring)
        if response.status_code == 200:
            return response.json().get("data", [])[:results]
        else:
            return []
    except Exception:
        return []

def fetch_company_by_domain(domain: str):
    """Fetch company information by domain using LinkedIn Data API"""
    url = f"https://linkedin-data-api.p.rapidapi.com/get-company-by-domain?domain={domain}"
    headers = {
        "X-RapidAPI-Key": RAPID_API_KEY,
        "X-RapidAPI-Host": "linkedin-data-api.p.rapidapi.com"
    }
    try:
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            return response.json()
        else:
            return None
    except Exception:
        return None

def unified_search(job_role, location, experience_level=None, job_type=None, foundit_experience=None):
    results = []

    # 1️⃣ Fetch live jobs from RapidAPI JSearch
    live_jobs = fetch_live_jobs(job_role, location, job_type=job_type, results=5)
    for job in live_jobs:
        results.append({
            "platform": "RapidAPI (Live)",
            "title": clean_html(job.get("job_title", "N/A")),
            "company": clean_html(job.get("employer_name", "Unknown")),
            "location": f"{job.get('job_city','')}, {job.get('job_country','')}",
            "salary": f"{job.get('job_min_salary','NA')} - {job.get('job_max_salary','NA')} {job.get('job_salary_currency','')}",
            "date": job.get("job_posted_at_datetime_utc", "N/A"),
            "type": job.get("job_employment_type","N/A"),
            "remote": "Remote" if job.get("job_is_remote") else "On-site",
            "publisher": clean_html(job.get("job_publisher","N/A")),
            "description": clean_html(job.get("job_description",""))[:200] + "...",
            "apply_link": job.get("job_apply_link", "#")
        })

    # 2️⃣ Add LinkedIn, Naukri, FoundIt links (existing function)
    external_links = search_jobs(job_role, location, experience_level, job_type, foundit_experience)
    for job in external_links:
        results.append({
            "platform": job["title"].split(":")[0],
            "title": job["title"].split(":")[1].strip(),
            "company": "N/A",
            "location": location,
            "salary": "Check site",
            "date": "N/A",
            "type": "N/A",
            "remote": "N/A",
            "publisher": job["title"].split(":")[0],
            "description": "Open this platform to view full details.",
            "apply_link": job["link"]
        })

    return results


# Database functions for job search history
def init_job_search_db():
    """Initialize the job search database and create user_jobs table if not exists"""
    try:
        conn = sqlite3.connect('resume_data.db')
        cursor = conn.cursor()

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS user_jobs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL,
                role TEXT NOT NULL,
                location TEXT NOT NULL,
                platform TEXT NOT NULL,
                url TEXT NOT NULL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Database initialization error: {e}")

def save_job_search(username, role, location, results):
    """Save job search results to database for logged-in user"""
    if not username:
        return

    try:
        conn = sqlite3.connect('resume_data.db')
        cursor = conn.cursor()

        for result in results:
            # Extract platform name from title or use platform field
            platform = result.get("platform", "Unknown")
            url = result.get("apply_link", "#")

            cursor.execute('''
                INSERT INTO user_jobs (username, role, location, platform, url, timestamp)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (username, role, location, platform, url, datetime.datetime.now()))

        conn.commit()
        conn.close()

    except Exception as e:
        st.error(f"Error saving job search: {e}")

def prune_old_searches(username):
    """Keep only the last 50 saved job searches per user (optional cleanup)"""
    if not username:
        return

    try:
        conn = sqlite3.connect('resume_data.db')
        cursor = conn.cursor()

        # Delete all but the most recent 50 searches for this user
        cursor.execute('''
            DELETE FROM user_jobs
            WHERE username = ? AND id NOT IN (
                SELECT id FROM user_jobs
                WHERE username = ?
                ORDER BY timestamp DESC
                LIMIT 50
            )
        ''', (username, username))

        conn.commit()
        conn.close()

    except Exception as e:
        st.error(f"Error pruning old searches: {e}")

def delete_saved_job_search(search_id):
    """Delete a saved job search by its ID"""
    try:
        conn = sqlite3.connect('resume_data.db')
        cursor = conn.cursor()

        cursor.execute('DELETE FROM user_jobs WHERE id = ?', (search_id,))

        conn.commit()
        conn.close()

    except Exception as e:
        st.error(f"Error deleting job search: {e}")

def get_saved_job_searches(username, limit=10, offset=0, platform_filter=None):
    """Get saved job searches for a user with filtering and pagination"""
    if not username:
        return []

    try:
        conn = sqlite3.connect('resume_data.db')
        cursor = conn.cursor()

        # Build the query with optional platform filter
        if platform_filter and platform_filter != "All":
            cursor.execute('''
                SELECT id, role, location, platform, url, timestamp
                FROM user_jobs
                WHERE username = ? AND platform = ?
                ORDER BY timestamp DESC
                LIMIT ? OFFSET ?
            ''', (username, platform_filter, limit, offset))
        else:
            cursor.execute('''
                SELECT id, role, location, platform, url, timestamp
                FROM user_jobs
                WHERE username = ?
                ORDER BY timestamp DESC
                LIMIT ? OFFSET ?
            ''', (username, limit, offset))

        results = cursor.fetchall()
        conn.close()

        return [
            {
                "id": row[0],
                "role": row[1],
                "location": row[2],
                "platform": row[3],
                "url": row[4],
                "timestamp": row[5]
            }
            for row in results
        ]
    except Exception as e:
        st.error(f"Error fetching saved searches: {e}")
        return []

def get_total_saved_searches_count(username, platform_filter=None):
    """Get total count of saved searches for pagination"""
    if not username:
        return 0

    try:
        conn = sqlite3.connect('resume_data.db')
        cursor = conn.cursor()

        if platform_filter and platform_filter != "All":
            cursor.execute('SELECT COUNT(*) FROM user_jobs WHERE username = ? AND platform = ?', (username, platform_filter))
        else:
            cursor.execute('SELECT COUNT(*) FROM user_jobs WHERE username = ?', (username,))

        count = cursor.fetchone()[0]
        conn.close()

        return count
    except Exception as e:
        st.error(f"Error getting search count: {e}")
        return 0

def get_available_platforms(username):
    """Get list of platforms that the user has searched on"""
    if not username:
        return []

    try:
        conn = sqlite3.connect('resume_data.db')
        cursor = conn.cursor()

        cursor.execute('SELECT DISTINCT platform FROM user_jobs WHERE username = ? ORDER BY platform', (username,))

        platforms = [row[0] for row in cursor.fetchall()]
        conn.close()

        return platforms
    except Exception as e:
        st.error(f"Error fetching platforms: {e}")
        return []

def slugify(text: str) -> str:
    """Convert text into a safe slug (lowercase, hyphenated, no special chars)."""
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"\s+", "-", text)
    return text

def render_job_card(title, link, platform_name, brand_color, platform_gradient, company=None, location=None, salary=None, description=None):
    """
    Reusable function to render a modern job card with consistent styling.

    Args:
        title: Job title or role
        link: Apply link URL
        platform_name: Name of the platform (LinkedIn, Naukri, etc.)
        brand_color: Platform brand color (hex)
        platform_gradient: CSS gradient for platform
        company: Company name (optional)
        location: Job location (optional)
        salary: Salary information (optional)
        description: Job description (optional)

    Returns:
        tuple: (html_string, estimated_height)
    """
    # Platform icon mapping
    icon_map = {
        "LinkedIn": "🔵",
        "Naukri": "🏢",
        "FoundIt (Monster)": "🌐",
        "RapidAPI (Live)": "⚡"
    }
    icon = icon_map.get(platform_name, "📄")

    # Build metadata section and calculate height
    metadata_html = ""
    estimated_height = 180  # Base height (platform + title + button + padding)

    if company:
        metadata_html += f"""
        <div style="color: #aaaaaa; font-size: 14px; margin-bottom: 8px; z-index: 2; position: relative;">
            🏢 <b>{company}</b>
        </div>
        """
        estimated_height += 30

    if location:
        metadata_html += f"""
        <div style="color: #aaaaaa; font-size: 14px; margin-bottom: 8px; z-index: 2; position: relative;">
            📍 {location}
        </div>
        """
        estimated_height += 30

    if salary and salary not in ["Check site", "N/A - N/A "]:
        metadata_html += f"""
        <div style="color: #aaaaaa; font-size: 14px; margin-bottom: 8px; z-index: 2; position: relative;">
            💰 {salary}
        </div>
        """
        estimated_height += 30

    if description and description != "Open this platform to view full details.":
        # Estimate height based on description length
        desc_lines = len(description) // 60 + 1
        estimated_height += (desc_lines * 22) + 15
        metadata_html += f"""
        <div style="color: #999999; font-size: 14px; margin-bottom: 15px; line-height: 1.6; z-index: 2; position: relative;">
            {description}
        </div>
        """

    # Create the job card HTML
    job_card_html = f"""
<!DOCTYPE html>
<html>
<head>
<style>
    * {{
        margin: 0;
        padding: 0;
        box-sizing: border-box;
    }}
    body {{
        background: transparent;
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    }}
    @keyframes shimmer {{
        0% {{ transform: translateX(-100%); }}
        100% {{ transform: translateX(100%); }}
    }}
    .shimmer-overlay {{
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.05), transparent);
        transform: translateX(-100%);
        animation: shimmer 3s infinite;
        z-index: 1;
    }}
    .job-result-card {{
        background: linear-gradient(135deg, #1e1e1e 0%, #2d2d2d 100%);
        padding: 22px;
        border-radius: 20px;
        border-left: 6px solid {brand_color};
        box-shadow: 0 8px 32px rgba(0,0,0,0.3), 0 0 20px {brand_color}40;
        position: relative;
        overflow: hidden;
        transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
    }}
    .job-result-card:hover {{
        transform: translateY(-3px);
        box-shadow: 0 12px 40px rgba(0,0,0,0.4), 0 0 30px {brand_color}60;
    }}
    .job-button {{
        background: {platform_gradient};
        color: white;
        padding: 12px 20px;
        border: none;
        border-radius: 12px;
        font-size: 16px;
        font-weight: bold;
        cursor: pointer;
        box-shadow: 0 4px 15px {brand_color}50;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
        text-decoration: none;
        display: inline-block;
    }}
    .job-button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 6px 20px {brand_color}70;
    }}
</style>
</head>
<body>
<div class="job-result-card">
    <div class="shimmer-overlay"></div>

    <!-- Platform Badge -->
    <div style="font-size: 20px; margin-bottom: 12px; z-index: 2; position: relative; font-weight: bold; color: {brand_color};">
        {icon} {platform_name}
    </div>

    <!-- Job Title -->
    <div style="color: #ffffff; font-size: 18px; margin-bottom: 12px; font-weight: bold; z-index: 2; position: relative; line-height: 1.4;">
        {title}
    </div>

    <!-- Metadata (company, location, salary, description) -->
    {metadata_html}

    <!-- Apply Button -->
    <a href="{link}" target="_blank" style="text-decoration: none; z-index: 2; position: relative;">
        <button class="job-button">
            <span style="position: relative; z-index: 2;">🚀 Apply Now →</span>
        </button>
    </a>
</div>
</body>
</html>
"""
    return job_card_html, estimated_height

def search_jobs(job_role, location, experience_level=None, job_type=None, foundit_experience=None):
    # Encode query values
    role_encoded = urllib.parse.quote_plus(job_role.strip())
    loc_encoded = urllib.parse.quote_plus(location.strip())

    # Slugs
    role_path_naukri = job_role.strip().lower().replace(" ", "-")
    city_part = location.strip().split(",")[0].strip()
    city_naukri = city_part.lower().replace(" ", "-")
    # Only encode what the user entered for the query
    city_query_naukri = urllib.parse.quote_plus(location.strip())

    # FoundIt slugs
    role_path_foundit = slugify(job_role)
    city_path_foundit = slugify(city_part)

    # Experience mappings
    experience_range_map = {
        "Internship": "0~0", "Entry Level": "1~1", "Associate": "2~3",
        "Mid-Senior Level": "4~7", "Director": "8~15", "Executive": "16~20"
    }
    experience_exact_map = {
        "Internship": "0", "Entry Level": "1", "Associate": "2",
        "Mid-Senior Level": "4", "Director": "8", "Executive": "16"
    }
    linkedin_exp_map = {
        "Internship": "1", "Entry Level": "2", "Associate": "3",
        "Mid-Senior Level": "4", "Director": "5", "Executive": "6"
    }
    job_type_map = {
        "Full-time": "F", "Part-time": "P", "Contract": "C",
        "Temporary": "T", "Volunteer": "V", "Internship": "I"
    }

    # LinkedIn URL
    linkedin_url = f"https://www.linkedin.com/jobs/search/?keywords={role_encoded}&location={loc_encoded}"
    if experience_level in linkedin_exp_map:
        linkedin_url += f"&f_E={linkedin_exp_map[experience_level]}"
    if job_type in job_type_map:
        linkedin_url += f"&f_JT={job_type_map[job_type]}"

    # Determine experience values
    if foundit_experience is not None:
        experience_range = f"{foundit_experience}~{foundit_experience}"
        experience_exact = str(foundit_experience)
    else:
        experience_range = experience_range_map.get(experience_level, "")
        experience_exact = experience_exact_map.get(experience_level, "")

    # Naukri URL – no forced "and-india"
    naukri_url = (
        f"https://www.naukri.com/{role_path_naukri}-jobs-in-{city_naukri}"
        f"?k={role_encoded}&l={city_query_naukri}"
    )
    if experience_exact:
        naukri_url += f"&experience={experience_exact}"
    naukri_url += "&nignbevent_src=jobsearchDeskGNB"

    # FoundIt URL
    search_id = uuid.uuid4()
    child_search_id = uuid.uuid4()
    if role_path_foundit and city_path_foundit:
        foundit_url = (
            f"https://www.foundit.in/search/{role_path_foundit}-jobs-in-{city_path_foundit}"
            f"?query={role_encoded}&locations={loc_encoded}"
            f"&experienceRanges={urllib.parse.quote_plus(experience_range)}"
            f"&experience={experience_exact}"
            f"&queryDerived=true"
            f"&searchId={search_id}&child_search_id={child_search_id}"
        )
    else:
        foundit_url = (
            f"https://www.foundit.in/search/result?query={role_encoded}&locations={loc_encoded}"
            f"&experienceRanges={urllib.parse.quote_plus(experience_range)}"
            f"&experience={experience_exact}"
            f"&queryDerived=true"
            f"&searchId={search_id}&child_search_id={child_search_id}"
        )

    return [
        {"title": f"LinkedIn: {job_role} jobs in {location}", "link": linkedin_url},
        {"title": f"Naukri: {job_role} jobs in {location}", "link": naukri_url},
        {"title": f"FoundIt (Monster): {job_role} jobs in {location}", "link": foundit_url}
    ]



def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A function to add a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color and underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    color_element = OxmlElement('w:color')
    color_element.set(qn('w:val'), color)
    rPr.append(color_element)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# Initialize database
init_job_search_db()

# Your existing tab3 code with enhanced CSS styling
with tab3:
    st.markdown("<h1 style='text-align: center; color: #ffffff; margin-bottom: 30px;'>🟦 Job Search Hub</h1>", unsafe_allow_html=True)

    # Initialize session state for search mode
    if 'search_mode' not in st.session_state:
        st.session_state.search_mode = "External Platforms"

    # Modern Toggle Switch with Circular Indicator
    is_external = st.session_state.search_mode == "External Platforms"

    toggle_html = f"""
    <style>
    .toggle-switch-container {{
        display: flex;
        justify-content: center;
        align-items: center;
        margin-bottom: 30px;
        gap: 0;
    }}
    .toggle-option {{
        background: rgba(40, 40, 40, 0.95);
        padding: 18px 35px;
        color: rgba(255, 255, 255, 0.4);
        font-size: 15px;
        font-weight: 600;
        border: 1px solid rgba(255, 255, 255, 0.15);
        display: flex;
        align-items: center;
        gap: 12px;
        transition: all 0.3s ease;
        cursor: pointer;
        position: relative;
    }}
    .toggle-option.left {{
        border-radius: 16px 0 0 16px;
        border-right: none;
    }}
    .toggle-option.right {{
        border-radius: 0 16px 16px 0;
        border-left: none;
    }}
    .toggle-option.active {{
        color: #ffffff;
    }}
    .toggle-option.active.external {{
        background: linear-gradient(135deg, #2196F3 0%, #1976D2 100%);
        border-color: #1976D2;
    }}
    .toggle-option.active.rapid {{
        background: linear-gradient(135deg, #00E676 0%, #00C853 100%);
        border-color: #00C853;
    }}
    .toggle-circle {{
        width: 16px;
        height: 16px;
        border-radius: 50%;
        border: 2px solid rgba(255, 255, 255, 0.4);
        background: transparent;
        transition: all 0.3s ease;
    }}
    .toggle-option.active .toggle-circle {{
        background: #ffffff;
        border-color: #ffffff;
    }}
    .toggle-option:hover:not(.active) {{
        background: rgba(55, 55, 55, 0.95);
        color: rgba(255, 255, 255, 0.7);
    }}
    .active-badge {{
        text-align: center;
        padding: 15px;
        margin-bottom: 25px;
    }}
    .badge {{
        background: linear-gradient(135deg, #2196F3 0%, #1976D2 100%);
        padding: 10px 25px;
        border-radius: 20px;
        color: white;
        font-weight: 600;
        font-size: 14px;
        display: inline-block;
    }}
    .badge.rapid {{
        background: linear-gradient(135deg, #00E676 0%, #00C853 100%);
    }}
    </style>

    <div class="toggle-switch-container">
        <div class="toggle-option left {'active external' if is_external else ''}" id="toggle-external">
            <div class="toggle-circle"></div>
            <span>🌐 External Platforms (LinkedIn, Naukri, FoundIt)</span>
        </div>
        <div class="toggle-option right {'active rapid' if not is_external else ''}" id="toggle-rapid">
            <div class="toggle-circle"></div>
            <span>⚡ RapidAPI Jobs (India Only)</span>
        </div>
    </div>

    <div class="active-badge">
        <span class="badge {'rapid' if not is_external else ''}">
            {'🌐 External Platforms Mode Active' if is_external else '⚡ RapidAPI Jobs Mode Active'}
        </span>
    </div>
    """

    st.markdown(toggle_html, unsafe_allow_html=True)

    # Create clickable buttons (hidden but functional)
    col_btn1, col_btn2 = st.columns(2)

    with col_btn1:
        if st.button("Switch to External Platforms", key="btn_external"):
            st.session_state.search_mode = "External Platforms"
            st.rerun()

    with col_btn2:
        if st.button("Switch to RapidAPI Jobs", key="btn_rapid"):
            st.session_state.search_mode = "RapidAPI Jobs"
            st.rerun()

    search_mode = st.session_state.search_mode

    if search_mode == "External Platforms":
        # External Platforms Section
        col1, col2 = st.columns(2)

        with col1:
            job_role = st.text_input("💼 Job Title / Skills", placeholder="e.g., Data Scientist", key="external_role")
            experience_level = st.selectbox(
                "📈 Experience Level",
                ["", "Internship", "Entry Level", "Associate", "Mid-Senior Level", "Director", "Executive"],
                key="external_exp"
            )

        with col2:
            location = st.text_input("📍 Location", placeholder="e.g., Bangalore, India", key="external_loc")
            job_type = st.selectbox(
                "📋 Job Type",
                ["", "Full-time", "Part-time", "Contract", "Temporary", "Volunteer", "Internship"],
                key="external_type"
            )

        foundit_experience = st.text_input("🔢 FoundIt Experience (Years)", placeholder="e.g., 1", key="external_foundit")

        search_clicked = st.button("🔎 Search External Jobs", key="search_external")

        if search_clicked:
            if job_role.strip() and location.strip():
                # Call search_jobs function for external platforms
                results = search_jobs(job_role, location, experience_level, job_type, foundit_experience)

                # Save search results if user is logged in
                if hasattr(st.session_state, 'username') and st.session_state.username:
                    # Convert results to format expected by save_job_search
                    formatted_results = []
                    for result in results:
                        platform_name = result["title"].split(":")[0]
                        formatted_results.append({
                            "platform": platform_name,
                            "apply_link": result["link"]
                        })
                    save_job_search(st.session_state.username, job_role, location, formatted_results)

                st.markdown("## 🎯 External Job Search Results")

                for job in results:
                    platform = job["title"].split(":")[0].lower()

                    # Platform styling
                    if "linkedin" in platform:
                        platform_name = "LinkedIn"
                        btn_color = "#0e76a8"
                        platform_gradient = "linear-gradient(135deg, #0e76a8 0%, #1a8cc8 100%)"
                    elif "naukri" in platform:
                        platform_name = "Naukri"
                        btn_color = "#ff5722"
                        platform_gradient = "linear-gradient(135deg, #ff5722 0%, #ff7043 100%)"
                    elif "foundit" in platform:
                        platform_name = "FoundIt (Monster)"
                        btn_color = "#7c4dff"
                        platform_gradient = "linear-gradient(135deg, #7c4dff 0%, #9c64ff 100%)"
                    else:
                        platform_name = platform.title()
                        btn_color = "#00c4cc"
                        platform_gradient = "linear-gradient(135deg, #00c4cc 0%, #26d0ce 100%)"

                    # Render card using reusable function
                    job_card_html, card_height = render_job_card(
                        title=job_role,
                        link=job['link'],
                        platform_name=platform_name,
                        brand_color=btn_color,
                        platform_gradient=platform_gradient,
                        location=location,
                        description="Open this platform to view full details."
                    )
                    st.components.v1.html(job_card_html, height=card_height, scrolling=False)
            else:
                st.warning("⚠️ Please enter both the Job Title and Location to perform the search.")

    else:
        # RapidAPI Jobs Section
        col1, col2 = st.columns(2)

        with col1:
            rapid_job_role = st.text_input("💼 Job Title / Skills", placeholder="e.g., Python Developer", key="rapid_role")

        with col2:
            rapid_location = st.text_input("📍 Location", placeholder="e.g., Mumbai", key="rapid_loc")

        # Number of results
        num_results = st.slider("📊 Number of Jobs to Fetch", min_value=5, max_value=50, value=10, step=5, key="rapid_num_results")

        # Advanced Filters
        with st.expander("🔧 Advanced Filters"):
            date_posted = st.selectbox(
                "📅 Date Posted",
                ["all", "today", "3days", "week", "month"],
                key="rapid_date"
            )
            rapid_job_type = st.selectbox(
                "📋 Job Type",
                ["", "Full-time", "Part-time", "Contract", "Internship"],
                key="rapid_type"
            )
            remote_only = st.checkbox("🏠 Remote Only", key="rapid_remote")
            radius = st.number_input("📏 Radius (km)", min_value=0, max_value=200, value=50, key="rapid_radius")
            job_requirements = st.multiselect(
                "📝 Job Requirements",
                ["under_3_years_experience", "more_than_3_years_experience", "no_experience", "no_degree"],
                key="rapid_req"
            )

        search_rapid_clicked = st.button("🔎 Search Rapid Jobs", key="search_rapid")

        if search_rapid_clicked:
            if rapid_job_role.strip() and rapid_location.strip():
                # Call fetch_live_jobs with parameters
                results = fetch_live_jobs(
                    rapid_job_role,
                    rapid_location,
                    job_type=rapid_job_type if rapid_job_type else None,
                    remote_only=remote_only,
                    results=num_results
                )

                # Save search results if user is logged in
                if hasattr(st.session_state, 'username') and st.session_state.username:
                    formatted_results = []
                    for job in results:
                        formatted_results.append({
                            "platform": "RapidAPI (Live)",
                            "apply_link": job.get("job_apply_link", "#")
                        })
                    save_job_search(st.session_state.username, rapid_job_role, rapid_location, formatted_results)

                st.markdown("## 🎯 RapidAPI Job Results")

                if results:
                    for job in results:
                        # Clean all job fields
                        job_title = clean_html(job.get("job_title", "N/A"))
                        job_company = clean_html(job.get("employer_name", "Unknown"))
                        job_location = f"{job.get('job_city','')}, {job.get('job_country','')}"
                        job_salary = f"{job.get('job_min_salary','None')} - {job.get('job_max_salary','None')} {job.get('job_salary_currency','')}"
                        job_type = job.get("job_employment_type", "N/A")
                        job_mode = "Remote" if job.get("job_is_remote") else "On-site"
                        job_publisher = clean_html(job.get("job_publisher", "N/A"))
                        job_description = clean_html(job.get("job_description", ""))[:250] + "..."

                        # Format date
                        formatted_date = "N/A"
                        if job.get("job_posted_at_datetime_utc") and job["job_posted_at_datetime_utc"] != "N/A":
                            try:
                                date_obj = datetime.datetime.fromisoformat(job["job_posted_at_datetime_utc"].replace('Z', '+00:00'))
                                formatted_date = date_obj.strftime("%b %d, %Y")
                            except:
                                formatted_date = job["job_posted_at_datetime_utc"]

                        # Colors
                        btn_color = "#00ff88"
                        platform_gradient = "linear-gradient(135deg, #00ff88 0%, #00cc6f 100%)"

                        # Custom HTML card
                        job_card_html = f"""
<div class="job-result-card" style="
    background: linear-gradient(135deg, #1e1e1e 0%, #2d2d2d 100%);
    padding: 25px;
    border-radius: 20px;
    margin-bottom: 25px;
    border-left: 6px solid {btn_color};
    box-shadow: 0 8px 32px rgba(0,0,0,0.3), 0 0 20px {btn_color}40;
    position: relative;
    overflow: hidden;
    transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
">
    <div class="shimmer-overlay"></div>

    <!-- Platform Badge -->
    <div style="font-size: 18px; margin-bottom: 15px; color: {btn_color}; font-weight: bold;">
        ⚡ RapidAPI (Live)
    </div>

    <!-- Job Title -->
    <div style="color: #ffffff; font-size: 22px; margin-bottom: 10px; font-weight: 600; line-height: 1.4;">
        {job_title}
    </div>

    <!-- Company -->
    <div style="color: #aaaaaa; font-size: 16px; margin-bottom: 15px;">
        🏢 <b>{job_company}</b>
    </div>

    <!-- Job Details Grid -->
    <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 10px; margin-bottom: 15px;">
        <div style="color: #cccccc; font-size: 14px;">📍 <b>Location:</b> {job_location}</div>
        <div style="color: #cccccc; font-size: 14px;">💰 <b>Salary:</b> {job_salary}</div>
        <div style="color: #cccccc; font-size: 14px;">📋 <b>Type:</b> {job_type}</div>
        <div style="color: #cccccc; font-size: 14px;">🌍 <b>Mode:</b> {job_mode}</div>
        <div style="color: #cccccc; font-size: 14px;">📅 <b>Posted:</b> {formatted_date}</div>
        <div style="color: #cccccc; font-size: 14px;">📰 <b>Source:</b> {job_publisher}</div>
    </div>

    <!-- Description -->
    <div style="color: #999999; font-size: 14px; margin-bottom: 20px; line-height: 1.6;">
        {job_description}
    </div>

    <!-- Apply Button -->
    <a href="{job.get('job_apply_link', '#')}" target="_blank" style="text-decoration: none;">
        <button class="job-button" style="
            background: {platform_gradient};
            color: white;
            padding: 12px 24px;
            border: none;
            border-radius: 12px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            box-shadow: 0 4px 15px {btn_color}50;
            transition: all 0.3s ease;
        ">
            🚀 Apply Now →
        </button>
    </a>
</div>
"""
                        
                        st.components.v1.html(job_card_html, height=450, scrolling=False)


                else:
                    st.info("No jobs found. Try adjusting your search criteria.")
            else:
                st.warning("⚠️ Please enter both the Job Title and Location to perform the search.")

    # Display saved job searches if user is logged in
    if hasattr(st.session_state, 'username') and st.session_state.username:
        # Get available platforms for filtering
        available_platforms = get_available_platforms(st.session_state.username)
        platform_options = ["All"] + available_platforms

        # Get total count of searches
        total_searches = get_total_saved_searches_count(st.session_state.username)

        st.markdown("### 📌 Your Saved Job Searches")

        if total_searches > 0:
            # Controls for filtering and pagination
            col1, col2 = st.columns([2, 1])

            with col1:
                platform_filter = st.selectbox(
                    "🔍 Filter by Platform",
                    platform_options,
                    key="platform_filter"
                )

            with col2:
                # Calculate pagination
                searches_per_page = 5
                filtered_count = get_total_saved_searches_count(st.session_state.username, platform_filter)
                max_pages = max(1, (filtered_count + searches_per_page - 1) // searches_per_page)

                if max_pages > 1:
                    current_page = st.slider(
                        "📄 Page",
                        min_value=1,
                        max_value=max_pages,
                        value=1,
                        key="page_slider"
                    )
                else:
                    current_page = 1

            # Calculate offset for pagination
            offset = (current_page - 1) * searches_per_page

            # Get filtered and paginated results
            saved_searches = get_saved_job_searches(
                st.session_state.username,
                limit=searches_per_page,
                offset=offset,
                platform_filter=platform_filter
            )

            if saved_searches:
                # Calculate and display search count info
                start_index = offset + 1
                end_index = min(offset + len(saved_searches), filtered_count)

                if platform_filter != "All":
                    st.markdown(f"**Showing {start_index}-{end_index} of {filtered_count} searches for {platform_filter}**")
                else:
                    st.markdown(f"**Showing {start_index}-{end_index} of {filtered_count} searches**")

                for search in saved_searches:
                    # Format timestamp - Convert UTC to IST
                    timestamp = datetime.datetime.strptime(search["timestamp"], "%Y-%m-%d %H:%M:%S.%f")
                    # Assume stored timestamp is in UTC, convert to IST
                    timestamp_utc = timestamp.replace(tzinfo=ZoneInfo('UTC'))
                    timestamp_ist = timestamp_utc.astimezone(ZoneInfo('Asia/Kolkata'))
                    formatted_time = timestamp_ist.strftime("%b %d, %Y at %I:%M %p IST")

                    # Platform styling
                    platform_lower = search["platform"].lower()
                    if "rapidapi" in platform_lower or "live" in platform_lower:
                        platform_color = "#00ff88"
                        platform_icon = "⚡"
                    elif platform_lower == "linkedin":
                        platform_color = "#0e76a8"
                        platform_icon = "🔵"
                    elif platform_lower == "naukri":
                        platform_color = "#ff5722"
                        platform_icon = "🏢"
                    elif "foundit" in platform_lower:
                        platform_color = "#7c4dff"
                        platform_icon = "🌐"
                    else:
                        platform_color = "#00c4cc"
                        platform_icon = "📄"

                    # Create columns for the card content and delete button
                    card_col, delete_col = st.columns([10, 1])

                    with card_col:
                        st.markdown(f"""
<div class="job-result-card" style="
    background: linear-gradient(135deg, #1a1a1a 0%, #2a2a2a 100%);
    padding: 20px;
    border-radius: 15px;
    margin-bottom: 15px;
    border-left: 4px solid {platform_color};
    box-shadow: 0 4px 16px rgba(0,0,0,0.2);
    position: relative;
    overflow: hidden;
">
    <div style="display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 15px;">
        <div>
            <div style="color: #ffffff; font-size: 16px; font-weight: 600; margin-bottom: 5px;">
                {platform_icon} {search['role']} in {search['location']}
            </div>
            <div style="color: {platform_color}; font-size: 14px; font-weight: 500;">
                {search['platform']}
            </div>
        </div>
        <div style="color: #888; font-size: 12px; text-align: right;">
            {formatted_time}
        </div>
    </div>
    <a href="{search['url']}" target="_blank" style="text-decoration: none;">
        <button class="job-button" style="
            background: linear-gradient(135deg, {platform_color} 0%, {platform_color}dd 100%);
            color: white;
            padding: 8px 16px;
            border: none;
            border-radius: 8px;
            font-size: 14px;
            font-weight: 500;
            cursor: pointer;
            transition: all 0.3s ease;
        ">
            🔗 View Jobs →
        </button>
    </a>
</div>
""", unsafe_allow_html=True)

                    with delete_col:
                        # Delete button
                        if st.button("🗑", key=f"delete_{search['id']}", help="Delete this search"):
                            delete_saved_job_search(search['id'])
                            st.rerun()
            else:
                # No results for the current filter
                st.markdown(f"""
<div style="
    background: linear-gradient(135deg, #1a1a1a 0%, #2a2a2a 100%);
    padding: 20px;
    border-radius: 15px;
    text-align: center;
    color: #888;
    border: 2px dashed #444;
">
    <div style="font-size: 24px; margin-bottom: 10px;">🔍</div>
    <div>No saved searches found for {platform_filter if platform_filter != 'All' else 'this page'}.</div>
</div>
""", unsafe_allow_html=True)
        else:
            # No saved searches at all
            st.markdown("""
<div style="
    background: linear-gradient(135deg, #1a1a1a 0%, #2a2a2a 100%);
    padding: 20px;
    border-radius: 15px;
    text-align: center;
    color: #888;
    border: 2px dashed #444;
">
    <div style="font-size: 24px; margin-bottom: 10px;">📭</div>
    <div>No saved job searches yet. Start searching to see your history here!</div>
</div>
""", unsafe_allow_html=True)

    # Enhanced CSS with advanced animations and effects
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

    /* Global Enhancements */
    .stApp {
        font-family: 'Inter', sans-serif;
    }

    /* Advanced Glow Animation */
    @keyframes glow {
        0% {
            box-shadow: 0 0 5px rgba(255,255,255,0.1), 0 0 10px rgba(0,255,255,0.1), 0 0 15px rgba(0,255,255,0.1);
        }
        50% {
            box-shadow: 0 0 10px rgba(255,255,255,0.2), 0 0 20px rgba(0,255,255,0.4), 0 0 30px rgba(0,255,255,0.3);
        }
        100% {
            box-shadow: 0 0 5px rgba(255,255,255,0.1), 0 0 10px rgba(0,255,255,0.1), 0 0 15px rgba(0,255,255,0.1);
        }
    }

    /* Shimmer Effect */
    @keyframes shimmer {
        0% {
            transform: translateX(-100%);
        }
        100% {
            transform: translateX(100%);
        }
    }

    .shimmer-overlay {
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.05), transparent);
        transform: translateX(-100%);
        animation: shimmer 3s infinite;
        z-index: 1;
    }

    /* Floating Animation */
    @keyframes float {
        0%, 100% {
            transform: translateY(0px);
        }
        50% {
            transform: translateY(-5px);
        }
    }

    /* Pulse Animation */
    @keyframes pulse {
        0%, 100% {
            transform: scale(1);
        }
        50% {
            transform: scale(1.02);
        }
    }

    /* Enhanced Company Cards */
    .company-card {
        background: linear-gradient(135deg, #1e1e1e 0%, #2d2d2d 100%);
        color: #ffffff;
        border-radius: 20px;
        padding: 25px;
        margin-bottom: 25px;
        box-shadow: 0 8px 32px rgba(0,0,0,0.3);
        transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        cursor: pointer;
        text-decoration: none;
        display: block;
        animation: glow 4s infinite alternate, float 6s ease-in-out infinite;
        position: relative;
        overflow: hidden;
        border: 1px solid rgba(255,255,255,0.1);
    }

    .company-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(135deg, rgba(0,255,255,0.1) 0%, rgba(255,0,255,0.1) 100%);
        opacity: 0;
        transition: opacity 0.3s ease;
        z-index: 1;
    }

    .company-card:hover::before {
        opacity: 1;
    }

    .company-card:hover {
        transform: translateY(-8px) scale(1.02);
        box-shadow: 0 20px 40px rgba(0,0,0,0.4), 0 0 30px rgba(0, 255, 255, 0.3);
        border-color: rgba(0,255,255,0.5);
    }

    /* Job Result Cards */
    .job-result-card:hover {
        transform: translateY(-5px) scale(1.01);
        box-shadow: 0 15px 40px rgba(0,0,0,0.4) !important;
    }

    /* Enhanced Buttons */
    .job-button::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent);
        transition: left 0.5s;
        z-index: 1;
    }

    .job-button:hover::before {
        left: 100%;
    }

    .job-button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.3);
    }

    /* Enhanced Pills */
    .pill {
        display: inline-block;
        background: linear-gradient(135deg, #333 0%, #444 100%);
        padding: 8px 16px;
        border-radius: 25px;
        margin: 6px 8px 0 0;
        font-size: 13px;
        font-weight: 500;
        border: 1px solid rgba(255,255,255,0.1);
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }

    .pill::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(135deg, rgba(0,255,255,0.2) 0%, rgba(255,0,255,0.2) 100%);
        opacity: 0;
        transition: opacity 0.3s ease;
    }

    .pill:hover::before {
        opacity: 1;
    }

    .pill:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,255,255,0.3);
    }

    /* Enhanced Title Headers */
    .title-header {
        color: #ffffff;
        font-size: 28px;
        margin-top: 50px;
        margin-bottom: 30px;
        font-weight: 700;
        text-align: center;
        background: linear-gradient(135deg, #00c4cc 0%, #7c4dff 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        position: relative;
        animation: pulse 3s infinite;
    }

    .title-header::after {
        content: '';
        position: absolute;
        bottom: -10px;
        left: 50%;
        transform: translateX(-50%);
        width: 60px;
        height: 3px;
        background: linear-gradient(135deg, #00c4cc 0%, #7c4dff 100%);
        border-radius: 2px;
    }

    /* Company Logo Enhancement */
    .company-logo {
        font-size: 28px;
        margin-right: 12px;
        filter: drop-shadow(0 0 8px rgba(255,255,255,0.3));
        animation: float 4s ease-in-out infinite;
    }

    .company-header {
        font-size: 24px;
        font-weight: 700;
        display: flex;
        align-items: center;
        margin-bottom: 15px;
        position: relative;
        z-index: 2;
    }

    /* Responsive Enhancements */
    @media (max-width: 768px) {
        .company-card, .job-result-card {
            padding: 20px;
            margin-bottom: 20px;
        }

        .title-header {
            font-size: 24px;
        }

        .company-header {
            font-size: 20px;
        }
    }

    /* Scrollbar Styling */
    ::-webkit-scrollbar {
        width: 8px;
    }

    ::-webkit-scrollbar-track {
        background: #1e1e1e;
    }

    ::-webkit-scrollbar-thumb {
        background: linear-gradient(135deg, #00c4cc 0%, #7c4dff 100%);
        border-radius: 4px;
    }

    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(135deg, #26d0ce 0%, #9c64ff 100%);
    }
    </style>
    """, unsafe_allow_html=True)

    # ---------- Company Lookup by Domain ----------


    # ---------- Featured Companies ----------
    st.markdown("### <div class='title-header'>🏢 Featured Companies</div>", unsafe_allow_html=True)

    selected_category = st.selectbox("📂 Browse Featured Companies By Category", ["All", "tech", "indian_tech", "global_corps"])
    companies_to_show = get_featured_companies() if selected_category == "All" else get_featured_companies(selected_category)

    for company in companies_to_show:
        category_tags = ''.join([f"<span class='pill'>{cat}</span>" for cat in company['categories']])
        st.markdown(f"""
        <a href="{company['careers_url']}" class="company-card" target="_blank">
            <div class="company-header">
                <span class="company-logo">{company.get('emoji', '🏢')}</span>
                {company['name']}
            </div>
            <p style="margin-bottom: 15px; line-height: 1.6; position: relative; z-index: 2;">{company['description']}</p>
            <div style="position: relative; z-index: 2;">{category_tags}</div>
        </a>
        """, unsafe_allow_html=True)

    # ---------- Market Insights ----------
    st.markdown("### <div class='title-header'>📈 Job Market Trends</div>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### <div style='color: #00c4cc; font-size: 20px; font-weight: 600; margin-bottom: 20px;'>🚀 Trending Skills</div>", unsafe_allow_html=True)
        for skill in JOB_MARKET_INSIGHTS["trending_skills"]:
            st.markdown(f"""
            <div class="company-card">
                <h4 style="color: #00c4cc; margin-bottom: 10px; position: relative; z-index: 2;">🔧 {skill['name']}</h4>
                <p style="position: relative; z-index: 2;">📈 Growth Rate: <span style="color: #4ade80; font-weight: 600;">{skill['growth']}</span></p>
            </div>
            """, unsafe_allow_html=True)

    with col2:
        st.markdown("#### <div style='color: #7c4dff; font-size: 20px; font-weight: 600; margin-bottom: 20px;'>🌍 Top Job Locations</div>", unsafe_allow_html=True)
        for loc in JOB_MARKET_INSIGHTS["top_locations"]:
            st.markdown(f"""
            <div class="company-card">
                <h4 style="color: #7c4dff; margin-bottom: 10px; position: relative; z-index: 2;">📍 {loc['name']}</h4>
                <p style="position: relative; z-index: 2;">💼 Openings: <span style="color: #fbbf24; font-weight: 600;">{loc['jobs']}</span></p>
            </div>
            """, unsafe_allow_html=True)

    # ---------- Salary Insights ----------
    st.markdown("### <div class='title-header'>💰 Salary Insights</div>", unsafe_allow_html=True)
    for role in JOB_MARKET_INSIGHTS["salary_insights"]:
        st.markdown(f"""
        <div class="company-card">
            <h4 style="color: #10b981; margin-bottom: 10px; position: relative; z-index: 2;">💼 {role['role']}</h4>
            <p style="margin-bottom: 8px; position: relative; z-index: 2;">📅 Experience: <span style="color: #60a5fa; font-weight: 500;">{role['experience']}</span></p>
            <p style="position: relative; z-index: 2;">💵 Salary Range: <span style="color: #34d399; font-weight: 600;">{role['range']}</span></p>
        </div>
        """, unsafe_allow_html=True)

def evaluate_interview_answer(answer: str, question: str = None):
    """
    Uses an LLM to strictly evaluate an interview answer.
    Returns (score out of 5, feedback string).
    """
    from llm_manager import call_llm
    import re
    import streamlit as st

    # Empty check
    if not answer.strip():
        return 0, "⚠️ No answer provided."

    # 🔹 LLM Prompt (STRICTER)
    prompt = f"""
    You are an expert technical interview evaluator.

    ### Task:
    Evaluate the candidate's answer to the question below.
    Be STRICT. Only give high scores if the answer is technically correct, relevant, and detailed.

    ### Question:
    {question if question else "N/A"}

    ### Candidate Answer:
    {answer}

    ### Strict Scoring Rubric:
    - 5 = Exceptional: Fully correct, highly relevant, clear, detailed, technically accurate.
    - 4 = Good: Mostly correct and relevant, but missing some depth/clarity.
    - 3 = Average: Partially correct OR generic, but somewhat relevant.
    - 2 = Weak: Mostly irrelevant, shallow, or major gaps in correctness.
    - 1 = Poor: Completely irrelevant, incoherent, or very wrong.
    - 0 = No answer / total nonsense.

    ### Output Format:
    Score: <number between 0 and 5>
    Feedback: <constructive feedback in 1–2 sentences>
    """

    try:
        # Call LLM
        response = call_llm(prompt, session=st.session_state).strip()

        # Extract Score
        score_match = re.search(r"Score:\s*(\d+)", response)
        score = int(score_match.group(1)) if score_match else 1  # stricter fallback

        # Extract Feedback
        feedback_match = re.search(r"Feedback:\s*(.+)", response)
        feedback = feedback_match.group(1).strip() if feedback_match else "Answer was unclear or irrelevant."

        # ✅ Keep score in 0–5 range
        score = max(0, min(score, 5))

    except Exception as e:
        score = 1
        feedback = f"⚠️ Evaluation fallback due to error: {e}"

    return score, feedback


def evaluate_interview_answer_for_scores(answer: str, question: str, difficulty: str, role: str = "", domain: str = ""):
    """
    UPGRADED: Intelligent evaluation with chain-of-thought reasoning and structured feedback.
    Uses JSON-based parsing for robustness and provides detailed, actionable feedback.

    Returns dict with keys: knowledge, communication, relevance, feedback (list), followup

    Features:
    - Chain-of-thought evaluation: extracts key concepts, identifies strengths/gaps
    - Structured feedback: detailed paragraph with specific, actionable insights
    - Difficulty calibration: Easy (encouraging), Medium (balanced), Hard (strict)
    - JSON-based parsing for reliability
    """
    from llm_manager import call_llm
    import json
    import streamlit as st

    # Empty check or junk answers
    if not answer.strip() or answer == "⚠️ No Answer" or len(answer.strip()) < 3:
        return {
            "knowledge": 0,
            "communication": 0,
            "relevance": 0,
            "feedback": "No answer provided. Try using the STAR method: Situation, Task, Action, Result. Provide specific examples from your experience to demonstrate your understanding and capabilities.",
            "followup": ""
        }

    # Check for obvious junk answers (single character, just symbols, etc.)
    if len(answer.strip()) == 1 or not any(c.isalnum() for c in answer):
        return {
            "knowledge": 0,
            "communication": 0,
            "relevance": 0,
            "feedback": "Answer appears incomplete or invalid. Please provide a meaningful response with technical details and structure your answer clearly with concrete examples from your experience.",
            "followup": ""
        }

    # STRICTER JUNK FILTERING: Check word count and meaningful tokens
    words = answer.strip().split()
    meaningful_words = [w for w in words if len(w) > 2 and any(c.isalpha() for c in w)]

    if len(words) < 5 or len(meaningful_words) < 2:
        return {
            "knowledge": 0,
            "communication": 0,
            "relevance": 0,
            "feedback": "Answer too short or lacks substance. Provide a detailed response with at least 3-4 sentences and include specific examples or technical details to demonstrate your understanding.",
            "followup": ""
        }

    # Difficulty-based evaluation guidance
    difficulty_guidance = {
        "Easy": {
            "tone": "encouraging and forgiving",
            "expectations": "basic understanding and general concepts",
            "scoring": "Give partial credit for effort. Score 5-10 for reasonable attempts, 3-4 for weak but present answers, 0-2 for irrelevant/junk.",
            "feedback_style": "positive and encouraging with gentle improvement tips"
        },
        "Medium": {
            "tone": "balanced and realistic",
            "expectations": "scenario-based thinking, some technical depth, and practical examples",
            "scoring": "Score 6-10 for good answers, 3-5 for incomplete/basic answers, 0-2 for poor/irrelevant.",
            "feedback_style": "constructive and specific with clear improvement areas"
        },
        "Hard": {
            "tone": "strict and technical",
            "expectations": "deep technical knowledge, system design thinking, edge cases, and comprehensive understanding",
            "scoring": "Score 7-10 for excellent answers, 4-6 for adequate but incomplete, 0-3 for weak/incorrect.",
            "feedback_style": "concise and technical with precise critique"
        }
    }

    guidance = difficulty_guidance.get(difficulty, difficulty_guidance["Medium"])

    # Build context for relevance checking
    context_info = f" for {role} in {domain}" if role and domain else ""

    # UPGRADED CHAIN-OF-THOUGHT EVALUATION PROMPT
    prompt = f"""You are an expert technical interviewer evaluating a candidate's answer{context_info}.

QUESTION: {question}
CANDIDATE'S ANSWER: {answer}
DIFFICULTY LEVEL: {difficulty}

EVALUATION APPROACH ({guidance['tone']}):
Expected: {guidance['expectations']}
Scoring: {guidance['scoring']}
Feedback Style: {guidance['feedback_style']}

STEP-BY-STEP EVALUATION PROCESS:

STEP 1 - EXTRACT KEY CONCEPTS FROM THE QUESTION:
List 3-5 technical concepts, keywords, or expected topics the question is asking about.

STEP 2 - ANALYZE THE ANSWER:
✅ STRENGTHS: What did the candidate do well? Which concepts did they cover? What was clear or correct?
⚠️ GAPS/IMPROVEMENTS: What's missing? What's incorrect? What could be clearer?

STEP 3 - SCORE THE ANSWER (1-10 scale):
- Knowledge: Technical correctness, depth, and completeness (did they cover key concepts?)
- Communication: Clarity, structure, and articulation (was it easy to follow?)
- Relevance: How directly does this answer the question? Is it on-topic?

STEP 4 - GENERATE DETAILED FEEDBACK (2-4 comprehensive paragraphs):
Provide detailed, flowing feedback that covers:
- Specific strengths: What they did well, which concepts they covered correctly, and what was clear
- Specific gaps or areas for improvement: What key concepts or details they missed, what could be more accurate
- Actionable recommendations: Concrete suggestions for improvement with examples
- Overall assessment: A brief summary of their understanding level

Write feedback as natural, flowing paragraphs (not bullet points). Make it detailed, specific to their answer, and constructive.

{"STEP 5 - FOLLOW-UP QUESTION: Generate ONE probing follow-up question that digs deeper based on their answer." if difficulty == "Hard" else ""}

OUTPUT FORMAT (strict JSON):
{{
  "key_concepts": ["concept1", "concept2", "concept3"],
  "strengths": ["strength1", "strength2"],
  "gaps": ["gap1", "gap2"],
  "knowledge": <number 1-10>,
  "communication": <number 1-10>,
  "relevance": <number 1-10>,
  "feedback": "Detailed, comprehensive feedback in 2-4 flowing paragraphs. Be specific about what the candidate did well, what they missed, and how they can improve. Reference actual content from their answer. Make it constructive, actionable, and personalized."{',\n  "followup": "One probing follow-up question"' if difficulty == "Hard" else ''}
}}

IMPORTANT RULES:
- If answer is off-topic or from wrong domain, set relevance to 0-2
- If answer is junk/minimal, set all scores to 0-2
- Feedback must be specific to THIS answer, not generic templates
- Reference actual content from the candidate's answer in feedback
- Each feedback point should feel personalized and human

Provide ONLY the JSON output, no additional text."""

    try:
        response = call_llm(prompt, session=st.session_state).strip()

        # Clean response - remove markdown code blocks if present
        if response.startswith("```"):
            response = response.split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()

        # Parse JSON response
        result = json.loads(response)

        # Extract and validate scores
        knowledge = int(result.get("knowledge", 1))
        communication = int(result.get("communication", 1))
        relevance = int(result.get("relevance", 1))

        # Clamp scores to 0-10 range
        knowledge = max(0, min(10, knowledge))
        communication = max(0, min(10, communication))
        relevance = max(0, min(10, relevance))

        # Extract feedback (should be a detailed string, not a list)
        feedback = result.get("feedback", "")

        # If feedback comes as a list (fallback), join it into paragraphs
        if isinstance(feedback, list):
            feedback = "\n\n".join(feedback)

        # Ensure we have substantial feedback
        if not feedback or len(feedback.strip()) < 50:
            feedback = "Your answer shows some understanding, but could benefit from more technical depth and specific examples. Consider structuring your response more clearly and providing concrete details from your experience. Focus on addressing all aspects of the question comprehensively."

        # Extract follow-up question
        followup = result.get("followup", "") if difficulty == "Hard" else ""

        return {
            "knowledge": knowledge,
            "communication": communication,
            "relevance": relevance,
            "feedback": feedback,  # Now a string, not a list
            "followup": followup
        }

    except json.JSONDecodeError as e:
        # Fallback: try to extract scores from non-JSON response
        import re
        try:
            knowledge = int(re.search(r'"?knowledge"?\s*:\s*(\d+)', response, re.IGNORECASE).group(1))
            communication = int(re.search(r'"?communication"?\s*:\s*(\d+)', response, re.IGNORECASE).group(1))
            relevance = int(re.search(r'"?relevance"?\s*:\s*(\d+)', response, re.IGNORECASE).group(1))

            # Extract feedback (try both string and array format)
            feedback_match = re.search(r'"feedback"\s*:\s*"([^"]+)"', response, re.DOTALL)
            if feedback_match:
                feedback = feedback_match.group(1)
            else:
                # Fallback: try array format and join
                feedback_array_match = re.search(r'"feedback"\s*:\s*\[(.*?)\]', response, re.DOTALL)
                if feedback_array_match:
                    feedback_text = feedback_array_match.group(1)
                    feedback_items = [f.strip(' "\'') for f in re.findall(r'"([^"]+)"', feedback_text)]
                    feedback = "\n\n".join(feedback_items) if feedback_items else "Answer evaluated but formatting unclear. Provide more structured responses with clear examples and explanations."
                else:
                    feedback = "Answer evaluated but formatting unclear. Provide more structured responses with clear examples and explanations."

            return {
                "knowledge": max(0, min(10, knowledge)),
                "communication": max(0, min(10, communication)),
                "relevance": max(0, min(10, relevance)),
                "feedback": feedback if isinstance(feedback, str) else "\n\n".join(feedback[:5]),
                "followup": ""
            }
        except:
            pass

    except Exception as e:
        pass

    # Final fallback based on difficulty
    fallback_scores = {"Easy": 3, "Medium": 2, "Hard": 1}
    fallback_score = fallback_scores.get(difficulty, 2)

    return {
        "knowledge": fallback_score,
        "communication": fallback_score,
        "relevance": fallback_score,
        "feedback": "Unable to evaluate properly. Please provide a clear, structured answer. Use the STAR method for behavioral questions and include technical details and examples for technical questions.",
        "followup": ""
    }


def get_ist_time():
    """Get current time in IST timezone"""
    try:
        from datetime import datetime
        import pytz
        ist = pytz.timezone('Asia/Kolkata')
        return datetime.now(ist).strftime('%Y-%m-%d %H:%M:%S')
    except:
        from datetime import datetime
        return datetime.now().strftime('%Y-%m-%d %H:%M:%S')


def log_user_action(username: str, action: str):
    """Log user actions - placeholder for compatibility"""
    pass


def create_interview_database():
    """Create interview_results table if not exists"""
    import sqlite3
    try:
        conn = sqlite3.connect('resume_data.db')
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS interview_results (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL,
                role TEXT,
                domain TEXT,
                avg_score REAL,
                total_questions INTEGER,
                completed_on TEXT,
                feedback_summary TEXT
            )
        """)
        conn.commit()
        conn.close()
    except Exception as e:
        import streamlit as st
        st.error(f"Database error: {e}")


def save_interview_result(username: str, role: str, domain: str, avg_score: float, total_questions: int, feedback_summary: str):
    """Save interview result to database"""
    import sqlite3
    try:
        conn = sqlite3.connect('resume_data.db')
        cursor = conn.cursor()
        completed_on = get_ist_time()
        cursor.execute("""
            INSERT INTO interview_results (username, role, domain, avg_score, total_questions, completed_on, feedback_summary)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (username, role, domain, avg_score, total_questions, completed_on, feedback_summary))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        import streamlit as st
        st.error(f"Failed to save interview result: {e}")
        return False


def format_feedback_text(feedback):
    """
    Format feedback text into bullet points for clean display
    """
    import re
    import html
    sentences = re.split(r'(?<=\.)\s+', feedback.strip())
    sentences = [s.strip() for s in sentences if len(s.strip()) > 0]
    formatted = "<b>💡 Improvement Tips:</b><br><ul style='margin-top:5px;'>"
    for s in sentences:
        # Escape HTML special characters to display tags like <header>, <section>, etc.
        safe_sentence = html.escape(s)
        formatted += f"<li>{safe_sentence}</li>"
    formatted += "</ul>"
    return formatted


def generate_interview_pdf_report(username, role, domain, completed_on, questions, answers, scores, feedbacks, overall_avg, badge, difficulty="Medium"):
    """
    Generate PDF report for interview using xhtml2pdf

    FIXED: Now shows full answers (up to 2000 chars) instead of truncating at 500
    FIXED: Added follow-up questions for Hard difficulty interviews
    """
    try:
        from xhtml2pdf import pisa
        from io import BytesIO

        # Build XHTML content
        xhtml = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #00c3ff; text-align: center; }}
                h2 {{ color: #0099cc; margin-top: 20px; }}
                .header {{ background: #f0f0f0; padding: 15px; border-radius: 8px; margin-bottom: 20px; }}
                .question-block {{ margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 8px; page-break-inside: avoid; }}
                .score {{ font-weight: bold; color: #00c3ff; }}
                .feedback {{ color: #666; margin-top: 10px; padding: 10px; background: #f9f9f9; border-left: 3px solid #00c3ff; }}
                .feedback ul {{ margin: 5px 0 0 0; padding-left: 20px; }}
                .feedback li {{ margin: 8px 0; line-height: 1.5; }}
                .summary {{ background: #fffacd; padding: 15px; border-radius: 8px; margin: 20px 0; }}
                .answer-text {{ white-space: pre-wrap; word-wrap: break-word; margin: 10px 0; }}
            </style>
        </head>
        <body>
            <h1>Interview Practice Report</h1>
            <div class="header">
                <p><strong>Candidate:</strong> {username}</p>
                <p><strong>Role:</strong> {role}</p>
                <p><strong>Domain:</strong> {domain}</p>
                <p><strong>Date:</strong> {completed_on}</p>
            </div>
            <div class="summary">
                <h2>Overall Performance</h2>
                <p class="score">Average Score: {overall_avg:.1f}/10</p>
                <p><strong>Badge Earned:</strong> {badge}</p>
            </div>
            <h2>Detailed Q&A Review</h2>
        """

        # CRITICAL FIX: Add each question/answer/score/feedback with FULL answer (no truncation)
        for i, (q, a, score_dict, f) in enumerate(zip(questions, answers, scores, feedbacks), 1):
            # Ensure score_dict is a dictionary
            if isinstance(score_dict, dict):
                avg_q_score = (score_dict.get('knowledge', 5) + score_dict.get('communication', 5) + score_dict.get('relevance', 5)) / 3
            else:
                # Fallback if score_dict is not a dict
                avg_q_score = 5.0
                score_dict = {'knowledge': 5, 'communication': 5, 'relevance': 5}

            # Escape HTML special characters to prevent rendering issues
            q_escaped = q.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            a_escaped = a.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            # Handle feedback as string (convert list to paragraphs if needed)
            if isinstance(f, list):
                f_text = "\n\n".join(f)
            else:
                f_text = str(f)

            # Format feedback into bullet points
            import re
            sentences = re.split(r'(?<=\.)\s+', f_text.strip())
            sentences = [sent.strip() for sent in sentences if len(sent.strip()) > 0]
            bullet_feedback = "<b>💡 Improvement Tips:</b><ul>"
            for sent in sentences:
                sent_escaped = sent.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                bullet_feedback += f"<li>{sent_escaped}</li>"
            bullet_feedback += "</ul>"
            f_escaped = bullet_feedback

            # SHOW FULL ANSWER - NO TRUNCATION IN PDF
            answer_display = a_escaped

            # Get follow-up question if exists (for Hard difficulty)
            followup_text = ""
            if difficulty == "Hard" and isinstance(score_dict, dict) and score_dict.get('followup'):
                followup_escaped = score_dict['followup'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                followup_text = f"""<div style="margin-top: 10px; padding: 10px; background: #fff3cd; border-radius: 5px;">
                    <strong>Follow-up Question (for Hard interviews):</strong><br/>
                    {followup_escaped}
                </div>"""

            xhtml += f"""
            <div class="question-block">
                <h3>Question {i}</h3>
                <p><strong>Q:</strong> {q_escaped}</p>
                <div class="answer-text"><strong>Your Answer:</strong><br/>{answer_display}</div>
                <p class="score">Knowledge: {score_dict.get('knowledge', 0)}/10 | Communication: {score_dict.get('communication', 0)}/10 | Relevance: {score_dict.get('relevance', 0)}/10</p>
                <p class="score">Question Score: {avg_q_score:.1f}/10</p>
                <div class="feedback">{f_escaped}</div>
                {followup_text}
            </div>
            """

        xhtml += """
        </body>
        </html>
        """

        # Convert to PDF
        pdf_out = BytesIO()
        pisa_status = pisa.CreatePDF(xhtml, dest=pdf_out)
        pdf_out.seek(0)

        if pisa_status.err:
            return None

        return pdf_out.getvalue()

    except Exception as e:
        import streamlit as st
        st.error(f"PDF generation failed: {e}")
        return None


import streamlit as st
import plotly.graph_objects as go
from courses import COURSES_BY_CATEGORY, RESUME_VIDEOS, INTERVIEW_VIDEOS, get_courses_for_role
from llm_manager import call_llm
import time
import threading

with tab4:
    # Inject CSS styles (keeping existing styles)
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
        
        * {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
        }

        .header-box {
            background: linear-gradient(135deg, #0a0e27 0%, #1a1f3a 25%, #2d3561 50%, #3f4787 75%, #5158ae 100%);
            border: 2px solid transparent;
            background-clip: padding-box;
            position: relative;
            padding: 25px;
            border-radius: 20px;
            text-align: center;
            margin-bottom: 35px;
            box-shadow: 
                0 8px 32px rgba(0, 195, 255, 0.15),
                0 4px 16px rgba(0, 195, 255, 0.1),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
            backdrop-filter: blur(10px);
            overflow: hidden;
        }

        .header-box::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(45deg, #00c3ff, #0066cc, #00c3ff, #0066cc);
            background-size: 400% 400%;
            animation: gradientShift 8s ease infinite;
            z-index: -1;
            border-radius: 20px;
            padding: 2px;
            mask: linear-gradient(#fff 0 0) content-box, linear-gradient(#fff 0 0);
            mask-composite: exclude;
        }

        @keyframes gradientShift {
            0%, 100% { background-position: 0% 50%; }
            50% { background-position: 100% 50%; }
        }

        .header-box h2 {
            font-size: 32px;
            color: #ffffff;
            margin: 0;
            font-weight: 700;
            text-shadow: 
                0 0 20px rgba(0, 195, 255, 0.5),
                0 2px 4px rgba(0, 0, 0, 0.3);
            letter-spacing: -0.5px;
        }

        .glow-header {
            font-size: 24px;
            text-align: center;
            color: #00c3ff;
            text-shadow: 
                0 0 20px rgba(0, 195, 255, 0.8),
                0 0 40px rgba(0, 195, 255, 0.4);
            margin: 20px 0 15px 0;
            font-weight: 600;
            letter-spacing: -0.3px;
            animation: pulse 3s ease-in-out infinite;
        }

        @keyframes pulse {
            0%, 100% { opacity: 1; transform: scale(1); }
            50% { opacity: 0.9; transform: scale(1.02); }
        }

        .stRadio > div {
            flex-direction: row !important;
            justify-content: center !important;
            gap: 16px;
            flex-wrap: wrap;
        }

        .stRadio label {
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            border: 2px solid #00c3ff;
            color: #00c3ff;
            padding: 14px 24px;
            margin: 6px;
            border-radius: 12px;
            cursor: pointer;
            transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            font-weight: 500;
            min-width: 190px;
            text-align: center;
            position: relative;
            overflow: hidden;
            box-shadow: 
                0 4px 15px rgba(0, 195, 255, 0.1),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
        }

        .stRadio label::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(0, 195, 255, 0.2), transparent);
            transition: left 0.5s;
        }

        .stRadio label:hover {
            background: linear-gradient(135deg, #00c3ff15 0%, #00c3ff25 100%);
            transform: translateY(-2px);
            box-shadow: 
                0 8px 25px rgba(0, 195, 255, 0.2),
                inset 0 1px 0 rgba(255, 255, 255, 0.2);
        }

        .stRadio label:hover::before {
            left: 100%;
        }

        .stRadio input:checked + div > label {
            background: linear-gradient(135deg, #00c3ff 0%, #0099cc 100%);
            color: #000000;
            font-weight: 600;
            transform: scale(1.05);
            box-shadow: 
                0 8px 30px rgba(0, 195, 255, 0.4),
                inset 0 1px 0 rgba(255, 255, 255, 0.3);
        }

        .card {
            background: linear-gradient(135deg, #0f1419 0%, #1a2332 25%, #253447 50%, #30455c 75%, #3b5671 100%);
            border: 2px solid transparent;
            border-radius: 16px;
            padding: 20px 25px;
            margin: 12px 0;
            position: relative;
            overflow: hidden;
            transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            box-shadow: 
                0 4px 20px rgba(0, 195, 255, 0.1),
                inset 0 1px 0 rgba(255, 255, 255, 0.05);
        }

        .card::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(45deg, #00c3ff, #0066cc);
            z-index: -1;
            border-radius: 16px;
            padding: 2px;
            mask: linear-gradient(#fff 0 0) content-box, linear-gradient(#fff 0 0);
            mask-composite: exclude;
            opacity: 0.8;
        }

        .card::after {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.1), transparent);
            transition: left 0.6s;
        }

        .card:hover {
            transform: translateY(-4px) scale(1.02);
            box-shadow: 
                0 12px 40px rgba(0, 195, 255, 0.25),
                0 8px 20px rgba(0, 195, 255, 0.15),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
        }

        .card:hover::after {
            left: 100%;
        }

        .card a {
            color: #00c3ff;
            font-weight: 600;
            font-size: 16px;
            text-decoration: none;
            display: flex;
            align-items: center;
            gap: 8px;
            transition: all 0.3s ease;
            text-shadow: 0 0 10px rgba(0, 195, 255, 0.3);
        }

        .card a:hover {
            color: #ffffff;
            text-decoration: none;
            text-shadow: 
                0 0 15px rgba(255, 255, 255, 0.5),
                0 0 30px rgba(0, 195, 255, 0.3);
            transform: translateX(4px);
        }

        /* Enhanced selectbox styling */
        .stSelectbox > div > div {
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            border: 2px solid #00c3ff;
            border-radius: 10px;
            color: #00c3ff;
        }

        .stSelectbox > div > div:hover {
            box-shadow: 0 0 15px rgba(0, 195, 255, 0.3);
        }

        /* Enhanced subheader styling */
        .stApp h3 {
            color: #00c3ff;
            text-shadow: 0 0 10px rgba(0, 195, 255, 0.5);
            font-weight: 600;
            margin-bottom: 20px;
        }

        /* Learning path container */
        .learning-path-container {
            text-align: center;
            margin: 30px 0 20px 0;
            padding: 15px;
            background: linear-gradient(135deg, rgba(0, 195, 255, 0.05) 0%, rgba(0, 195, 255, 0.1) 100%);
            border-radius: 12px;
            border: 1px solid rgba(0, 195, 255, 0.2);
        }

        .learning-path-text {
            color: #00c3ff;
            font-weight: 600;
            font-size: 20px;
            text-shadow: 0 0 15px rgba(0, 195, 255, 0.6);
            letter-spacing: -0.3px;
        }

        /* Video container enhancements */
        .stVideo {
            border-radius: 12px;
            overflow: hidden;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
            transition: transform 0.3s ease;
        }

        .stVideo:hover {
            transform: scale(1.02);
        }

        /* Info message styling */
        .stAlert {
            background: linear-gradient(135deg, rgba(0, 195, 255, 0.1) 0%, rgba(0, 195, 255, 0.05) 100%);
            border: 1px solid rgba(0, 195, 255, 0.3);
            border-radius: 10px;
        }

        /* New styles for quiz and interview sections */
        .quiz-card {
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            border: 2px solid #00c3ff;
            border-radius: 15px;
            padding: 20px;
            margin: 15px 0;
            box-shadow: 0 4px 20px rgba(0, 195, 255, 0.15);
        }

        .badge-container {
            text-align: center;
            padding: 30px;
            background: linear-gradient(135deg, rgba(0, 195, 255, 0.12) 0%, rgba(0, 195, 255, 0.06) 100%);
            backdrop-filter: blur(10px);
            -webkit-backdrop-filter: blur(10px);
            border-radius: 16px;
            border: 1px solid rgba(0, 195, 255, 0.25);
            margin: 20px 0;
            box-shadow: 0 8px 32px rgba(0, 0, 0, 0.2), inset 0 1px 0 rgba(255, 255, 255, 0.08);
        }

        .score-display {
            font-size: 64px;
            font-weight: bold;
            color: #00d4ff;
            text-shadow: 0 0 30px rgba(0, 212, 255, 0.6);
            letter-spacing: 2px;
        }

        .role-selector {
            background: linear-gradient(135deg, rgba(0, 195, 255, 0.05) 0%, rgba(0, 195, 255, 0.1) 100%);
            border: 1px solid rgba(0, 195, 255, 0.2);
            border-radius: 12px;
            padding: 20px;
            margin: 15px 0;
        }

        /* Course tile styling */
        .course-tile {
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            border: 2px solid #00c3ff;
            border-radius: 15px;
            padding: 20px;
            margin: 15px 0;
            transition: all 0.3s ease;
            position: relative;
            overflow: hidden;
        }

        .course-tile:hover {
            transform: translateY(-5px);
            box-shadow: 0 10px 30px rgba(0, 195, 255, 0.3);
        }

        .course-title {
            color: #00c3ff;
            font-size: 18px;
            font-weight: 600;
            margin-bottom: 10px;
        }

        .course-description {
            color: #ffffff;
            font-size: 14px;
            margin-bottom: 15px;
            line-height: 1.4;
        }

        .difficulty-badge {
            display: inline-block;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: 500;
            margin-bottom: 15px;
        }

        .difficulty-beginner {
            background: linear-gradient(135deg, #4CAF50, #45a049);
            color: white;
        }

        .difficulty-intermediate {
            background: linear-gradient(135deg, #FF9800, #f57c00);
            color: white;
        }

        .difficulty-advanced {
            background: linear-gradient(135deg, #f44336, #d32f2f);
            color: white;
        }

        .course-link-btn {
            background: linear-gradient(135deg, #00c3ff, #0099cc);
            color: white;
            padding: 8px 16px;
            border-radius: 8px;
            text-decoration: none;
            font-weight: 500;
            display: inline-block;
            transition: all 0.3s ease;
        }

        .course-link-btn:hover {
            background: linear-gradient(135deg, #0099cc, #007acc);
            transform: translateX(2px);
            text-decoration: none;
            color: white;
        }

        /* Radar chart container */
        .radar-container {
            background: linear-gradient(135deg, rgba(0, 195, 255, 0.05) 0%, rgba(0, 195, 255, 0.1) 100%);
            border: 1px solid rgba(0, 195, 255, 0.2);
            border-radius: 15px;
            padding: 20px;
            margin: 20px 0;
        }

        /* Timer styling */
        .timer-container {
            background: linear-gradient(135deg, rgba(255, 193, 7, 0.1) 0%, rgba(255, 193, 7, 0.05) 100%);
            border: 1px solid rgba(255, 193, 7, 0.3);
            border-radius: 12px;
            padding: 15px;
            margin: 15px 0;
            text-align: center;
        }

        .timer-display {
            font-size: 24px;
            font-weight: bold;
            color: #ffd700;
            text-shadow: 0 0 10px rgba(255, 215, 0, 0.5);
        }

        .timer-urgent {
            color: #ff4444;
            text-shadow: 0 0 15px rgba(255, 68, 68, 0.8);
            animation: pulse 1s ease-in-out infinite;
        }
        </style>
    """, unsafe_allow_html=True)

    # Header (keeping existing)
    st.markdown("""
        <div class="header-box">
            <h2>📚 Recommended Learning Hub</h2>
        </div>
    """, unsafe_allow_html=True)

    # Subheader (keeping existing)
    st.markdown('<div class="glow-header">🎓 Explore Career Resources</div>', unsafe_allow_html=True)
    st.markdown("<p style='text-align:center; color:#ccc; font-size: 16px; margin-bottom: 25px;'>Curated courses and videos for your career growth, resume tips, and interview success.</p>", unsafe_allow_html=True)

    # Learning path label (keeping existing)
    st.markdown("""
        <div class="learning-path-container">
            <span class="learning-path-text">
                🧭 Choose Your Learning Path
            </span>
        </div>
    """, unsafe_allow_html=True)

    # Updated Radio buttons with new options
    st.markdown("""
        <div style="display: flex; justify-content: center; width: 100%;">
            <div style="display: flex; justify-content: center; gap: 16px;">
    """, unsafe_allow_html=True)

    # Check if page changed away from AI Interview Coach - stop interview if so
    previous_page = st.session_state.get('previous_page_selection', None)

    page = st.radio(
        label="Select Learning Option",
        options=["Courses by Role", "Resume Videos", "Interview Videos",  "AI Interview Coach 🤖"],
        horizontal=True,
        key="page_selection",
        label_visibility="collapsed"
    )

    # STOP INTERVIEW ON TAB CHANGE
    if previous_page == "AI Interview Coach 🤖" and page != "AI Interview Coach 🤖":
        # User switched away from AI Interview Coach - reset interview state
        if st.session_state.get('dynamic_interview_started', False) and not st.session_state.get('dynamic_interview_completed', False):
            st.session_state.dynamic_interview_started = False
            st.session_state.dynamic_interview_completed = True

    # Update previous page for next comparison
    st.session_state.previous_page_selection = page

    st.markdown("</div></div>", unsafe_allow_html=True)

    # NEW: Index-based difficulty function (replaces keyword-based)
    def get_course_difficulty_by_index(index):
        if index == 0:
            return "Beginner"
        elif index in [1, 2]:
            return "Intermediate"
        else:
            return "Advanced"

    # Helper functions for dynamic question generation
    def generate_career_quiz_questions(domain, role):
        """Generate role-specific career quiz questions"""
        questions = []
        
        # Role-specific question templates
        role_templates = {
            "Software Development and Engineering": {
                "Frontend Developer": [
                    {
                        "question": "Which aspect of web development excites you most?",
                        "options": [
                            "Creating beautiful, interactive user interfaces",
                            "Building responsive designs that work on all devices", 
                            "Optimizing website performance and accessibility",
                            "Working with modern JavaScript frameworks"
                        ]
                    },
                    {
                        "question": "What's your preferred approach to styling?",
                        "options": [
                            "Writing custom CSS from scratch",
                            "Using CSS frameworks like Bootstrap or Tailwind",
                            "CSS-in-JS solutions for component-based styling", 
                            "CSS preprocessors like Sass or Less"
                        ]
                    },
                    {
                        "question": "Which tools do you enjoy working with most?",
                        "options": [
                            "React, Vue, or Angular for building SPAs",
                            "HTML5, CSS3, and vanilla JavaScript",
                            "Design tools like Figma or Adobe XD",
                            "Build tools like Webpack, Vite, or Parcel"
                        ]
                    }
                ],
                "Backend Developer": [
                    {
                        "question": "What backend architecture interests you most?",
                        "options": [
                            "RESTful API design and implementation",
                            "Microservices architecture and distributed systems",
                            "Database design and optimization",
                            "Server-side security and authentication"
                        ]
                    },
                    {
                        "question": "Which programming paradigm do you prefer?",
                        "options": [
                            "Object-oriented programming with Java/.NET",
                            "Functional programming with languages like Scala",
                            "Dynamic languages like Python or JavaScript",
                            "Systems programming with Go or Rust"
                        ]
                    },
                    {
                        "question": "What type of backend challenges excite you?",
                        "options": [
                            "Scaling applications to handle millions of users",
                            "Integrating complex third-party services",
                            "Optimizing database queries and performance",
                            "Building robust error handling and monitoring"
                        ]
                    }
                ],
                "Full Stack Developer": [
                    {
                        "question": "What full-stack aspect appeals to you most?",
                        "options": [
                            "Building end-to-end features from UI to database",
                            "Managing the entire application development lifecycle",
                            "Working with both frontend and backend technologies",
                            "Understanding how all system components interact"
                        ]
                    },
                    {
                        "question": "Which tech stack interests you most?",
                        "options": [
                            "MERN (MongoDB, Express, React, Node.js)",
                            "MEAN (MongoDB, Express, Angular, Node.js)",
                            "Django + React/Vue for Python development",
                            "Ruby on Rails with modern frontend frameworks"
                        ]
                    }
                ],
                "Mobile App Developer": [
                    {
                        "question": "What type of mobile development interests you?",
                        "options": [
                            "Native iOS development with Swift",
                            "Native Android development with Kotlin/Java",
                            "Cross-platform development with React Native",
                            "Hybrid app development with Flutter"
                        ]
                    },
                    {
                        "question": "Which mobile development aspect excites you most?",
                        "options": [
                            "Creating intuitive mobile user experiences",
                            "Integrating with device hardware and sensors",
                            "Optimizing app performance and battery usage",
                            "Publishing apps to App Store and Google Play"
                        ]
                    }
                ],
                "Game Developer": [
                    {
                        "question": "What type of game development interests you?",
                        "options": [
                            "3D game development with Unity or Unreal Engine",
                            "2D indie game development and pixel art",
                            "Mobile gaming and casual game mechanics",
                            "VR/AR game development and immersive experiences"
                        ]
                    },
                    {
                        "question": "Which game development aspect excites you most?",
                        "options": [
                            "Game design and player experience",
                            "Graphics programming and visual effects",
                            "Game physics and realistic simulations",
                            "Multiplayer networking and real-time systems"
                        ]
                    }
                ]
            },
            "Data Science and Analytics": {
                "Data Scientist": [
                    {
                        "question": "Which data science task excites you most?",
                        "options": [
                            "Building predictive models and machine learning algorithms",
                            "Exploring large datasets to discover hidden patterns",
                            "Creating data visualizations and storytelling with data",
                            "Designing experiments and A/B testing strategies"
                        ]
                    },
                    {
                        "question": "What's your preferred approach to data analysis?",
                        "options": [
                            "Statistical modeling and hypothesis testing",
                            "Deep learning and neural networks",
                            "Feature engineering and data preprocessing",
                            "Time series analysis and forecasting"
                        ]
                    },
                    {
                        "question": "Which tools do you enjoy working with most?",
                        "options": [
                            "Python with pandas, scikit-learn, and TensorFlow",
                            "R for statistical computing and analysis",
                            "SQL for database querying and data manipulation",
                            "Jupyter notebooks for exploratory data analysis"
                        ]
                    }
                ],
                "Data Analyst": [
                    {
                        "question": "Which type of analysis interests you most?",
                        "options": [
                            "Business intelligence and performance dashboards",
                            "Customer behavior analysis and segmentation",
                            "Financial analysis and risk assessment",
                            "Market research and competitive analysis"
                        ]
                    },
                    {
                        "question": "What's your preferred way to present insights?",
                        "options": [
                            "Interactive dashboards with Tableau or Power BI",
                            "Statistical reports with clear recommendations",
                            "Data visualizations and infographics",
                            "Executive summaries and business presentations"
                        ]
                    }
                ],
                "Machine Learning Engineer": [
                    {
                        "question": "Which ML engineering task excites you most?",
                        "options": [
                            "Deploying models to production at scale",
                            "Building ML pipelines and automation systems",
                            "Optimizing model performance and efficiency",
                            "Implementing MLOps and model monitoring"
                        ]
                    },
                    {
                        "question": "What type of ML problems interest you?",
                        "options": [
                            "Computer vision and image processing",
                            "Natural language processing and text analysis",
                            "Recommendation systems and personalization",
                            "Reinforcement learning and autonomous systems"
                        ]
                    }
                ]
            },
            "Cloud Computing and DevOps": {
                "Cloud Architect": [
                    {
                        "question": "Which cloud architecture aspect interests you most?",
                        "options": [
                            "Designing scalable, fault-tolerant systems",
                            "Multi-cloud and hybrid cloud strategies",
                            "Cloud security and compliance frameworks",
                            "Cost optimization and resource management"
                        ]
                    },
                    {
                        "question": "What type of cloud solutions excite you?",
                        "options": [
                            "Serverless architectures and event-driven systems",
                            "Container orchestration with Kubernetes",
                            "Data lakes and analytics platforms",
                            "AI/ML platforms and managed services"
                        ]
                    }
                ],
                "DevOps Engineer": [
                    {
                        "question": "Which DevOps practice interests you most?",
                        "options": [
                            "Building CI/CD pipelines and automation",
                            "Infrastructure as Code with Terraform/CloudFormation",
                            "Container orchestration and microservices",
                            "Monitoring, logging, and observability"
                        ]
                    },
                    {
                        "question": "What type of automation excites you?",
                        "options": [
                            "Deployment automation and release management",
                            "Infrastructure provisioning and configuration",
                            "Testing automation and quality gates",
                            "Incident response and self-healing systems"
                        ]
                    }
                ],
                "Site Reliability Engineer": [
                    {
                        "question": "Which SRE responsibility interests you most?",
                        "options": [
                            "Maintaining system reliability and uptime",
                            "Performance optimization and capacity planning",
                            "Incident management and post-mortem analysis",
                            "Service level objectives and error budgets"
                        ]
                    },
                    {
                        "question": "What aspect of system reliability excites you?",
                        "options": [
                            "Building robust monitoring and alerting systems",
                            "Designing disaster recovery and backup strategies",
                            "Automating operational tasks and runbooks",
                            "Analyzing system performance and bottlenecks"
                        ]
                    }
                ]
            },
            "Cybersecurity": {
                "Security Analyst": [
                    {
                        "question": "Which security area interests you most?",
                        "options": [
                            "Threat detection and incident response",
                            "Vulnerability assessment and risk management",
                            "Security monitoring and SIEM analysis",
                            "Compliance and security policy development"
                        ]
                    },
                    {
                        "question": "What type of security challenges excite you?",
                        "options": [
                            "Investigating security breaches and forensics",
                            "Analyzing malware and attack patterns",
                            "Network security and firewall management",
                            "Identity and access management systems"
                        ]
                    }
                ],
                "Penetration Tester": [
                    {
                        "question": "Which penetration testing approach interests you?",
                        "options": [
                            "Web application security testing",
                            "Network penetration testing and infrastructure",
                            "Social engineering and phishing simulations",
                            "Mobile application security testing"
                        ]
                    },
                    {
                        "question": "What aspect of ethical hacking excites you?",
                        "options": [
                            "Finding vulnerabilities before malicious actors",
                            "Using creative techniques to bypass security",
                            "Helping organizations improve their defenses",
                            "Staying updated on latest attack methods"
                        ]
                    }
                ]
            },
            "UI/UX Design": {
                "UI Designer": [
                    {
                        "question": "Which UI design aspect interests you most?",
                        "options": [
                            "Creating visually stunning interface designs",
                            "Designing consistent design systems and components",
                            "Working with typography, colors, and visual hierarchy",
                            "Prototyping interactions and micro-animations"
                        ]
                    },
                    {
                        "question": "What type of design work excites you?",
                        "options": [
                            "Mobile app interface design",
                            "Web application and dashboard design",
                            "Icon design and visual asset creation",
                            "Brand identity and visual design systems"
                        ]
                    }
                ],
                "UX Designer": [
                    {
                        "question": "Which UX design activity interests you most?",
                        "options": [
                            "User research and persona development",
                            "Information architecture and user flows",
                            "Wireframing and prototype development",
                            "Usability testing and design validation"
                        ]
                    },
                    {
                        "question": "What aspect of user experience excites you?",
                        "options": [
                            "Solving complex user problems with simple solutions",
                            "Understanding user behavior and psychology",
                            "Designing accessible and inclusive experiences",
                            "Measuring and optimizing user engagement"
                        ]
                    }
                ]
            },
            "Project Management": {
                "Project Manager": [
                    {
                        "question": "Which project management aspect interests you most?",
                        "options": [
                            "Planning and scheduling project timelines",
                            "Coordinating teams and stakeholder communication",
                            "Risk management and problem-solving",
                            "Budget management and resource allocation"
                        ]
                    },
                    {
                        "question": "What type of projects excite you?",
                        "options": [
                            "Large-scale software development projects",
                            "Cross-functional digital transformation initiatives",
                            "Product launches and go-to-market strategies",
                            "Process improvement and organizational change"
                        ]
                    }
                ],
                "Product Manager": [
                    {
                        "question": "Which product management activity interests you most?",
                        "options": [
                            "Product strategy and roadmap development",
                            "User research and market analysis",
                            "Feature prioritization and requirement gathering",
                            "Go-to-market strategy and product launches"
                        ]
                    },
                    {
                        "question": "What aspect of product development excites you?",
                        "options": [
                            "Identifying user needs and pain points",
                            "Defining product vision and strategy",
                            "Working with engineering and design teams",
                            "Analyzing product metrics and user feedback"
                        ]
                    }
                ]
            }
        }

        # Get role-specific questions or generate generic ones
        if domain in role_templates and role in role_templates[domain]:
            questions = role_templates[domain][role]
        else:
            # Generate generic questions based on role name
            questions = [
                {
                    "question": f"How interested are you in pursuing a career as a {role}?",
                    "options": [
                        "Very interested - it's my dream job",
                        "Somewhat interested - I want to learn more",
                        "Moderately interested - it seems challenging",
                        "Not very interested - but I'm curious"
                    ]
                },
                {
                    "question": f"What attracts you most about the {role} role?",
                    "options": [
                        "The technical challenges and problem-solving",
                        "The creative aspects and innovation opportunities", 
                        "The career growth potential and salary",
                        "The impact on users and business outcomes"
                    ]
                }
            ]
        
        return questions

    # Helper function to generate fallback questions
    def self_generate_fallback_questions(role, domain, difficulty, count):
        """Generate fallback questions when LLM doesn't return enough"""
        if difficulty == "Easy":
            base_questions = [
                f"What interests you most about the {role} position?",
                f"Describe your basic understanding of {role} responsibilities.",
                f"What are the fundamental skills needed for {role}?",
                f"How do you stay updated with trends in {domain}?",
                f"Why do you want to work as a {role}?",
                f"What do you know about the {role} role?",
                f"Tell me about yourself and your interest in {role}.",
                f"What motivates you to pursue a career in {domain}?",
                f"Describe a project you've worked on related to {role}.",
                f"What are your career goals as a {role}?"
            ]
        elif difficulty == "Hard":
            base_questions = [
                f"Design a scalable system architecture for a {role} project handling millions of users.",
                f"Explain the trade-offs between different approaches in {domain} and when to use each.",
                f"How would you troubleshoot a critical production issue in a {role} context?",
                f"Describe your approach to mentoring junior team members as a {role}.",
                f"What are the biggest technical challenges facing {role} professionals today?",
                f"How would you architect a distributed system for {domain}?",
                f"Explain how you would optimize performance in a complex {role} project.",
                f"What advanced techniques do you use in {domain}?",
                f"Describe a time you made a critical technical decision as a {role}.",
                f"How do you approach system design for high availability in {domain}?"
            ]
        else:  # Medium
            base_questions = [
                f"Describe a challenging project you've worked on relevant to {role}.",
                f"How do you approach problem-solving in {domain}?",
                f"What tools and technologies are you most comfortable with for {role}?",
                f"Tell me about a time you had to learn a new skill for {role}.",
                f"How do you prioritize tasks when working as a {role}?",
                f"Describe your experience with {domain} technologies.",
                f"How do you handle tight deadlines as a {role}?",
                f"What's your approach to code quality in {domain}?",
                f"Tell me about a technical challenge you solved as a {role}.",
                f"How do you collaborate with team members in {domain}?"
            ]
        return base_questions[:count]

    # UPDATED: AI-Generated Questions using LLM with DIFFICULTY SUPPORT
    def generate_interview_questions_with_llm(domain, role, interview_type, num_questions, difficulty="Medium"):
        """
        Generate interview questions using LLM based on domain, role, type, and difficulty.

        FIXED: Now difficulty is passed into LLM prompt and affects question complexity
        """
        # Define difficulty-specific instructions
        difficulty_instructions = {
            "Easy": "Generate BASIC and INTRODUCTORY level questions. Focus on fundamental concepts, definitions, and simple scenarios. Questions should be suitable for entry-level candidates or those new to the field.",
            "Medium": "Generate SCENARIO-BASED and MODERATELY TECHNICAL questions. Include situational questions that require practical thinking and intermediate technical knowledge. Suitable for candidates with some experience.",
            "Hard": "Generate DEEP TECHNICAL, SYSTEM DESIGN, and COMPLEX PROBLEM-SOLVING questions. Include architecture decisions, trade-offs, scalability concerns, and advanced concepts. Suitable for senior-level candidates."
        }

        prompt = f"""You are an expert interviewer.

Generate EXACTLY {num_questions} unique {interview_type} interview questions
for the role of {role} in {domain}.

DIFFICULTY LEVEL: {difficulty}
{difficulty_instructions.get(difficulty, difficulty_instructions["Medium"])}

CRITICAL REQUIREMENTS:
- Generate EXACTLY {num_questions} questions - no more, no less
- Keep each question concise (1-2 sentences max)
- Avoid duplicates
- Match the difficulty level specified above
- Output ONLY the questions, one per line
- DO NOT add numbering, bullet points, or any prefixes
- DO NOT add any introductory text or explanations

Output format example:
What is your experience with cloud technologies?
How would you handle a system outage?
Describe your approach to code reviews.

Generate exactly {num_questions} questions now:
"""

        try:
            response = call_llm(prompt, session=st.session_state)

            # Split by newlines and clean up
            raw_questions = [q.strip() for q in response.split('\n') if q.strip()]

            # Remove any numbering or bullet points more aggressively
            import re
            cleaned_questions = []
            for q in raw_questions:
                # Remove various prefixes: "1. ", "1) ", "- ", "• ", "* ", "Question 1:", etc.
                clean_q = re.sub(r'^[\d\)\.\-•\*]+\s*', '', q).strip()
                clean_q = re.sub(r'^Question\s*\d*\s*:?\s*', '', clean_q, flags=re.IGNORECASE).strip()

                # Only add if it's a meaningful question
                if clean_q and len(clean_q) > 15 and not clean_q.lower().startswith('generate') and not clean_q.lower().startswith('here'):
                    cleaned_questions.append(clean_q)

                # Stop if we have enough questions
                if len(cleaned_questions) >= num_questions:
                    break

            # If we got fewer questions than requested, try to pad with fallback
            if len(cleaned_questions) < num_questions:
                st.warning(f"Only generated {len(cleaned_questions)} questions, padding with fallback questions...")
                # Add fallback questions to meet the requirement
                fallback_needed = num_questions - len(cleaned_questions)
                fallback_qs = self_generate_fallback_questions(role, domain, difficulty, fallback_needed)
                cleaned_questions.extend(fallback_qs)

            # EXACT QUESTION COUNT: Enforce exact count
            cleaned_questions = cleaned_questions[:num_questions]
            return cleaned_questions

        except Exception as e:
            st.error(f"Failed to generate questions with LLM: {e}")
            # Fallback to static questions appropriate for difficulty
            if difficulty == "Easy":
                fallback_questions = [
                    f"What interests you most about the {role} position?",
                    f"Describe your basic understanding of {role} responsibilities.",
                    f"What are the fundamental skills needed for {role}?",
                    f"How do you stay updated with trends in {domain}?",
                    f"Why do you want to work as a {role}?"
                ]
            elif difficulty == "Hard":
                fallback_questions = [
                    f"Design a scalable system architecture for a {role} project handling millions of users.",
                    f"Explain the trade-offs between different approaches in {domain} and when to use each.",
                    f"How would you troubleshoot a critical production issue in a {role} context?",
                    f"Describe your approach to mentoring junior team members as a {role}.",
                    f"What are the biggest technical challenges facing {role} professionals today?"
                ]
            else:  # Medium
                fallback_questions = [
                    f"Describe a challenging project you've worked on relevant to {role}.",
                    f"How do you approach problem-solving in {domain}?",
                    f"What tools and technologies are you most comfortable with for {role}?",
                    f"Tell me about a time you had to learn a new skill for {role}.",
                    f"How do you prioritize tasks when working as a {role}?"
                ]
            return fallback_questions[:num_questions]

    # Badge system for gamification
    BADGE_CONFIG = {
        "career_quiz": {
            "novice": {"min_score": 0, "max_score": 40, "emoji": "🌱", "title": "Career Explorer"},
            "intermediate": {"min_score": 41, "max_score": 70, "emoji": "📚", "title": "Career Seeker"},
            "advanced": {"min_score": 71, "max_score": 100, "emoji": "🎯", "title": "Career Champion"}
        },
        "interview": {
            "needs_practice": {"min_score": 1.0, "max_score": 2.5, "emoji": "💪", "title": "Keep Practicing"},
            "good": {"min_score": 2.6, "max_score": 3.5, "emoji": "👍", "title": "Good Performer"},
            "excellent": {"min_score": 3.6, "max_score": 4.5, "emoji": "🌟", "title": "Star Performer"},
            "interview_ready": {"min_score": 4.6, "max_score": 5.0, "emoji": "🏆", "title": "Interview Ready"}
        }
    }

    def get_badge_for_score(score_type, score):
        """Get badge based on score type and value"""
        badges = BADGE_CONFIG.get(score_type, {})
        for badge_name, config in badges.items():
            if config["min_score"] <= score <= config["max_score"]:
                return config["emoji"], config["title"]
        return "🎖️", "Participant"

    def create_skill_radar_chart(skills_data):
        """Create a radar chart for skills using Plotly"""
        # Extract skills and values
        skills = list(skills_data.keys())
        values = list(skills_data.values())
        
        # Create radar chart
        fig = go.Figure()
        
        fig.add_trace(go.Scatterpolar(
            r=values,
            theta=skills,
            fill='toself',
            name='Skills',
            line=dict(color='#00c3ff', width=2),
            fillcolor='rgba(0, 195, 255, 0.2)',
            hovertemplate='<b>%{theta}</b><br>Importance: %{r}/10<br><extra></extra>'
        ))
        
        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, 10],
                    tickfont=dict(color='white', size=10),
                    gridcolor='rgba(255, 255, 255, 0.2)'
                ),
                angularaxis=dict(
                    tickfont=dict(color='white', size=12),
                    gridcolor='rgba(255, 255, 255, 0.2)'
                ),
                bgcolor='rgba(0, 0, 0, 0)'
            ),
            showlegend=False,
            title=dict(
                text="Skills Importance Radar",
                x=0.5,
                font=dict(color='#00c3ff', size=16)
            ),
            paper_bgcolor='rgba(0, 0, 0, 0)',
            plot_bgcolor='rgba(0, 0, 0, 0)',
            font=dict(color='white'),
            height=400
        )
        
        return fig

    def get_course_description(course_title, role):
        """Generate a short description for the course"""
        descriptions = {
            'Frontend Developer': f"Master modern frontend development with {course_title.split()[0]} and build responsive web applications.",
            'Backend Developer': f"Learn server-side development and API design to become a skilled backend developer.",
            'Full Stack Developer': f"Comprehensive full-stack development course covering both frontend and backend technologies.",
            'Data Scientist': f"Dive deep into data science methodologies, machine learning, and statistical analysis.",
            'Machine Learning Engineer': f"Build and deploy machine learning models at scale with industry best practices.",
            'Cloud Architect': f"Design scalable cloud infrastructure and learn enterprise-grade cloud solutions.",
            'DevOps Engineer': f"Master CI/CD pipelines, containerization, and infrastructure automation.",
            'UI Designer': f"Create stunning user interfaces with modern design principles and tools.",
            'UX Designer': f"Learn user research, wireframing, and create exceptional user experiences."
        }
        
        return descriptions.get(role, f"Comprehensive course to advance your skills in {role} role.")

    def display_courses_by_difficulty(courses, role):
        """Display courses grouped by difficulty using index-based mapping"""
        # Group courses by difficulty
        difficulty_groups = {"Beginner": [], "Intermediate": [], "Advanced": []}
        
        for idx, (title, url) in enumerate(courses):
            difficulty = get_course_difficulty_by_index(idx)
            description = get_course_description(title, role)
            difficulty_groups[difficulty].append((title, url, description))
        
        # Display each difficulty group
        for difficulty in ["Beginner", "Intermediate", "Advanced"]:
            if difficulty_groups[difficulty]:
                st.markdown(f"### 🎯 {difficulty} Level")
                for title, url, description in difficulty_groups[difficulty]:
                    st.markdown(f"""
                        <div class="course-tile">
                            <div class="course-title">{title}</div>
                            <div class="course-description">{description}</div>
                            <span class="difficulty-badge difficulty-{difficulty.lower()}">{difficulty}</span>
                            <br>
                            <a href="{url}" target="_blank" class="course-link-btn">
                                🚀 Start Learning
                            </a>
                        </div>
                    """, unsafe_allow_html=True)

    # UPDATED SECTIONS

    # Section 1: UPDATED Courses by Role with Index-based Difficulty
    if page == "Courses by Role":
        st.subheader("🎯 Courses by Career Role")
        
        col1, col2 = st.columns(2)
        with col1:
            category = st.selectbox(
                "Select Career Category",
                options=list(COURSES_BY_CATEGORY.keys()),
                key="category_selection"
            )
        
        with col2:
            if category:
                roles = list(COURSES_BY_CATEGORY[category].keys())
                role = st.selectbox(
                    "Select Role / Job Title",
                    options=roles,
                    key="role_selection"
                )
            else:
                role = None
        
        if category and role:
            # UPDATED: Add difficulty filter
            difficulty_filter = st.selectbox(
                "Filter by Difficulty Level",
                options=["All Levels", "Beginner", "Intermediate", "Advanced"],
                key="difficulty_filter"
            )
            
            st.subheader(f"📘 Courses for **{role}** in **{category}**:")
            courses = get_courses_for_role(category, role)
            
            if courses:
                # UPDATED: Display courses using index-based difficulty
                filtered_courses = []
                for idx, (title, url) in enumerate(courses):
                    difficulty = get_course_difficulty_by_index(idx)
                    
                    # Apply difficulty filter
                    if difficulty_filter == "All Levels" or difficulty == difficulty_filter:
                        filtered_courses.append((title, url, difficulty, idx))
                
                if filtered_courses:
                    for title, url, difficulty, idx in filtered_courses:
                        description = get_course_description(title, role)
                        
                        # UPDATED: Interactive course tile with index-based difficulty
                        st.markdown(f"""
                            <div class="course-tile">
                                <div class="course-title">{title}</div>
                                <div class="course-description">{description}</div>
                                <span class="difficulty-badge difficulty-{difficulty.lower()}">{difficulty}</span>
                                <br>
                                <a href="{url}" target="_blank" class="course-link-btn">
                                    🚀 Start Learning
                                </a>
                            </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("🚫 No courses found for this difficulty level.")
            else:
                st.info("🚫 No courses found for this role.")
        
        # Show skill radar chart for selected role
        if category and role:
            st.markdown("---")
            st.markdown('<div class="radar-container">', unsafe_allow_html=True)
            st.subheader("🎯 Skills Radar Chart")
            
            # Generate sample skills data based on role
            role_skills = {
                # ==== Software Development & Engineering ====
                "Frontend Developer": {
                    "JavaScript": 9, "React/Vue": 8, "CSS/HTML": 9,
                    "Responsive Design": 8, "Performance Optimization": 7, "Testing": 6
                },
                "Backend Developer": {
                    "API Design": 9, "Database Management": 8, "Security": 8,
                    "Scalability": 7, "Cloud Services": 7, "Testing": 6
                },
                "Full Stack Developer": {
                    "Frontend": 8, "Backend": 8, "Databases": 7,
                    "API Integration": 8, "DevOps Basics": 6, "Testing": 7
                },
                "Mobile App Developer": {
                    "Flutter/React Native": 8, "Swift/Kotlin": 8, "UI/UX": 8,
                    "APIs": 7, "Performance Optimization": 7, "App Deployment": 7
                },
                "Game Developer": {
                    "Unity/Unreal": 9, "C# / C++": 8, "Game Physics": 7,
                    "Graphics/Rendering": 8, "AI in Games": 6, "Multiplayer Systems": 7
                },
                # ==== Data Science & Analytics ====
                "Data Scientist": {
                    "Python/R": 9, "Machine Learning": 8, "Statistics": 9,
                    "Data Visualization": 7, "SQL": 8, "Domain Knowledge": 6
                },
                "Data Analyst": {
                    "SQL": 9, "Excel/Spreadsheets": 8, "Visualization": 8,
                    "Statistics": 8, "Python/R": 7, "Business Acumen": 7
                },
                "Machine Learning Engineer": {
                    "ML Algorithms": 9, "Deep Learning": 8, "MLOps": 7,
                    "Data Engineering": 8, "Python/Frameworks": 9, "Cloud Deployment": 7
                },
                # ==== Cloud Computing & DevOps ====
                "Cloud Architect": {
                    "AWS/Azure/GCP": 9, "System Design": 8, "Networking": 7,
                    "Security": 8, "Scalability": 9, "Cost Optimization": 7
                },
                "DevOps Engineer": {
                    "CI/CD": 9, "Containerization": 8, "Cloud Platforms": 8,
                    "Monitoring": 7, "Infrastructure as Code": 8, "Security": 7
                },
                "Site Reliability Engineer": {
                    "Reliability Engineering": 9, "Monitoring": 8, "Automation": 8,
                    "Incident Response": 8, "System Design": 7, "Security": 7
                },
                # ==== Cybersecurity ====
                "Security Analyst": {
                    "Threat Detection": 9, "Incident Response": 8, "Networking": 7,
                    "SIEM Tools": 8, "Risk Management": 7, "Compliance": 6
                },
                "Penetration Tester": {
                    "Ethical Hacking": 9, "Web Security": 8, "Exploitation": 8,
                    "Scripting": 7, "Reporting": 6, "Network Security": 7
                },
                # ==== UI/UX Design ====
                "UI Designer": {
                    "Design Tools": 9, "Visual Design": 8, "Typography": 7,
                    "Color Theory": 8, "Prototyping": 7, "User Research": 6
                },
                "UX Designer": {
                    "User Research": 9, "Wireframing": 8, "Prototyping": 8,
                    "Usability Testing": 7, "Accessibility": 8, "Design Thinking": 7
                },
                # ==== Project Management ====
                "Project Manager": {
                    "Planning": 9, "Communication": 8, "Risk Management": 8,
                    "Leadership": 7, "Agile/Scrum": 8, "Budgeting": 7
                },
                "Product Manager": {
                    "Market Research": 9, "Product Strategy": 8, "Analytics": 8,
                    "Communication": 8, "Agile Methods": 7, "User-Centered Design": 7
                }
            }
            
            skills_data = role_skills.get(role, {
                "Technical Skills": 8, "Problem Solving": 7, "Communication": 6,
                "Leadership": 5, "Domain Knowledge": 7, "Continuous Learning": 8
            })
            
            # Create and display radar chart
            radar_fig = create_skill_radar_chart(skills_data)
            st.plotly_chart(radar_fig, use_container_width=True)
            
            # Add hover tooltip information
            st.markdown("""
                <div style="text-align: center; color: #00c3ff; margin-top: 10px;">
                    💡 Hover over the chart points to see skill importance ratings!
                </div>
            """, unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)

    # Section 2: Resume Videos (unchanged)
    elif page == "Resume Videos":
        st.subheader("📄 Resume Writing Videos")
        categories = list(RESUME_VIDEOS.keys())
        selected_cat = st.selectbox(
            "Select Resume Video Category",
            options=categories,
            key="resume_vid_cat"
        )
        if selected_cat:
            st.subheader(f"📂 {selected_cat}")
            videos = RESUME_VIDEOS[selected_cat]
            cols = st.columns(2)
            for idx, (title, url) in enumerate(videos):
                with cols[idx % 2]:
                    st.markdown(f"**{title}**")
                    st.video(url)

    # Section 3: Interview Videos (unchanged)
    elif page == "Interview Videos":
        st.subheader("🗣️ Interview Preparation Videos")
        categories = list(INTERVIEW_VIDEOS.keys())
        selected_cat = st.selectbox(
            "Select Interview Video Category",
            options=categories,
            key="interview_vid_cat"
        )
        if selected_cat:
            st.subheader(f"📂 {selected_cat}")
            videos = INTERVIEW_VIDEOS[selected_cat]
            cols = st.columns(2)
            for idx, (title, url) in enumerate(videos):
                with cols[idx % 2]:
                    st.markdown(f"**{title}**")
                    st.video(url)

    # Section 4: UPDATED AI Interview Coach 🤖 with Mock Interview and Enhanced Features
    elif page == "AI Interview Coach 🤖":
        st.subheader("🤖 AI Interview Coach")
        st.markdown("Practice role-specific interview questions with our AI coach. Get instant feedback on your answers and discover recommended courses!")

        # Create database table if not exists
        create_interview_database()

        # Domain and Role selection
        st.markdown('<div class="role-selector">', unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            selected_domain = st.selectbox(
                "Select Career Domain",
                options=list(COURSES_BY_CATEGORY.keys()),
                key="interview_domain_selection"
            )

        with col2:
            if selected_domain:
                roles = list(COURSES_BY_CATEGORY[selected_domain].keys())
                selected_role = st.selectbox(
                    "Select Target Role",
                    options=roles,
                    key="interview_role_selection"
                )
            else:
                selected_role = None

        st.markdown('</div>', unsafe_allow_html=True)
        
        if selected_domain and selected_role:
            # Initialize interview state
            if 'dynamic_interview_questions' not in st.session_state:
                st.session_state.dynamic_interview_questions = []
            if 'current_dynamic_interview_question' not in st.session_state:
                st.session_state.current_dynamic_interview_question = 0
            if 'dynamic_interview_answers' not in st.session_state:
                st.session_state.dynamic_interview_answers = []
            if 'dynamic_interview_scores' not in st.session_state:
                st.session_state.dynamic_interview_scores = []
            if 'dynamic_interview_feedbacks' not in st.session_state:
                st.session_state.dynamic_interview_feedbacks = []
            if 'dynamic_interview_completed' not in st.session_state:
                st.session_state.dynamic_interview_completed = False
            if 'dynamic_interview_started' not in st.session_state:
                st.session_state.dynamic_interview_started = False
            if 'dynamic_answer_submitted' not in st.session_state:
                st.session_state.dynamic_answer_submitted = False
            if 'current_interview_question_text' not in st.session_state:
                st.session_state.current_interview_question_text = ""
            if 'interview_domain' not in st.session_state or st.session_state.interview_domain != selected_domain:
                st.session_state.interview_domain = selected_domain
                st.session_state.interview_role = selected_role
                st.session_state.dynamic_interview_started = False
                st.session_state.dynamic_interview_completed = False
            if 'question_timer_start' not in st.session_state:
                st.session_state.question_timer_start = None
            if 'timer_seconds' not in st.session_state:
                st.session_state.timer_seconds = 120
            if 'interview_difficulty' not in st.session_state:
                st.session_state.interview_difficulty = "Medium"
            if 'original_num_questions' not in st.session_state:
                st.session_state.original_num_questions = 6

            # Start interview setup
            if not st.session_state.dynamic_interview_started:
                st.markdown(f"### Practice interview for: {selected_role}")

                col1, col2 = st.columns(2)

                with col1:
                    interview_type = st.selectbox(
                        "Interview Type",
                        options=["technical", "behavioral", "mixed"],
                        format_func=lambda x: x.title() + (" (Technical + Behavioral)" if x == "mixed" else ""),
                        key="dynamic_interview_type_select"
                    )

                with col2:
                    interview_difficulty = st.selectbox(
                        "Interview Difficulty",
                        options=["Easy", "Medium", "Hard"],
                        key="interview_difficulty_select",
                        index=1
                    )

                col3, col4 = st.columns(2)
                with col3:
                    num_questions = st.slider("Number of questions:", 5, 10, 6)

                with col4:
                    timer_seconds = st.slider("Time per question (seconds):", 60, 300, 120, step=30)

                if st.button("🚀 Start Mock Interview"):
                    with st.spinner("Generating personalized questions using AI..."):
                        # FIXED: Pass difficulty to question generation
                        selected_questions = generate_interview_questions_with_llm(
                            selected_domain,
                            selected_role,
                            interview_type,
                            num_questions,
                            interview_difficulty  # Now passing difficulty
                        )

                        if selected_questions:
                            # FIXED: Reset ALL interview state variables properly
                            # EXACT QUESTION COUNT: Enforce exact number of questions
                            selected_questions = selected_questions[:num_questions]

                            st.session_state.dynamic_interview_questions = selected_questions
                            st.session_state.original_num_questions = num_questions
                            st.session_state.current_dynamic_interview_question = 0
                            st.session_state.dynamic_interview_answers = []
                            st.session_state.dynamic_interview_scores = []
                            st.session_state.dynamic_interview_feedbacks = []
                            st.session_state.dynamic_interview_completed = False
                            st.session_state.dynamic_interview_started = True
                            st.session_state.dynamic_answer_submitted = False
                            st.session_state.current_interview_question_text = selected_questions[0]
                            st.session_state.question_timer_start = time.time()
                            st.session_state.timer_seconds = timer_seconds
                            st.session_state.interview_difficulty = interview_difficulty
                            st.success("Questions generated! Starting your mock interview...")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("Failed to generate questions. Please try again.")
            
            # Interview in progress
            elif st.session_state.dynamic_interview_started and not st.session_state.dynamic_interview_completed:
                # CRITICAL FIX: Properly count answered questions
                questions_answered = len(st.session_state.dynamic_interview_answers)
                total_questions = len(st.session_state.dynamic_interview_questions)
                current_index = st.session_state.current_dynamic_interview_question + 1

                # Display progress with correct counts in glassmorphism box
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, rgba(0, 195, 255, 0.08) 0%, rgba(0, 195, 255, 0.04) 100%);
                            backdrop-filter: blur(10px);
                            -webkit-backdrop-filter: blur(10px);
                            border: 1px solid rgba(0, 195, 255, 0.2);
                            border-radius: 12px;
                            padding: 16px 24px;
                            margin: 20px 0;
                            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1), inset 0 1px 0 rgba(255, 255, 255, 0.05);">
                    <p style="color: #ffffff; font-size: 16px; margin: 0; font-weight: 500;">
                        📊 Progress: Answered {questions_answered}/{st.session_state.original_num_questions} questions | Current Index: {current_index} of {st.session_state.original_num_questions}
                    </p>
                </div>
                """, unsafe_allow_html=True)

                if questions_answered < st.session_state.original_num_questions:
                    question = st.session_state.current_interview_question_text or st.session_state.dynamic_interview_questions[st.session_state.current_dynamic_interview_question]

                    # TIMER RESET: Reset timer every time a new question loads
                    if st.session_state.question_timer_start is None:
                        st.session_state.question_timer_start = time.time()

                    # Calculate remaining time
                    elapsed_time = time.time() - st.session_state.question_timer_start
                    remaining_time = max(0, st.session_state.timer_seconds - elapsed_time)

                    # Display timer
                    timer_minutes = int(remaining_time // 60)
                    timer_seconds_display = int(remaining_time % 60)
                    timer_urgent_class = "timer-urgent" if remaining_time <= 30 else ""

                    st.markdown(f"""
                    <div class="timer-container">
                        <div class="timer-display {timer_urgent_class}">
                            ⏰ Time Remaining: {timer_minutes:02d}:{timer_seconds_display:02d}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

                    # Timer progress bar
                    progress_value = (st.session_state.timer_seconds - remaining_time) / st.session_state.timer_seconds
                    st.progress(progress_value)

                    # Question display
                    st.markdown(f"""
                    <div class="quiz-card">
                        <h3 style="color: #00c3ff;">Question {questions_answered + 1} of {st.session_state.original_num_questions}</h3>
                        <h4 style="color: #ffffff; margin: 15px 0;">Role: {selected_role} | Difficulty: {st.session_state.interview_difficulty}</h4>
                        <p style="font-size: 18px; color: #ffffff; margin: 15px 0;">{question}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    # Add refresh button for regenerating all interview questions
                    col1, col2 = st.columns([3, 1])
                    with col2:
                        if st.button("🔄 Refresh Interview"):
                            # Clear all interview state
                            st.session_state.dynamic_interview_questions = []
                            st.session_state.current_dynamic_interview_question = 0
                            st.session_state.dynamic_interview_answers = []
                            st.session_state.dynamic_interview_scores = []
                            st.session_state.dynamic_interview_feedbacks = []
                            st.session_state.dynamic_interview_completed = False
                            st.session_state.dynamic_interview_started = False
                            st.session_state.dynamic_answer_submitted = False
                            st.session_state.current_interview_question_text = ""
                            st.session_state.question_timer_start = None
                            # Force regeneration
                            st.rerun()

                    # Answer input with character limit
                    answer_key = f"dynamic_interview_answer_{st.session_state.current_dynamic_interview_question}"
                    answer = st.text_area(
                        "Your answer:",
                        placeholder="Type your detailed answer here... (Use STAR method: Situation, Task, Action, Result)",
                        height=150,
                        max_chars=2000,
                        key=answer_key,
                        help="Maximum 2000 characters"
                    )

                    # Auto-submit logic when timer expires
                    if remaining_time <= 0 and not st.session_state.dynamic_answer_submitted:
                        if not answer.strip():
                            answer = "⚠️ No Answer"

                        # Evaluate answer using enhanced evaluation with role/domain context
                        with st.spinner("Evaluating your answer..."):
                            eval_result = evaluate_interview_answer_for_scores(
                                answer,
                                question,
                                st.session_state.interview_difficulty,
                                role=selected_role,
                                domain=selected_domain
                            )

                        # FIXED: Store answer, scores, and feedback - ensuring all are tracked properly
                        st.session_state.dynamic_interview_answers.append(answer)
                        st.session_state.dynamic_interview_scores.append(eval_result)
                        st.session_state.dynamic_interview_feedbacks.append(eval_result["feedback"])
                        st.session_state.dynamic_answer_submitted = True

                        # FIXED: Handle follow-up for Hard difficulty without breaking indexing
                        # Follow-ups are added but don't count toward original_num_questions
                        if st.session_state.interview_difficulty == "Hard" and eval_result.get("followup") and eval_result["followup"].strip():
                            # Only add follow-up if we haven't reached the end
                            if questions_answered < st.session_state.original_num_questions - 1:
                                st.session_state.dynamic_interview_questions.insert(
                                    st.session_state.current_dynamic_interview_question + 1,
                                    eval_result["followup"]
                                )

                        st.warning("⏰ Time's up! Answer auto-submitted.")
                        st.rerun()

                    # Submit answer button
                    if not st.session_state.dynamic_answer_submitted and remaining_time > 0:
                        if st.button("Submit Answer & Get Feedback"):
                            if answer.strip():
                                with st.spinner("Evaluating your answer..."):
                                    # Evaluate answer using enhanced evaluation with role/domain context
                                    eval_result = evaluate_interview_answer_for_scores(
                                        answer,
                                        question,
                                        st.session_state.interview_difficulty,
                                        role=selected_role,
                                        domain=selected_domain
                                    )

                                    # FIXED: Store answer, scores, and feedback ensuring proper tracking
                                    st.session_state.dynamic_interview_answers.append(answer)
                                    st.session_state.dynamic_interview_scores.append(eval_result)
                                    st.session_state.dynamic_interview_feedbacks.append(eval_result["feedback"])
                                    st.session_state.dynamic_answer_submitted = True

                                    # FIXED: Handle follow-up for Hard difficulty without breaking indexing
                                    if st.session_state.interview_difficulty == "Hard" and eval_result.get("followup") and eval_result["followup"].strip():
                                        # Only add follow-up if we haven't reached the end
                                        if questions_answered < st.session_state.original_num_questions - 1:
                                            st.session_state.dynamic_interview_questions.insert(
                                                st.session_state.current_dynamic_interview_question + 1,
                                                eval_result["followup"]
                                            )

                                    st.rerun()
                            else:
                                st.warning("Please provide an answer before proceeding.")

                    # Show feedback after answer submitted
                    if st.session_state.dynamic_answer_submitted:
                        current_score_dict = st.session_state.dynamic_interview_scores[-1]
                        avg_q_score = (current_score_dict["knowledge"] + current_score_dict["communication"] + current_score_dict["relevance"]) / 3

                        # Format feedback for display
                        feedback_text = current_score_dict["feedback"] if isinstance(current_score_dict["feedback"], str) else chr(10).join(current_score_dict["feedback"])
                        formatted_feedback = format_feedback_text(feedback_text)

                        st.markdown(f"""
                        <div style="background: linear-gradient(135deg, rgba(0, 195, 255, 0.1) 0%, rgba(0, 195, 255, 0.05) 100%);
                                    border: 1px solid rgba(0, 195, 255, 0.3); border-radius: 10px; padding: 15px; margin: 15px 0;">
                            <h4 style="color: #00c3ff;">Immediate Feedback:</h4>
                            <p style="color: #ffffff;">📊 Knowledge: {current_score_dict["knowledge"]}/10 | Communication: {current_score_dict["communication"]}/10 | Relevance: {current_score_dict["relevance"]}/10</p>
                            <p style="color: #ffffff;">⭐ Question Score: {avg_q_score:.1f}/10</p>
                            <div style="color: #ffffff; margin-top: 10px;">
                                {formatted_feedback}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                        # Show follow-up question for Hard difficulty
                        if st.session_state.interview_difficulty == "Hard" and current_score_dict.get("followup"):
                            st.info(f"🔎 Follow-Up Question: {current_score_dict['followup']}")

                        # Continue/Complete button
                        # CRITICAL FIX: Check if we've answered all original questions
                        if questions_answered >= st.session_state.original_num_questions:
                            # All questions answered, mark as complete
                            if st.button("Complete Interview 🏁"):
                                st.session_state.dynamic_interview_completed = True
                                st.rerun()
                        else:
                            # More questions to go
                            if st.button("Continue to Next Question ➡️"):
                                st.session_state.current_dynamic_interview_question += 1
                                st.session_state.dynamic_answer_submitted = False
                                if st.session_state.current_dynamic_interview_question < len(st.session_state.dynamic_interview_questions):
                                    st.session_state.current_interview_question_text = st.session_state.dynamic_interview_questions[st.session_state.current_dynamic_interview_question]
                                else:
                                    # Safety check - if we're out of questions but haven't answered all, generate one
                                    st.session_state.current_interview_question_text = f"Additional question for {selected_role}"
                                # TIMER RESET: Reset timer for next question
                                st.session_state.question_timer_start = time.time()
                                st.rerun()

                    # Progress bar for interview completion
                    interview_progress = questions_answered / st.session_state.original_num_questions
                    st.markdown("### Interview Progress")
                    st.progress(interview_progress)

                    # CRITICAL FIX: Review Previous Answers - show all properly
                    if len(st.session_state.dynamic_interview_answers) > 0:
                        with st.expander("📖 Review Previous Answers"):
                            # Show all submitted answers
                            num_to_show = len(st.session_state.dynamic_interview_answers)
                            for i in range(num_to_show):
                                if i < len(st.session_state.dynamic_interview_questions) and i < len(st.session_state.dynamic_interview_scores):
                                    prev_question = st.session_state.dynamic_interview_questions[i]
                                    prev_answer = st.session_state.dynamic_interview_answers[i]
                                    prev_scores = st.session_state.dynamic_interview_scores[i]
                                    prev_avg = (prev_scores["knowledge"] + prev_scores["communication"] + prev_scores["relevance"]) / 3

                                    # Show full answer (up to 500 chars in review, full in final)
                                    answer_preview = prev_answer[:500]
                                    if len(prev_answer) > 500:
                                        answer_preview += "..."

                                    st.markdown(f"**Question {i+1}:** {prev_question}")
                                    st.markdown(f"**Your Answer:** {answer_preview}")
                                    st.markdown(f"**Score:** {prev_avg:.1f}/10")
                                    if i < num_to_show - 1:  # Don't add separator after last item
                                        st.markdown("---")

                    # Auto-refresh for timer
                    if remaining_time > 0 and not st.session_state.dynamic_answer_submitted:
                        time.sleep(1)
                        st.rerun()
                else:
                    # CRITICAL FIX: All questions answered, move to completion automatically
                    st.session_state.dynamic_interview_completed = True
                    st.success(f"✅ Completed all {st.session_state.original_num_questions} questions!")
                    time.sleep(1)
                    st.rerun()
            
            # UNIFIED: Interview completed + Course Recommendations + DB + PDF
            elif st.session_state.dynamic_interview_completed:
                # Calculate average scores for each dimension
                knowledge_scores = [s["knowledge"] for s in st.session_state.dynamic_interview_scores]
                communication_scores = [s["communication"] for s in st.session_state.dynamic_interview_scores]
                relevance_scores = [s["relevance"] for s in st.session_state.dynamic_interview_scores]

                avg_knowledge = sum(knowledge_scores) / len(knowledge_scores)
                avg_communication = sum(communication_scores) / len(communication_scores)
                avg_relevance = sum(relevance_scores) / len(relevance_scores)
                overall_avg = (avg_knowledge + avg_communication + avg_relevance) / 3

                # Determine badge based on overall average
                if overall_avg >= 8.5:
                    badge = "Interview Ready"
                    badge_emoji = "🏆"
                elif overall_avg >= 7.0:
                    badge = "Excellent"
                    badge_emoji = "🌟"
                elif overall_avg >= 5.0:
                    badge = "Good"
                    badge_emoji = "👍"
                else:
                    badge = "Needs Practice"
                    badge_emoji = "💪"

                st.markdown(f"""
                <div class="badge-container">
                    <h2 style="margin: 0; color: #ffffff; font-size: 28px; font-weight: 600;">🎉 Mock Interview Complete!</h2>
                    <div style="margin: 30px 0;">
                        <div class="score-display">{overall_avg:.1f}/10</div>
                        <h3 style="color: #ffffff; margin: 15px 0; font-size: 24px; font-weight: 500;">{badge_emoji} {badge}</h3>
                    </div>
                    <p style="color: rgba(255, 255, 255, 0.85); font-size: 16px; margin: 8px 0;">Role: {selected_role} in {selected_domain}</p>
                    <p style="color: rgba(255, 255, 255, 0.85); font-size: 16px; margin: 8px 0;">Difficulty: {st.session_state.interview_difficulty}</p>
                </div>
                """, unsafe_allow_html=True)

                # Create radar chart for skills
                st.markdown('<div class="radar-container">', unsafe_allow_html=True)
                st.subheader("📊 Performance Radar Chart")

                radar_data = {
                    "Communication": avg_communication,
                    "Knowledge": avg_knowledge,
                    "Confidence": avg_relevance
                }

                fig = go.Figure()
                fig.add_trace(go.Scatterpolar(
                    r=list(radar_data.values()),
                    theta=list(radar_data.keys()),
                    fill='toself',
                    name='Performance',
                    line=dict(color='#00c3ff', width=2),
                    fillcolor='rgba(0, 195, 255, 0.2)'
                ))

                fig.update_layout(
                    polar=dict(
                        radialaxis=dict(
                            visible=True,
                            range=[0, 10],
                            tickfont=dict(color='white', size=10),
                            gridcolor='rgba(255, 255, 255, 0.2)'
                        ),
                        angularaxis=dict(
                            tickfont=dict(color='white', size=12),
                            gridcolor='rgba(255, 255, 255, 0.2)'
                        ),
                        bgcolor='rgba(0, 0, 0, 0)'
                    ),
                    showlegend=False,
                    title=dict(
                        text="Interview Performance Metrics",
                        x=0.5,
                        font=dict(color='#00c3ff', size=16)
                    ),
                    paper_bgcolor='rgba(0, 0, 0, 0)',
                    plot_bgcolor='rgba(0, 0, 0, 0)',
                    font=dict(color='white'),
                    height=400
                )

                st.plotly_chart(fig, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)

                # Strengths and Weaknesses
                st.subheader("💡 Performance Analysis")
                col1, col2 = st.columns(2)

                metrics = [("Communication", avg_communication), ("Knowledge", avg_knowledge), ("Confidence", avg_relevance)]
                metrics_sorted = sorted(metrics, key=lambda x: x[1], reverse=True)

                with col1:
                    st.markdown("**🌟 Strengths:**")
                    for name, score in metrics_sorted[:2]:
                        st.markdown(f"- {name}: {score:.1f}/10")

                with col2:
                    st.markdown("**📈 Areas to Improve:**")
                    for name, score in metrics_sorted[-2:]:
                        st.markdown(f"- {name}: {score:.1f}/10")

                # FIXED: Show detailed Q&A results with full answers and proper matching
                st.markdown("---")
                st.subheader("📋 Detailed Q&A Review:")

                # Ensure we only show as many Q&A pairs as we have complete data for
                num_complete_qa = min(
                    len(st.session_state.dynamic_interview_scores),
                    len(st.session_state.dynamic_interview_answers),
                    len(st.session_state.dynamic_interview_feedbacks),
                    len(st.session_state.dynamic_interview_questions)
                )

                for i in range(num_complete_qa):
                    score_dict = st.session_state.dynamic_interview_scores[i]
                    answer = st.session_state.dynamic_interview_answers[i]
                    feedback = st.session_state.dynamic_interview_feedbacks[i]
                    question = st.session_state.dynamic_interview_questions[i]

                    q_avg = (score_dict["knowledge"] + score_dict["communication"] + score_dict["relevance"]) / 3

                    with st.expander(f"Question {i+1}: Score {q_avg:.1f}/10"):
                        st.write(f"**Question:** {question}")
                        st.write(f"**Your Answer:** {answer}")  # Show full answer
                        st.write(f"**Scores:** Knowledge: {score_dict['knowledge']}/10 | Communication: {score_dict['communication']}/10 | Relevance: {score_dict['relevance']}/10")

                        # Format and display feedback as bullet points
                        feedback_text = "\n".join(feedback) if isinstance(feedback, list) else feedback
                        formatted_feedback = format_feedback_text(feedback_text)
                        st.markdown(formatted_feedback, unsafe_allow_html=True)

                # Save to database
                username = st.session_state.get("username", "Guest")
                feedback_summary = f"Strengths: {metrics_sorted[0][0]}, {metrics_sorted[1][0]}. Weaknesses: {metrics_sorted[-1][0]}, {metrics_sorted[-2][0]}."

                if save_interview_result(username, selected_role, selected_domain, overall_avg, st.session_state.original_num_questions, feedback_summary):
                    log_user_action(username, "completed_interview")

                # Generate PDF report
                st.markdown("---")
                st.subheader("📄 Download Interview Report")

                completed_on = get_ist_time()

                # CRITICAL FIX: Ensure all arrays have same length for PDF generation
                num_complete = min(
                    len(st.session_state.dynamic_interview_questions),
                    len(st.session_state.dynamic_interview_answers),
                    len(st.session_state.dynamic_interview_scores),
                    len(st.session_state.dynamic_interview_feedbacks)
                )

                pdf_bytes = generate_interview_pdf_report(
                    username,
                    selected_role,
                    selected_domain,
                    completed_on,
                    st.session_state.dynamic_interview_questions[:num_complete],
                    st.session_state.dynamic_interview_answers[:num_complete],
                    st.session_state.dynamic_interview_scores[:num_complete],
                    st.session_state.dynamic_interview_feedbacks[:num_complete],
                    overall_avg,
                    badge,
                    difficulty=st.session_state.interview_difficulty
                )

                if pdf_bytes:
                    st.download_button(
                        label="📄 Download Interview Report",
                        data=pdf_bytes,
                        file_name=f"interview_report_{username}_{selected_role.replace(' ', '_')}_{completed_on.split()[0]}.pdf",
                        mime="application/pdf"
                    )
                else:
                    st.warning("PDF generation failed. You can still review your results above.")

                # UNIFIED: Display recommended courses by difficulty
                st.markdown("---")
                st.subheader("📚 Recommended Courses for Your Career Growth")
                st.markdown(f"Based on your interview practice for **{selected_role}** in **{selected_domain}**, here are our course recommendations organized by difficulty level:")

                courses = get_courses_for_role(selected_domain, selected_role)
                if courses:
                    display_courses_by_difficulty(courses, selected_role)
                else:
                    st.info("No specific courses found for this role. Explore our course categories to find relevant learning resources!")

                # FIXED: Restart button - properly resets ALL interview state
                if st.button("🔄 Practice Again"):
                    # Reset all interview-related session state variables
                    st.session_state.dynamic_interview_started = False
                    st.session_state.dynamic_interview_completed = False
                    st.session_state.dynamic_interview_questions = []
                    st.session_state.current_dynamic_interview_question = 0
                    st.session_state.dynamic_interview_answers = []
                    st.session_state.dynamic_interview_scores = []
                    st.session_state.dynamic_interview_feedbacks = []
                    st.session_state.dynamic_answer_submitted = False
                    st.session_state.current_interview_question_text = ""
                    st.session_state.question_timer_start = None
                    st.session_state.timer_seconds = 120
                    st.session_state.interview_difficulty = "Medium"
                    st.session_state.original_num_questions = 6
                    st.rerun()
        else:
            st.info("Please select both a career domain and target role to start the interview practice.")
if tab5:
	with tab5:
		import sqlite3
		import pandas as pd
		import matplotlib.pyplot as plt
		import numpy as np
		import streamlit as st
		from datetime import datetime, timedelta
		import plotly.express as px
		import plotly.graph_objects as go
		from plotly.subplots import make_subplots
		import time
		import glob, os

		# Import enhanced database manager functions
		from db_manager import (
			get_top_domains_by_score,
			get_resume_count_by_day,
			get_average_ats_by_domain,
			get_domain_distribution,
			get_bias_distribution,
			filter_candidates_by_date,
			delete_candidate_by_id,
			get_all_candidates,
			get_candidate_by_id,
			get_domain_performance_stats,
			get_daily_ats_stats,
			get_flagged_candidates,
			get_database_stats,
			analyze_domain_transitions,
			export_to_csv,
			cleanup_old_records,
			DatabaseManager
		)

		# Initialize enhanced database manager
		@st.cache_resource
		def get_db_manager():
			return DatabaseManager()

		db_manager = get_db_manager()

		def create_enhanced_pie_chart(df, values_col, labels_col, title):
			"""Create an enhanced pie chart with better styling"""
			fig = px.pie(
				df, 
				values=values_col, 
				names=labels_col,
				title=title,
				color_discrete_sequence=px.colors.qualitative.Set3
			)
			fig.update_traces(
				textposition='inside', 
				textinfo='percent+label',
				hovertemplate='<b>%{label}</b><br>Count: %{value}<br>Percentage: %{percent}<extra></extra>'
			)
			fig.update_layout(
				showlegend=True,
				legend=dict(orientation="v", yanchor="middle", y=0.5, xanchor="left", x=1.01),
				margin=dict(t=50, b=50, l=50, r=150)
			)
			return fig

		def create_enhanced_bar_chart(df, x_col, y_col, title, orientation='v'):
			"""Create enhanced bar chart with better interactivity"""
			if orientation == 'v':
				fig = px.bar(df, x=x_col, y=y_col, title=title, 
							color=y_col, color_continuous_scale='viridis')
				fig.update_xaxes(tickangle=45)
			else:
				fig = px.bar(df, x=y_col, y=x_col, title=title, orientation='h',
							color=y_col, color_continuous_scale='viridis')
			
			fig.update_traces(
				hovertemplate='<b>%{y if orientation == "v" else x}</b><br>Value: %{x if orientation == "v" else y}<extra></extra>'
			)
			fig.update_layout(showlegend=False, margin=dict(t=50, b=50, l=50, r=50))
			return fig

		def load_domain_distribution():
			"""Enhanced domain distribution loading with error handling"""
			try:
				df = get_domain_distribution()
				if not df.empty:
					df = df.sort_values(by="count", ascending=False).reset_index(drop=True)
					return df
			except Exception as e:
				st.error(f"Error loading domain distribution: {e}")
			return pd.DataFrame()

		# Enhanced Data Loading with Caching
		@st.cache_data(ttl=300)  # Cache for 5 minutes
		def load_all_candidates():
			try:
				return get_all_candidates()
			except Exception as e:
				st.error(f"Error loading candidates: {e}")
				return pd.DataFrame()

		# -------- Glassmorphism Styles with Shimmer --------
		st.markdown("""
		<style>
		.glass-box {
			background: rgba(10, 20, 40, 0.55);
			border-radius: 18px;
			padding: 2rem;
			backdrop-filter: blur(14px);
			border: 1px solid rgba(0, 200, 255, 0.35);
			box-shadow: 0 8px 32px rgba(0, 200, 255, 0.25);
			position: relative;
			overflow: hidden;
			text-align: center;
			margin-bottom: 2rem;
		}
		.glass-box::before {
			content: "";
			position: absolute;
			top: -50%;
			left: -50%;
			width: 200%;
			height: 200%;
			background: linear-gradient(
				120deg,
				rgba(255,255,255,0.15) 0%,
				rgba(255,255,255,0.05) 40%,
				transparent 60%
			);
			transform: rotate(25deg);
			animation: shimmer 6s infinite;
		}
		@keyframes shimmer {
			0% { top: -50%; left: -50%; }
			50% { top: 100%; left: 100%; }
			100% { top: -50%; left: -50%; }
		}
		.glass-box h1, .glass-box h2 {
			color: #4da6ff;
			text-shadow: 0 0 12px rgba(0,200,255,0.7);
			margin: 0 0 0.5rem 0;
			font-weight: 600;
		}
		.glass-box p {
			color: #cce6ff;
			margin: 0;
			font-size: 0.95rem;
		}

		/* Glassy input fields */
		.stTextInput > div > div > input {
			background: rgba(255, 255, 255, 0.08) !important;
			border: 1px solid rgba(0, 200, 255, 0.3) !important;
			border-radius: 12px !important;
			padding: 10px !important;
			color: #e6f7ff !important;
			font-weight: 500 !important;
			backdrop-filter: blur(10px) !important;
		}
		.stTextInput > div > div > input:focus {
			border: 1px solid rgba(0, 200, 255, 0.8) !important;
			box-shadow: 0 0 12px rgba(0, 200, 255, 0.6) !important;
			outline: none !important;
		}

		/* Glassy button */
		.stButton > button {
			background: rgba(0, 200, 255, 0.15);
			border: 1px solid rgba(0, 200, 255, 0.4);
			border-radius: 12px;
			color: #e6f7ff;
			padding: 0.6rem 1.2rem;
			font-weight: bold;
			backdrop-filter: blur(8px);
			transition: all 0.3s ease;
		}
		.stButton > button:hover {
			background: rgba(0, 200, 255, 0.3);
			box-shadow: 0 0 16px rgba(0, 200, 255, 0.7);
			transform: translateY(-2px);
		}
		</style>
		""", unsafe_allow_html=True)

		# ---------------- Enhanced Authentication System ----------------
		if "admin_logged_in" not in st.session_state:
			st.session_state.admin_logged_in = False

		if not st.session_state.admin_logged_in:
			st.markdown("""
			<div class="glass-box">
				<h2>🔐 Admin Authentication Required</h2>
				<p>Please enter your email and password to access the admin dashboard</p>
			</div>
			""", unsafe_allow_html=True)
			
			col1, col2, col3 = st.columns([1, 2, 1])
			with col2:
				email = st.text_input("📧 Enter Admin Email", placeholder="Enter email...")
				password = st.text_input("🔑 Enter Admin Password", type="password", placeholder="Enter password...")
				login_clicked = st.button("🚀 Login", use_container_width=True)

				if login_clicked:
					valid_email = "admin@example.com"
					valid_password = "Swagato@2002"

					if email == valid_email and password == valid_password:
						st.session_state.admin_logged_in = True
						st.success("✅ Authentication successful! Redirecting to dashboard...")
						st.rerun()
					else:
						msg_placeholder = st.empty()
						msg_placeholder.markdown("""
							<div style='
								background-color: #ff4d4d;
								color: white;
								padding: 10px 15px;
								border-radius: 10px;
								text-align: center;
								animation: slideDown 0.5s ease-in-out;
							'>❌ Invalid credentials. Please try again.</div>
							<style>
							@keyframes slideDown {
								0% {transform: translateY(-50px); opacity: 0;}
								100% {transform: translateY(0); opacity: 1;}
							}
							</style>
						""", unsafe_allow_html=True)
						time.sleep(3)
						msg_placeholder.empty()

			st.stop()

		# ---------------- Enhanced Header with Database Stats ----------------
		st.markdown("""
		<div class="glass-box">
			<h1>🛡️ Enhanced Admin Database Panel</h1>
			<p>Advanced Resume Analysis System Dashboard</p>
		</div>
		""", unsafe_allow_html=True)

		# Enhanced Control Panel
		col1, col2, col3, col4 = st.columns(4)
		with col1:
			if st.button("🔄 Refresh All Data", use_container_width=True):
				st.cache_data.clear()
				st.rerun()
		with col2:
			if st.button("📊 Database Stats", use_container_width=True):
				st.session_state.show_db_stats = True
		with col3:
			if st.button("🧹 Cleanup Old Records", use_container_width=True):
				st.session_state.show_cleanup = True
		with col4:
			if st.button("🚪 Secure Logout", use_container_width=True):
				st.session_state.admin_logged_in = False
				st.success("👋 Logged out successfully.")
				st.rerun()

		# Database Statistics Panel
		if st.session_state.get('show_db_stats', False):
			with st.expander("📈 Database Statistics", expanded=True):
				try:
					stats = get_database_stats()
					if stats:
						col1, col2, col3, col4 = st.columns(4)
						with col1:
							st.metric("Total Candidates", stats.get('total_candidates', 0))
						with col2:
							st.metric("Average ATS Score", f"{stats.get('avg_ats_score', 0):.2f}")
						with col3:
							st.metric("Unique Domains", stats.get('unique_domains', 0))
						with col4:
							st.metric("Database Size", f"{stats.get('database_size_mb', 0):.2f} MB")
						
						col5, col6 = st.columns(2)
						with col5:
							st.metric("Earliest Record", stats.get('earliest_date', 'N/A'))
						with col6:
							st.metric("Latest Record", stats.get('latest_date', 'N/A'))
				except Exception as e:
					st.error(f"Error loading database statistics: {e}")

		# Cleanup Panel
		if st.session_state.get('show_cleanup', False):
			with st.expander("🧹 Database Cleanup", expanded=True):
				days_to_keep = st.slider("Days to Keep", 30, 730, 365)
				if st.button("⚠️ Cleanup Old Records"):
					try:
						deleted_count = cleanup_old_records(days_to_keep)
						if deleted_count > 0:
							st.success(f"✅ Cleaned up {deleted_count} old records")
						else:
							st.info("ℹ️ No old records found to cleanup")
					except Exception as e:
						st.error(f"Error during cleanup: {e}")

		st.markdown("<hr style='border-top: 2px solid #bbb; margin: 2rem 0;'>", unsafe_allow_html=True)

		df = load_all_candidates()

		# Enhanced Search and Filter Section
		st.markdown("### 🔍 Advanced Search & Filters")
		
		col1, col2 = st.columns(2)
		with col1:
			search = st.text_input("🔍 Search by Candidate Name", placeholder="Enter candidate name...")
			if search:
				df = df[df["candidate_name"].str.contains(search, case=False, na=False)]
		
		with col2:
			domain_filter = st.selectbox("🏢 Filter by Domain", 
									options=["All Domains"] + list(df["domain"].unique()) if not df.empty else ["All Domains"])
			if domain_filter != "All Domains":
				df = df[df["domain"] == domain_filter]

		# Enhanced Date Filter
		st.markdown("#### 📅 Date Range Filter")
		col1, col2, col3 = st.columns(3)
		with col1:
			start_date = st.date_input("📅 Start Date", value=datetime.now() - timedelta(days=30))
		with col2:
			end_date = st.date_input("📅 End Date", value=datetime.now())
		with col3:
			if st.button("🎯 Apply Filters", use_container_width=True):
				try:
					df = filter_candidates_by_date(str(start_date), str(end_date))
					if domain_filter != "All Domains":
						df = df[df["domain"] == domain_filter]
					if search:
						df = df[df["candidate_name"].str.contains(search, case=False, na=False)]
					st.success(f"✅ Filters applied. Found {len(df)} candidates.")
				except Exception as e:
					st.error(f"Error applying filters: {e}")

		# Enhanced Candidates Display
		if df.empty:
			st.info("ℹ️ No candidate data available with current filters.")
		else:
			st.markdown(f"### 📋 Candidates Overview ({len(df)} records)")
			
			# Enhanced metrics
			col1, col2, col3, col4 = st.columns(4)
			with col1:
				st.metric("Total Candidates", len(df))
			with col2:
				st.metric("Avg ATS Score", f"{df['ats_score'].mean():.2f}")
			with col3:
				st.metric("Avg Bias Score", f"{df['bias_score'].mean():.3f}")
			with col4:
				st.metric("Unique Domains", df['domain'].nunique())

			# Enhanced data display with sorting
			sort_column = st.selectbox("📊 Sort by", 
								options=['timestamp', 'ats_score', 'bias_score', 'candidate_name', 'domain'])
			sort_order = st.radio("Sort Order", ["Descending", "Ascending"], horizontal=True)
			
			df_sorted = df.sort_values(by=sort_column, ascending=(sort_order == "Ascending"))
			
			# Display with enhanced formatting
			st.dataframe(
				df_sorted.style.format({
					'ats_score': '{:.0f}',
					'edu_score': '{:.0f}',
					'exp_score': '{:.0f}',
					'skills_score': '{:.0f}',
					'lang_score': '{:.0f}',
					'keyword_score': '{:.0f}',
					'bias_score': '{:.3f}'
				}),
				use_container_width=True,
				height=400
			)

			# Enhanced Export Options
			col1, col2 = st.columns(2)
			with col1:
				csv_data = df_sorted.to_csv(index=False)
				st.download_button(
					label="📥 Download Filtered Data (CSV)",
					data=csv_data,
					file_name=f"candidates_filtered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
					mime="text/csv",
					use_container_width=True
				)
			with col2:
				if st.button("📤 Export All Data", use_container_width=True):
					try:
						filename = f"full_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
						if export_to_csv(filename):
							st.success(f"✅ Data exported to {filename}")
						else:
							st.error("❌ Export failed")
					except Exception as e:
						st.error(f"Export error: {e}")

			st.markdown("### 📂 Export Archive")
			export_files = sorted(glob.glob("full_export_*.csv"), reverse=True)

			if export_files:
				for file in export_files:
					with open(file, "rb") as f:
						st.download_button(
							label=f"⬇️ Download {os.path.basename(file)}",
							data=f,
							file_name=os.path.basename(file),
							mime="text/csv",
							use_container_width=True
						)
			else:
				st.info("📭 No export files found yet.")

			# Enhanced Delete Functionality
			with st.expander("🗑️ Delete Candidate", expanded=False):
				st.warning("⚠️ This action cannot be undone!")
				delete_id = st.number_input("Enter Candidate ID", min_value=1, step=1, key="delete_id")
				
				if delete_id in df["id"].values:
					candidate_info = get_candidate_by_id(delete_id)
					if not candidate_info.empty:
						st.info("📄 Candidate to be deleted:")
						st.dataframe(candidate_info, use_container_width=True)
						
						if st.button("❌ Confirm Delete", type="primary"):
							try:
								if delete_candidate_by_id(delete_id):
									st.success(f"✅ Candidate with ID {delete_id} deleted successfully.")
									st.cache_data.clear()
									st.rerun()
								else:
									st.error("❌ Failed to delete candidate.")
							except Exception as e:
								st.error(f"Delete error: {e}")
				elif delete_id > 0:
					st.error("❌ Candidate ID not found.")

		# Enhanced Analytics Section
		st.markdown("<hr style='border-top: 2px solid #bbb; margin: 2rem 0;'>", unsafe_allow_html=True)
		st.markdown("## 📊 Advanced Analytics Dashboard")

		# Enhanced Top Domains Analysis
		st.markdown("### 🏆 Top Performing Domains")
		
		try:
			top_domains = get_top_domains_by_score(limit=10)
			if top_domains:
				df_top = pd.DataFrame(top_domains, columns=["domain", "avg_ats", "count"])
				
				col1, col2 = st.columns([1, 2])
				with col1:
					sort_order = st.radio("📊 Sort by ATS", ["⬆️ High to Low", "⬇️ Low to High"], horizontal=True)
					limit = st.slider("Show Top N Domains", 1, len(df_top), value=min(8, len(df_top)))
				
				ascending = sort_order == "⬇️ Low to High"
				df_sorted = df_top.sort_values(by="avg_ats", ascending=ascending).head(limit)
				
				# Interactive chart
				fig = create_enhanced_bar_chart(df_sorted, "domain", "avg_ats", 
										"Average ATS Score by Domain", orientation='h')
				st.plotly_chart(fig, use_container_width=True)
				
				# Enhanced domain cards with glassmorphism
				st.markdown("""
				<style>
				@keyframes tab5-shimmer {
					0% { background-position: -200% 0; }
					100% { background-position: 200% 0; }
				}
				.tab5-domain-card {
					background: rgba(10, 20, 40, 0.3);
					backdrop-filter: blur(10px);
					border: 1px solid rgba(0, 200, 255, 0.2);
					box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3);
					border-radius: 15px;
					padding: 15px;
					margin-bottom: 15px;
					transition: all 0.3s ease;
					cursor: pointer;
					position: relative;
					overflow: hidden;
				}
				.tab5-domain-card::before {
					content: "";
					position: absolute;
					top: 0;
					left: 0;
					width: 100%;
					height: 100%;
					background: linear-gradient(
						120deg,
						transparent 0%,
						rgba(255, 255, 255, 0.08) 50%,
						transparent 100%
					);
					background-size: 200% 100%;
					opacity: 0;
					transition: opacity 0.3s ease;
				}
				.tab5-domain-card:hover::before {
					opacity: 1;
					animation: tab5-shimmer 1.5s ease-in-out infinite;
				}
				.tab5-domain-card:hover {
					transform: translateY(-2px);
					border-color: rgba(0, 200, 255, 0.35);
					background: rgba(10, 20, 40, 0.4);
				}
				</style>
				""", unsafe_allow_html=True)

				for i, row in df_sorted.iterrows():
					progress_value = row['avg_ats'] / 100
					st.markdown(f"""
					<div class="tab5-domain-card">
						<div style="display: flex; justify-content: space-between; align-items: center; position: relative; z-index: 1;">
							<h4 style="margin: 0; color: #5eb8ff;">📁 {row['domain']}</h4>
							<span style="
								background: rgba(0, 200, 255, 0.1);
								border: 1px solid rgba(0, 200, 255, 0.25);
								color: #5eb8ff;
								padding: 5px 10px;
								border-radius: 20px;
								font-size: 12px;
								font-weight: bold;
								backdrop-filter: blur(8px);
							">
								Rank #{i+1}
							</span>
						</div>
						<div style="margin: 10px 0; position: relative; z-index: 1;">
							<div style="
								background: rgba(255, 255, 255, 0.05);
								border-radius: 10px;
								height: 8px;
								overflow: hidden;
							">
								<div style="
									background: linear-gradient(90deg, rgba(0, 200, 255, 0.4), rgba(0, 255, 200, 0.5));
									height: 100%;
									width: {progress_value*100}%;
									transition: width 0.3s ease;
								"></div>
							</div>
						</div>
						<div style="display: flex; justify-content: space-between; margin-top: 10px; position: relative; z-index: 1;">
							<span style="color: #cce6ff;"><b>🧠 Avg ATS:</b> <span style="color: #5eb8ff; font-weight: bold;">{row['avg_ats']:.2f}</span></span>
							<span style="color: #cce6ff;"><b>📄 Resumes:</b> <span style="color: #5eb8ff; font-weight: bold;">{row['count']}</span></span>
						</div>
					</div>
					""", unsafe_allow_html=True)
			else:
				st.info("ℹ️ No domain performance data available.")
		except Exception as e:
			st.error(f"Error loading top domains: {e}")

		# Enhanced Domain Distribution
		st.markdown("### 📊 Domain Distribution Analysis")

		try:
			df_domain_dist = load_domain_distribution()
			if not df_domain_dist.empty:
				col1, col2 = st.columns(2)
				with col1:
					chart_type = st.radio(
						"📊 Visualization Type:",
						["📈 Interactive Bar Chart", "🥧 Interactive Pie Chart"],
						horizontal=True
					)
				with col2:
					max_val = len(df_domain_dist)
					if max_val <= 5:
						show_top_n = max_val  # No slider, just show all available domains
					else:
						show_top_n = st.slider(
							"Show Top N Domains",
							min_value=5,
							max_value=max_val,
							value=min(10, max_val)
						)

				df_top_domains = df_domain_dist.head(show_top_n)

				if chart_type == "📈 Interactive Bar Chart":
					fig = create_enhanced_bar_chart(df_top_domains, "domain", "count", 
											"Resume Count by Domain")
					st.plotly_chart(fig, use_container_width=True)
				else:
					fig = create_enhanced_pie_chart(df_top_domains, "count", "domain", 
											"Domain Distribution")
					st.plotly_chart(fig, use_container_width=True)

				# Summary statistics
				with st.expander("📋 Domain Statistics Summary"):
					st.dataframe(
						df_domain_dist.style.format({'percentage': '{:.2f}%'}),
						use_container_width=True
					)
			else:
				st.info("ℹ️ No domain distribution data available.")
		except Exception as e:
			st.error(f"Error loading domain distribution: {e}")

		# Enhanced ATS Performance Analysis
		st.markdown("### 📈 ATS Performance Analysis")
		
		try:
			df_ats = get_average_ats_by_domain()
			if not df_ats.empty:
				col1, col2 = st.columns(2)
				with col1:
					chart_orientation = st.radio("Chart Style", ["Vertical", "Horizontal"], horizontal=True)
				with col2:
					color_scheme = st.selectbox("Color Scheme", 
										["plasma", "viridis", "inferno", "magma", "turbo"])
				
				orientation = 'v' if chart_orientation == "Vertical" else 'h'
				fig = px.bar(df_ats, 
							x="domain" if orientation == 'v' else "avg_ats_score",
							y="avg_ats_score" if orientation == 'v' else "domain",
							title="Average ATS Score by Domain",
							orientation=orientation,
							color="avg_ats_score",
							color_continuous_scale=color_scheme,
							text="avg_ats_score",
							template="plotly_dark")  # Use dark theme for better readability
				
				fig.update_traces(texttemplate='%{text:.1f}', textposition='outside')
				if orientation == 'v':
					fig.update_xaxes(tickangle=45)
				
				# Enhanced layout for better readability
				fig.update_layout(
					showlegend=False,
					plot_bgcolor='rgba(0,0,0,0.1)',
					paper_bgcolor='rgba(0,0,0,0.05)',
					font=dict(color='white', size=12),
					title=dict(font=dict(size=16, color='white')),
					xaxis=dict(
						gridcolor='rgba(255,255,255,0.2)',
						tickfont=dict(color='white')
					),
					yaxis=dict(
						gridcolor='rgba(255,255,255,0.2)',
						tickfont=dict(color='white')
					),
					margin=dict(t=60, b=80, l=80, r=50)
				)
				
				st.plotly_chart(fig, use_container_width=True)
			else:
				st.info("ℹ️ No ATS performance data available.")
		except Exception as e:
			st.error(f"Error loading ATS performance data: {e}")

		# Enhanced Timeline Analysis
		st.markdown("### 📈 Resume Upload Timeline & Trends")
		
		try:
			df_timeline = get_resume_count_by_day()
			df_daily_ats = get_daily_ats_stats(days_limit=90)
			
			if not df_timeline.empty:
				df_timeline = df_timeline.sort_values("day")
				df_timeline["7_day_avg"] = df_timeline["count"].rolling(window=7, min_periods=1).mean()
				df_timeline["30_day_avg"] = df_timeline["count"].rolling(window=30, min_periods=1).mean()
				
				# Create subplot with proper spacing and formatting
				fig = make_subplots(
					rows=2, cols=1,
					subplot_titles=('Daily Upload Count with Moving Averages', 'Daily Average ATS Score Trend'),
					vertical_spacing=0.25,
					specs=[[{"secondary_y": False}], [{"secondary_y": False}]]
				)
				
				# Convert day column to datetime for proper spacing
				df_timeline['day'] = pd.to_datetime(df_timeline['day'])
				
				# Upload count plot
				fig.add_trace(
					go.Scatter(x=df_timeline["day"], y=df_timeline["count"], 
								mode='lines+markers', name='Daily Uploads',
								line=dict(color='#1f77b4', width=2),
								marker=dict(size=6)),
					row=1, col=1
				)
				
				fig.add_trace(
					go.Scatter(x=df_timeline["day"], y=df_timeline["7_day_avg"], 
								mode='lines', name='7-Day Average',
								line=dict(color='#ff7f0e', width=2, dash='dash')),
					row=1, col=1
				)
				
				fig.add_trace(
					go.Scatter(x=df_timeline["day"], y=df_timeline["30_day_avg"], 
								mode='lines', name='30-Day Average',
								line=dict(color='#2ca02c', width=2, dash='dot')),
					row=1, col=1
				)
				
				# ATS trend plot
				if not df_daily_ats.empty:
					df_daily_ats['date'] = pd.to_datetime(df_daily_ats['date'])
					fig.add_trace(
						go.Scatter(x=df_daily_ats["date"], y=df_daily_ats["avg_ats"], 
									mode='lines+markers', name='Daily Avg ATS',
									line=dict(color='#d62728', width=2),
									marker=dict(size=6)),
						row=2, col=1
					)
				
				# Update layout for better spacing and readability
				fig.update_layout(
					height=700, 
					showlegend=True,
					legend=dict(
						orientation="h",
						yanchor="bottom",
						y=1.02,
						xanchor="right",
						x=1
					),
					margin=dict(t=80, b=70, l=50, r=50)
				)
				
				# Update x-axes for proper date formatting and spacing
				fig.update_xaxes(title_text="Date", row=2, col=1)
				fig.update_xaxes(
					tickformat="%Y-%m-%d",
					tickangle=30,
					dtick="D1" if len(df_timeline) <= 30 else "D7",
					row=1, col=1
				)
				fig.update_xaxes(
					tickformat="%Y-%m-%d",
					tickangle=30,
					dtick="D1" if len(df_daily_ats) <= 30 else "D7",
					row=2, col=1
				)
				
				fig.update_yaxes(title_text="Upload Count", row=1, col=1)
				fig.update_yaxes(title_text="Average ATS Score", row=2, col=1)
				
				st.plotly_chart(fig, use_container_width=True)
				
				# Timeline statistics
				col1, col2, col3, col4 = st.columns(4)
				with col1:
					st.metric("Total Days", len(df_timeline))
				with col2:
					st.metric("Peak Daily Uploads", df_timeline["count"].max())
				with col3:
					st.metric("Avg Daily Uploads", f"{df_timeline['count'].mean():.1f}")
				with col4:
					if not df_daily_ats.empty:
						st.metric("Avg ATS Trend", f"{df_daily_ats['avg_ats'].mean():.2f}")
			else:
				st.info("ℹ️ No timeline data available.")
		except Exception as e:
			st.error(f"Error loading timeline data: {e}")

		# Enhanced Bias Analysis
		st.markdown("### 🧠 Advanced Bias Analysis")
		
		col1, col2 = st.columns(2)
		with col1:
			bias_threshold_pie = st.slider("Bias Detection Threshold", 
									min_value=0.0, max_value=1.0, value=0.6, step=0.05)
		with col2:
			analysis_type = st.radio("Analysis Type", ["Distribution", "Flagged Candidates"], horizontal=True)
		
		try:
			if analysis_type == "Distribution":
				df_bias = get_bias_distribution(threshold=bias_threshold_pie)
				if not df_bias.empty and "bias_category" in df_bias.columns:
					fig = create_enhanced_pie_chart(df_bias, "count", "bias_category", 
											f"Bias Distribution (Threshold: {bias_threshold_pie})")
					st.plotly_chart(fig, use_container_width=True)
					
					# Bias statistics
					col1, col2 = st.columns(2)
					with col1:
						total_candidates = df_bias["count"].sum()
						biased_count = df_bias[df_bias["bias_category"] == "Biased"]["count"].iloc[0] if len(df_bias[df_bias["bias_category"] == "Biased"]) > 0 else 0
						st.metric("Total Analyzed", total_candidates)
					with col2:
						bias_percentage = (biased_count / total_candidates * 100) if total_candidates > 0 else 0
						st.metric("Bias Percentage", f"{bias_percentage:.1f}%")
				else:
					st.info("📭 No bias distribution data available.")
			
			else:  # Flagged Candidates
				flagged_df = get_flagged_candidates(threshold=bias_threshold_pie)
				if not flagged_df.empty:
					st.markdown(f"**🚩 {len(flagged_df)} candidates flagged with bias score > {bias_threshold_pie}**")
					
					# Enhanced flagged candidates display
					display_df = flagged_df.copy()
					display_df = display_df.sort_values('bias_score', ascending=False)
					
					st.dataframe(
						display_df.style.format({'bias_score': '{:.3f}', 'ats_score': '{:.0f}'}),
						use_container_width=True,
						height=300
					)
					
					# Flagged candidates statistics
					col1, col2, col3 = st.columns(3)
					with col1:
						st.metric("Flagged Count", len(flagged_df))
					with col2:
						st.metric("Avg Bias Score", f"{flagged_df['bias_score'].mean():.3f}")
					with col3:
						st.metric("Avg ATS Score", f"{flagged_df['ats_score'].mean():.1f}")
				else:
					st.success("✅ No candidates flagged above the selected threshold.")
		except Exception as e:
			st.error(f"Error in bias analysis: {e}")

		# Enhanced Domain Performance Deep Dive
		with st.expander("🔍 Domain Performance Deep Dive", expanded=False):
			try:
				df_performance = get_domain_performance_stats()
				if not df_performance.empty:
					st.markdown("#### Comprehensive Domain Performance Metrics")
					
					# Performance heatmap
					performance_cols = ['avg_ats_score', 'avg_edu_score', 'avg_exp_score', 
								'avg_skills_score', 'avg_lang_score', 'avg_keyword_score']
					
					if all(col in df_performance.columns for col in performance_cols):
						heatmap_data = df_performance[['domain'] + performance_cols].set_index('domain')
						
						fig = px.imshow(heatmap_data.T, 
									title="Domain Performance Heatmap",
									color_continuous_scale="RdYlGn",
									aspect="auto")
						fig.update_layout(height=400)
						st.plotly_chart(fig, use_container_width=True)
					
					# Detailed performance table
					st.dataframe(
						df_performance.style.format({
							col: '{:.2f}' for col in performance_cols + ['avg_bias_score']
						}),
						use_container_width=True
					)
				else:
					st.info("ℹ️ No detailed performance data available.")
			except Exception as e:
				st.error(f"Error loading performance deep dive: {e}")

		# Footer with system information
		st.markdown("<hr style='border-top: 1px solid #ddd; margin: 2rem 0;'>", unsafe_allow_html=True)
		st.markdown("""
		<div style='text-align: center; color: #666; font-size: 0.9em; padding: 1rem;'>
			<p>🛡️ Enhanced Admin Dashboard | Powered by Advanced Database Manager</p>
			<p>Last updated: {}</p>
		</div>
		""".format(datetime.now().strftime("%Y-%m-%d %H:%M:%S")), unsafe_allow_html=True)
